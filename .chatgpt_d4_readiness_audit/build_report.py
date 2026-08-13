from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parents[1]
D4 = ROOT / 'Project 1 Data Quality Assessment' / 'D4 Parallel-redundancy Temporal Consistency'
FIG = D4 / 'outputs' / 'figures'
OUT = ROOT / 'deliverables'
OUT.mkdir(exist_ok=True)
DOCX = OUT / '260813_D4_顶刊SCI发表与最终DQR聚合_方法学缺口_四项问题审查_修改完善方案_专家版.docx'

FIGURES = [
    ('FigD4_2_pair_mechanism_profile.png', 'PR #19 FigD4_2：Pair-level mechanism profile'),
    ('FigD4_3_burden_coverage_calibration.png', 'PR #19 FigD4_3：Burden–coverage–calibration provenance'),
    ('FigD4_4_formal_episode_cases.png', 'PR #19 FigD4_4：Formal episode cases'),
    ('FigD4_5_mechanism_specificity.png', 'PR #19 FigD4_5：Mechanism specificity'),
    ('FigD4_6_ablation_and_lag_resolution.png', 'PR #19 FigD4_6：Ablation and lag-resolution limits'),
]
for fn, _ in FIGURES:
    if not (FIG / fn).exists():
        raise FileNotFoundError(FIG / fn)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def borders(cell, color='D9E2F3', sz='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = 'w:' + edge
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag); tcBorders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),sz); el.set(qn('w:color'),color)


def set_cell_text(cell, text, bold=False, color=None, size=8.7):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold; r.font.size = Pt(size); r.font.name = 'Arial'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    borders(cell)


def add_table(doc, headers, rows, widths=None, font_size=8.2):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for j,h in enumerate(headers):
        set_cell_text(t.rows[0].cells[j], h, True, 'FFFFFF', font_size)
        shade(t.rows[0].cells[j], '17365D')
        if widths: t.rows[0].cells[j].width = Cm(widths[j])
    for i,row in enumerate(rows):
        cells = t.add_row().cells
        for j,val in enumerate(row):
            set_cell_text(cells[j], val, False, None, font_size)
            if widths: cells[j].width = Cm(widths[j])
            if i % 2: shade(cells[j], 'F7F9FC')
    doc.add_paragraph()
    return t


def add_bullets(doc, items, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    for item in items:
        p = doc.add_paragraph(style=style); p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number'); p.add_run(item)


def add_callout(doc, title, text, fill='EAF2F8', border='B4C6E7'):
    t = doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,fill); borders(c,border,'6'); c.text=''
    p=c.paragraphs[0]
    r=p.add_run(title+'：'); r.bold=True; r.font.name='Arial'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
    r2=p.add_run(text); r2.font.name='Arial'; r2._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
    doc.add_paragraph()


def add_figure(doc, filename, caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG/filename), width=Cm(16.6))
    cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=cp.add_run(caption+'（仓库原图，基于 PR #19 head 3c58c8f）'); r.bold=True; r.font.size=Pt(8.2); r.font.color.rgb=RGBColor(79,129,189)


def add_source(doc, text):
    p=doc.add_paragraph()
    r=p.add_run('证据来源：'); r.bold=True; r.font.color.rgb=RGBColor(31,78,121)
    p.add_run(text)


def add_code(doc, text):
    p=doc.add_paragraph(style='CodeBlock')
    p.add_run(text)


doc = Document()
sec=doc.sections[0]
sec.top_margin=Cm(1.65); sec.bottom_margin=Cm(1.55); sec.left_margin=Cm(1.85); sec.right_margin=Cm(1.75)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(10.1); normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
normal.paragraph_format.line_spacing=1.28; normal.paragraph_format.space_after=Pt(4.5)
for sty,size,color in [('Title',24,'1F4E78'),('Heading 1',16,'17365D'),('Heading 2',13,'2F5597'),('Heading 3',11,'1F7A8C')]:
    s=styles[sty]; s.font.name='Arial'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
if 'CodeBlock' not in styles:
    cb=styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
else:
    cb=styles['CodeBlock']
cb.font.name='Consolas'; cb.font.size=Pt(8.5); cb.paragraph_format.left_indent=Cm(0.6); cb.paragraph_format.right_indent=Cm(0.4); cb.paragraph_format.space_after=Pt(5)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(78)
r=p.add_run('北岸厂 Class C-minDQR'); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor(23,54,93); r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('D4 平行冗余时序一致性\n顶刊 SCI 发表与最终 DQR 聚合方法学审查'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor(31,78,121); r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('PR #19 对照｜四项核心问题逐项核验｜统计修订方案｜聚合重跑条件｜发表前 P0/P1 路线图'); r.font.size=Pt(11.5); r.font.color.rgb=RGBColor(79,129,189); r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(110)
p.add_run('审查对象：GitHub main + PR #19\nmain 基线：60bd6783c257117ca70c2f923f1784db0b2e65c0\nPR #19 head：3c58c8fa0565a8dcbb24c1ac5b718c0389d0b9f2\n审查日期：2026-08-13')
doc.add_page_break()

# Executive

doc.add_heading('执行摘要', 1)
doc.add_paragraph('本审查的核心结论是：D4 已完成“回顾性科学评分与既定 v2.0 聚合”的工程闭环，但距顶刊 SCI 论文所要求的严格方法学证据链仍存在数个实质性缺口。PR #19 对图件体系进行了明显升级——由旧 8 图重构为 6 幅主图 + 3 幅补充图，加入 calibration provenance、正式 episode case、机制特异性、paired-window cluster bootstrap ablation 和源数据工作簿——但 PR #19 明确不重算、不改写 production D4_raw，也没有改动核心 scoring / calibration / validation 逻辑。因此，图件改进不能替代对评分内核的再审查。')

add_table(doc, ['审查问题','是否存在','严重性','专家结论'], [
    ('1. W1 与 KS 联合使用的增量价值尚未被直接消融证明','是','P0-Methods','当前消融删除的是整个 Q_dist；PR #19 仍只有 no_dist，而没有 W1-only、KS-only、full(W1+KS) 的成对比较。0.60/0.40 组合的必要性尚未由当前验证直接证明。'),
    ('2. 后段注入与 calibration thresholds 未完全隔离','是','P0-Publication','validation 取研究期后 30% 窗口，但 quantile thresholds 来自已用全研究期高质量窗口拟合的 D4_mapping_params。故属于“后段注入 stress test”，不是严格 untouched holdout。'),
    ('3. DO14 少数但更长 episode 的推断尚未正式统计验证','是（推断层面）','P1 / 若写入主结论则升 P0','现有汇总显示 DO14 低分率高而 ≥3 h 正式事件数仅 79，支持“持续型事件”假设；PR #19 只展示一个最长 formal case，未比较 episode-duration distribution 与 block-bootstrap CI。'),
    ('4. W1/KS 在 target/reference 各自 finite support 上计算，而非 common support','是','P0-Scoring','scoring.py 对两侧分别 drop NaN 后比较；D2 bilateral gate 只能保证两侧各自可观察，不能消除非对称缺失时段造成的人为分布差异。common-support 应成为 primary analysis。'),
], [5.3,2.0,2.6,7.0], 7.8)

add_callout(doc,'总审查判断','如果目标只是维持当前 v2.0 retrospective scientific aggregation，仓库合同已经将 D4_raw 判定为 score-ready；但如果要把 common-support 修正、development-only calibration 和 W1/KS 构件验证纳入最终论文主方法，则必须形成新的 D4 版本/Calibration ID，并重新生成 D4_raw 后再跑最终 composite。不能一边修改评分内核，一边继续把旧 D4V14 聚合结果称为最终主分析。','FFF2CC','D6B656')

# Scope

doc.add_heading('1 审查范围、证据层级与当前项目状态',1)
doc.add_heading('1.1 GitHub main 与 PR #19 的边界',2)
doc.add_paragraph('main 的 D4 expert audit 将当前 D4 定义为 homologous-pair temporal consistency：输入是 Section 1.1 de-periodised residuals；D2 负责可评价性；D1 仅解释；D5 只提供独立 report/gate interface；聚合数值只能使用 D4_raw。当前 run 为 D4V14_20260726_100717，42,847 pair-hours，37,987 finalizable，D4_forDQR 与 D4_raw 最大数值调整为 0。')
add_source(doc,'D4/docs/D4_EXPERT_AUDIT_2026-07.md；FINAL_FIVE_DIMENSION_AGGREGATION_CONTRACT.md。')
doc.add_paragraph('PR #19 的 changed-file scope 集中在 README、figure expert review、figure source-data workbooks、图件、figure QA、figure scripts 和 figure tests。PR 描述也明确 production D4_raw 不重算、不改写。换言之，PR #19 是 publication-figure revision，不是 scoring-core revision。')
add_source(doc,'PR #19 “Revise D4 figures for publication”，head 3c58c8f。')

add_table(doc,['状态层','当前结论','是否仍需工作'],[
('回顾性 D4 数值评分','已完成；D4_raw 为受保护独立维度','若不改核心算法，可直接用于现有 retrospective composite。'),
('PR #19 论文图件层','明显升级；6 main + 3 supplementary，source workbook + QA 完整','图件本身已接近高水平，但不能修复评分内核的统计偏差。'),
('严格内部验证','未完全成立','需要 development-only fit + frozen validation protocol。'),
('untouched terminal test','未完成','当前研究期已被反复查看/用于开发，不得事后重新命名为 untouched terminal。'),
('现场因果真值','不足','需要 maintenance/operator/process records；当前 D4 只能判 pair asymmetry。'),
('跨厂/未来期迁移','未完成','若无第二厂，至少需要真正未见的未来数据段。'),
('最终 DQR 聚合','回顾性 v2.0 已执行','若接受本报告 P0 内核修改，必须新版本重跑。'),
], [3.2,7.5,4.2])

# Issue 1

doc.add_heading('2 问题 1：Q_dist 中 W1 与 KS 联合使用的增量价值是否已证明？',1)
add_callout(doc,'结论','问题真实存在。当前证据只能证明“Distribution component 整体有价值”，不能证明“0.60×W1_norm + 0.40×KS 比 W1-only 或 KS-only 更优”。','FCE4D6','C65911')
doc.add_heading('2.1 当前实现事实',2)
add_code(doc,'risk_dist = 0.60 * (d_w1 / pooled_scale) + 0.40 * d_ks')
doc.add_paragraph('main/scoring.py 明确以 W1/pooled_scale 和 KS statistic 线性混合构造 risk_dist。当前 validation.py 的消融条件只有 full、no_dist、no_trend、no_var、no_cp、no_deadband；no_dist 会将整个 Q_dist 删除。PR #19 新的机制特异性消融也沿用 no_dist，而没有 W1-only / KS-only。')
add_source(doc,'D4/src/d4/scoring.py；D4/src/d4/validation.py；PR #19 src/d4/figures.py。')

doc.add_heading('2.2 为什么这在顶刊审稿中会被追问',2)
doc.add_paragraph('W1 和 KS 对“分布差异”的感知并不等价：KS 是两条 ECDF 的最大垂直差，强调局部最大 CDF discrepancy；一维 W1 可理解为 ECDF 间面积/quantile displacement，能保留整体位移幅度信息。文献中已反复指出，不同 distribution distances 对 location、shape、non-overlap 等替代情景具有不同灵敏度。因此若采用固定 0.60/0.40 混合，至少需要证明两者不是纯冗余拼接。')

doc.add_heading('2.3 必须新增的 confirmatory ablation',2)
add_table(doc,['模型','定义','用途'],[
('Full','0.60 W1_norm + 0.40 KS','现行主模型。'),
('W1-only','W1_norm；必须重新在 development 拟合其 own q50/q75/q90/q97.5','检验 KS 的增量贡献。'),
('KS-only','KS；必须重新在 development 拟合 own quantiles','检验 W1 的增量贡献。'),
('Optional equal-weight sensitivity','0.50/0.50；仅 development/internal validation','检验 0.60/0.40 权重是否高度敏感；不得用 terminal 调权。'),
], [3.2,8.5,3.5])
add_bullets(doc,[
'在完全相同的 window_id 上做 paired comparison，不能分别抽样。',
'主要性能：AUROC + AUPRC；机制至少分 drift、step、freeze，并增加专门的 distribution-shape challenges（location shift、scale inflation、tail contamination、mixture/bimodality），否则不能真正区分 W1 与 KS 的结构贡献。',
'负对照：equal common-mode process change 的 conditional new FAR；不能只比较正类 AUC。',
'不确定性：以 window_id 为 cluster 做 paired cluster bootstrap；输出 ΔAUROC_full-vs-W1、ΔAUROC_full-vs-KS、ΔAUPRC、ΔFAR 及 95% CI。',
'模型选择只能在 internal validation 做一次：若 Full 对至少一个机制有稳定增益且对 common-mode FAR 不劣，保留 Full；若无可重复增益，应优先更简洁的单一指标。',
'如果仍保留 0.60/0.40，需将其标为 development-frozen design choice，而不是事后根据 terminal performance 优化出的权重。'
])
add_figure(doc,'FigD4_6_ablation_and_lag_resolution.png','图 1  PR #19 已将 component-level ablation 做到机制特异性，但仍未拆开 W1 与 KS，故不能回答 Distribution 内部的增量价值问题。')

# Issue 2

doc.add_heading('3 问题 2：后段注入是否存在 calibration leakage？是否需要 development / internal validation / terminal test？',1)
add_callout(doc,'结论','存在实质性的“评估隔离不足”。当前 validation 窗口虽来自时间后段，但 quantile thresholds 由全期高质量 benchmark 形成，因此不能称为严格 out-of-sample terminal validation。建议升级为时间顺序、参数冻结、带 embargo 的三阶段协议；且当前已被反复分析的 2025-08–2026-04 数据不能事后重新命名为 untouched terminal test。','FCE4D6','C65911')

doc.add_heading('3.1 代码证据',2)
add_code(doc,"candidates = main[(main['usable_for_D4']) & (main['timestamp'] >= main['timestamp'].quantile(0.70))]")
doc.add_paragraph('这说明 injection challenge 的确被放在较后段窗口；但 _score_window() 使用的 thresholds 直接来自 D4_mapping_params.xlsx，而这些 public quantiles 在 pipeline 的 high-quality benchmark 上拟合。当前 pipeline 没有 phase filter 把 benchmark calibration 限制到 development。')
add_source(doc,'D4/src/d4/validation.py；D4/src/d4/pipeline.py。')

doc.add_heading('3.2 哪些对象必须只在 development 拟合/确定',2)
add_bullets(doc,[
'Section 1.1 decomposition 中所有需要估计或选择的参数：harmonic order、ARMA/GARCH order、whitening 参数、任何数据驱动频率/阈值选择。后段只能 apply frozen transform。',
'D1 regime templates：cluster number、feature preprocessing、cluster centers/templates；internal validation / terminal 只能 assign 到已冻结 regime，不得重新 k-means。',
'D4 benchmark eligibility rule：D1 minimum、D2 continuity、common-support minimum、regime stability rule、min stratum support 和 fallback hierarchy。',
'D4 q50/q75/q90/q97.5 for Q_dist/Q_trend/Q_var，包括 W1-only / KS-only 若参加模型比较。',
'D4 fixed rule：deadband、Q_cp timing table、aggregation weights、lambda_blend、D4<3 与 ≥3 h event rule。',
'任何候选结构选择（例如 Full vs W1-only vs KS-only）最多在 internal validation 选择一次；terminal 不得再调。'
])

doc.add_heading('3.3 推荐的三阶段合同',2)
add_table(doc,['阶段','允许做什么','禁止做什么','论文定位'],[
('Development','拟合 decomposition、regime、benchmark library、quantile mapping；建立候选模型；确定 QA 与 acceptance rule','不得报告其性能作为最终泛化性能','method development / calibration'),
('Internal validation','一次性比较已冻结候选；允许按预注册规则选最终结构','不得反复调阈值、权重、分位点；不得把多轮试验中的最好结果当 unbiased performance','model selection / internal confirmation'),
('Terminal test','只运行最终冻结 pipeline；一次性输出所有主要指标和 CI','禁止任何参数、规则、图形筛选后的返调；禁止按 terminal 结果改模型','final internal generalization estimate'),
], [3.0,5.5,6.0,3.0], 7.7)

doc.add_paragraph('由于当前整个 2025-08-01–2026-04-13 期间已经被 D1/D4 开发、审计、图件和 confirmatory analyses 多次使用，本报告不建议在该期间内“重新切一段”然后称其为真正 untouched terminal test。仓库 FINAL_FIVE_DIMENSION_READINESS_AUDIT 已明确指出：已在方法开发中查看过的数据段不能重新标记为 untouched test。')
add_callout(doc,'推荐实践','将现有研究期作为 development + internal evidence；在本次方法冻结之后，预先声明一个真正未来时间段作为 terminal test。若暂时无法获得新未来数据，应在稿件中诚实写为 single-plant retrospective/internal validation study，而不是制造伪独立 terminal。','E2F0D9','70AD47')

doc.add_heading('3.4 如果必须为下一轮新数据预注册，建议的时间隔离规则',2)
add_bullets(doc,[
'时间顺序切分，不随机打散；terminal 应是最新且从未用于任何规则选择的连续时段。',
'在 development→validation、validation→terminal 边界加入至少 7 d burn-in/embargo，因为 D4 change-point auxiliary window 为 7 d；边界后的前 7 d 可用于 causal history warm-up，但不计入独立评价指标。',
'所有 24 h rolling windows 必须有 phase provenance；禁止一个 scored window 同时跨越两个 phase 的拟合边界。',
'冻结 calibration manifest：fit_start、fit_end、input SHA256、regime template ID、D4 calibration ID、code commit、config commit。',
'终端评价前先锁定 primary endpoints、subgroup endpoints、multiplicity policy、CI method 和 fail/pass rule。'
])

# Issue 3

doc.add_heading('4 问题 3：DO14“少数但持续更长的 asymmetry episodes”是否需要正式验证？',1)
add_callout(doc,'结论','需要。如果该表述只作为探索性解释，可以保留为 hypothesis；如果要写入 Results/Discussion 的主要生物池空间结论，必须用正式 episode-duration distribution 与依赖结构感知的 CI 证明。','FFF2CC','D6B656')
doc.add_paragraph('当前 main expert audit 显示：DO14 mean D4_raw=2.978、low-score rate=47.8%、evaluable=98.8%，但 ≥3 h formal events 仅 79 个。相比 DO13 低分率 40.6% 却有 162 个事件，这一组合与“DO14 更偏持续型而非频繁短脉冲型 asymmetry”一致，但事件数和低分率不能单独证明 duration distribution 显著不同。PR #19 仅从 D4_event_windows 中为 DO14/ORP13 选取最长 formal case 做示例。')
add_source(doc,'D4_EXPERT_AUDIT_2026-07.md；PR #19 figure_4_field_cases()。')
add_figure(doc,'FigD4_4_formal_episode_cases.png','图 2  PR #19 的正式 episode case 图提高了可解释性，但 case study 不能替代 duration-distribution inference。')

doc.add_heading('4.1 推荐的 episode-level 统计输出',2)
add_table(doc,['指标','推荐统计量','目的'],[
('持续时间','median / IQR / P90 / P95 / restricted mean duration','刻画典型和长尾事件。'),
('事件发生频率','events per 1000 evaluable pair-hours','区分“高暴露来自频繁短事件”还是“少量长事件”。'),
('时间负担','fraction of evaluable hours with D4_raw<3','保留现有 burden。'),
('事件深度','min D4_raw、mean D4_raw、area below 3','增加严重度，不把时长和深度混为一谈。'),
('复合事件负荷','duration × mean deficit 或 area-below-threshold','用于 DO14 vs DO11–13 的稳健比较。'),
], [3.0,7.0,6.0])

doc.add_heading('4.2 CI 与检验方法',2)
add_bullets(doc,[
'不要把 1 h rolling windows 当 i.i.d.。当前 window 24 h、step 1 h，相邻窗口高度重叠。',
'优先在原始 hourly status sequence 上做 moving-block bootstrap，并在每个 bootstrap replicate 内重新运行 event extraction，而不是只对已经识别的 episode 简单 i.i.d. bootstrap。',
'为保持跨项目一致性，可用 7 d block 作为 primary，并用 48 h 与 14 d 做 sensitivity；若要完全数据驱动，可预注册 Politis–White block-length estimator。',
'输出 DO14 与 DO11–13 的 paired contrasts：Δmedian duration、duration ratio、Δevent rate、Δtime burden，并给 95% block-bootstrap CI。',
'图形建议：panel a 为 duration ECDF/CCDF，panel b 为 median duration forest + CI，panel c 为 event rate 与 time burden 二维定位；DO14 若位于“低 event rate + 高 burden/长 duration”象限，才形成正式证据。'
])

# Issue 4

doc.add_heading('5 问题 4：W1/KS 是否应改为 common-support primary analysis？',1)
add_callout(doc,'结论','应当修改，而且这是本报告最明确的 scoring-correctness P0。当前分别删除 NaN 后比较会允许两侧在不同时间段取样，D2 bilateral observability gate 不能完全修复这种 support mismatch。','FCE4D6','C65911')

doc.add_heading('5.1 当前代码事实',2)
add_code(doc,'finite_target = np.isfinite(target)\nfinite_reference = np.isfinite(reference)\nt = target[finite_target]\nr = reference[finite_reference]\nd_w1 = wasserstein_distance(t, r)\nd_ks = ks_2samp(t, r).statistic')
doc.add_paragraph('只要 target 和 reference 各自的 valid_fraction 均 ≥0.80，即可通过 data_ok。若 target 缺失集中在窗口前段、reference 缺失集中在后段，两侧仍可各有 80% 数据，但 W1/KS 比较的是不同时间支持；真实过程有趋势/工况切换时，这种不对称 support 本身就可能生成 distribution difference。')
add_source(doc,'D4/src/d4/scoring.py；D4/src/d4/pipeline.py。')

doc.add_heading('5.2 建议的 primary implementation',2)
add_code(doc,"common = np.isfinite(target) & np.isfinite(reference)\ncommon_fraction = common.mean()\nt = target[common]\nr = reference[common]\n# primary W1/KS are computed only on synchronous timestamps")
add_bullets(doc,[
'新增 n_common、valid_fraction_common、asymmetric_missing_fraction、support_jaccard 等 provenance 字段。',
'将 common_fraction≥0.80 作为 Distribution/Variability primary 的直接可评价条件；原“bilateral each≥0.80”保留为 D2/availability contextual evidence。',
'Q_dist 必须在 common-support risk distribution 上重新做 development-only q50/q75/q90/q97.5 calibration；不能直接沿用旧 quantiles。',
'Q_var 建议也在 common support 上计算 IQR，以维持同一时间支持。Q_trend 至少应在 hourly common-support medians 上估计，或者增加 common-hour coverage gate；否则 trend 仍可能由不同分钟支持形成。',
'Q_cp 可以继续允许单侧变点作为证据，但 adjacent segments 应报告 bilateral common coverage，以区分“真实单侧变化”和“单侧缺数导致的检出差异”。'
])

doc.add_heading('5.3 必须做的旧版 vs common-support sensitivity audit',2)
add_table(doc,['层级','比较指标'],[
('pair-hour score','median |ΔD4|、P95 |ΔD4|、Spearman、raw-status agreement、not-evaluable transition rate'),
('pair profile','mean/median D4、low-score burden、evaluable rate 的变化'),
('events','event count、duration distribution、event-window Jaccard、onset/offset displacement'),
('validation','drift/step/freeze AUROC/AUPRC、equal common-mode FAR、direction accuracy'),
('aggregation','formal pair-row count、Full/Basic coverage、pair composite delta、7 d block-bootstrap CI overlap'),
], [3.2,13.5])
add_callout(doc,'版本治理','common-support 若被采纳为 primary，就属于 scoring semantics change，应发布新 config_version + calibration_id + run_id；旧 D4V14 作为 legacy sensitivity 留档，不宜静默覆盖。','E2F0D9','70AD47')

# Additional gaps

doc.add_heading('6 除四项问题外，距顶刊论文仍需完善的关键工作',1)
add_table(doc,['优先级','工作项','当前证据','建议完成标准'],[
('P0','严格 phase freezing / terminal test','当前只有后段 injection；全期 quantile calibration','新未来期一次性 terminal；参数 SHA/ID 全冻结。'),
('P0','common-support scoring','当前独立 finite support','新版本 calibration + 全量 rerun + old/new audit。'),
('P0','W1–KS internal construct ablation','只有 no_dist','W1-only/KS-only/Full paired cluster-bootstrap；不得 terminal 调权。'),
('P0-aggregation','采用新 D4 后重跑 composite','现有 composite 基于旧 D4V14','新 D4 lineage 写入 aggregation manifest；重新算 pair score/coverage/uncertainty。'),
('P1','ORP exact-stratum support','expert audit 明确存在 fallback','逐 regime 报 exact_n、independent blocks、fallback share；增加未来高质量 ORP evidence 或预注册 shrinkage。'),
('P1','DO14 duration inference','只有汇总与 case','episode-duration distribution + block-bootstrap CI。'),
('P1','field truth / causal adjudication','D4 只能判 pair asymmetry','maintenance/operator/process records；至少双专家 adjudication。'),
('P1','3 h event rule validation','当前固定 ≥3 h','独立现场事件或预注册 sensitivity，报告 event Jaccard/precision burden。'),
('P1','change-point timing','0–180 min severity rho=0','保持 sub-hour claim 禁用；若需要亚小时结论，另建更高频 CP 设计。'),
('P1','overlap-aware effective support','100 windows 实际高度重叠','报告 unique days/blocks 或 effective sample support，不将 rolling rows 当独立 n。'),
('P2','second-plant / second-line transfer','尚无外部站点','优先 second plant；若无则明确 single-plant scope。'),
], [2.1,4.3,5.4,5.0], 7.2)

add_figure(doc,'FigD4_3_burden_coverage_calibration.png','图 3  PR #19 已把 evaluability 与 calibration provenance 画进主图，这是正确方向；但 provenance 可视化并不能替代 development-only calibration 与 independent-block support。')
add_figure(doc,'FigD4_5_mechanism_specificity.png','图 4  PR #19 已增强 drift/step/freeze、common-mode 与 pair heterogeneity 表达；但其 underlying thresholds 仍来自当前全期 calibration。')

# Aggregation

doc.add_heading('7 对“最终数据质量评分聚合”的具体影响',1)
doc.add_heading('7.1 当前合同下：为什么 D4 已经可以用于 retrospective aggregation',2)
doc.add_paragraph('FINAL_FIVE_DIMENSION_AGGREGATION_CONTRACT 已明确：D4 是 score-ready through D4_raw；D1 解释、D2 可评价、D5 report/gate 均不能重写 D4_raw。现有 composite 规则为：node score = eligible D1/D2/D5 等权均值；pair score = target-node score、reference-node score 与 D4_raw 三者等权均值，且三者必须全部 evaluable；D3 以非补偿 Safety Gate 独立存在。现有 v2.0 retrospective composite 已执行。')
add_source(doc,'FINAL_FIVE_DIMENSION_AGGREGATION_CONTRACT.md；FINAL_FIVE_DIMENSION_READINESS_AUDIT_2026-07.md。')

doc.add_heading('7.2 但若本报告 P0 修正被采纳，为什么必须重新聚合',2)
add_bullets(doc,[
'common-support 会改变 risk_dist 的数据支持，并要求重新拟合 quantile thresholds，因此 D4_raw 的数值语义发生变化。',
'development-only calibration 会改变 q50/q75/q90/q97.5，即使 risk metric 不变，Q_dist/Q_trend/Q_var 仍可能变化。',
'W1/KS 构件选择若导致 Q_dist 公式调整，同样属于核心 scoring version change。',
'聚合的 pair score 直接包含 D4_raw，所以新旧 D4 不能混在同一个“最终”结果中；必须重新生成 formal pair rows、Full/Basic coverage 和 composite CI。'
])

doc.add_heading('7.3 最终 aggregation release 必须新增的 lineage 字段',2)
add_table(doc,['字段','要求'],[
('d4_config_version','明确 common-support / phase-freeze 版本。'),
('d4_calibration_id','development-only quantile mapping 的 SHA-bound ID。'),
('fit_period','development start/end，不得为空。'),
('validation_period / terminal_period','作为只读评价范围写入 manifest。'),
('common_support_contract','min_common_fraction、common-hour rule、NaN policy。'),
('distribution_component_version','Full/W1-only/KS-only 中最终冻结选择及权重。'),
('benchmark_provenance','variable×regime exact_n、independent block count、fallback scope、calibration_quality。'),
('aggregation_source_sha','新 D4_main_scores / integration file 的 SHA256。'),
], [4.6,12.0])

add_callout(doc,'聚合准入判定','建议把“当前 retrospective aggregation complete”与“论文最终主分析 aggregation locked”分成两个 release 标签。前者保留历史可重复性；后者必须等 common-support、development-only calibration、W1/KS construct decision 冻结后再生成。','EAF2F8','5B9BD5')

# Implementation plan

doc.add_heading('8 文件级修改完善方案',1)
add_table(doc,['文件/模块','必须修改内容'],[
('configs/d4.yaml','新增 phase contract、fit period、terminal policy、common_support_min_fraction、distribution_component_version、固定权重来源与版本。'),
('src/d4/scoring.py','common-support W1/KS primary；同步输出 n_common/common_fraction/asymmetric_missing_fraction；Q_var/Q_trend 同步 support policy。'),
('src/d4/pipeline.py','benchmark 只允许 development rows；quantile fit 添加 fit_phase/provenance；后段只 apply frozen mapping。'),
('src/d4/validation.py','W1-only/KS-only/Full；paired cluster-bootstrap incremental value；phase guard；terminal read-only enforcement。'),
('src/d4/episode_validation.py（建议新增）','formal event duration、event rate、severity、moving-block bootstrap contrasts。'),
('outputs/data/D4_mapping_params.xlsx','增加 fit_start/end、exact_n、independent_blocks、mapping_scope、calibration_quality、component_version。'),
('outputs/data/D4_main_scores.xlsx','增加 phase_id、common support provenance、mapping_scope/calibration_quality 行级可追踪字段。'),
('outputs/data/D4_benchmark_results.xlsx','增加 distribution_internal_ablation、internal_validation、terminal_metrics、common_support_sensitivity。'),
('outputs/data/D4_event_duration_validation.xlsx','增加 pair-level duration distribution、contrasts、CI、block sensitivity。'),
('outputs/figures','PR #19 图组保留；在 FigD4_6 或 Supplement 增加 W1-vs-KS internal ablation；DO14 duration 若为主要结论应增加专门 panel。'),
('tests/','增加 phase leakage、common-support NaN pattern、terminal no-fit-access、mapping lineage、event-bootstrap reproduction、aggregation schema tests。'),
], [5.5,11.0], 7.2)

# Tests

doc.add_heading('9 建议新增的强制回归测试（publication gates）',1)
add_numbers(doc,[
'`test_no_validation_or_terminal_rows_in_calibration_library`：任何 phase != development 的 timestamp 出现在 benchmark/calibration 即 FAIL。',
'`test_terminal_has_zero_fit_permissions`：terminal pipeline 只能读取 frozen artifacts，任何 fit/quantile/template update 调用即 FAIL。',
'`test_common_support_invariance_to_asymmetric_nan_positions`：在保持共同有效值不变时，仅改变单侧额外 NaN 位置，不应改变 primary W1/KS。',
'`test_distribution_internal_ablation_is_paired`：Full/W1-only/KS-only 必须共享相同 window_id 和 label。',
'`test_mapping_id_binds_fit_period_and_component_version`：calibration ID 必须对 fit period / component formula 改动敏感。',
'`test_event_duration_bootstrap_reextracts_events`：bootstrap replicate 必须重新抽取 episode，而不是直接 resample event table。',
'`test_new_d4_reaggregation_lineage`：最终 composite 中的 D4 SHA / calibration ID 必须与新 D4 release manifest 完全一致。'
])

# Publication claim matrix

doc.add_heading('10 顶刊论文中哪些结论现在可以写、哪些必须等完成后再写',1)
add_table(doc,['结论','当前是否可写','建议表述'],[
('D4 measures homologous-pair temporal asymmetry','可以','明确 detection ≠ sensor causality。'),
('PR #19 figures are publication-grade and traceable','可以','说明 6 main + 3 supplementary、source workbooks、SHA QA。'),
('Distribution component contributes materially','可以有限写','现有 no_dist ablation 支持整体 Q_dist 价值。'),
('W1 + KS combination is superior to either alone','不可以','完成 W1-only/KS-only/Full paired confirmatory ablation 后再写。'),
('Current late-window injection is unbiased terminal validation','不可以','只能称 internal chronological stress test / injection validation。'),
('DO14 has significantly longer asymmetry episodes','不可以作为确定性结论','当前只能写 hypothesis/observed pattern；需 duration distribution + block CI。'),
('D4 scores are robust to asymmetric missingness','不可以','完成 common-support primary rerun 与 sensitivity audit 后再写。'),
('D4 is aggregation-ready for retrospective composite','可以','仅限 current D4V14 retrospective contract。'),
('Final paper composite is locked after core revisions','尚不可','新 D4 version rerun composite 后方可。'),
('D4 identifies sensor faults','不可以','需要 D5/localization + external truth；当前只能 pair asymmetry。'),
], [7.2,3.0,6.5], 7.3)

# Priorities

doc.add_heading('11 推荐执行顺序：最小闭环到顶刊主分析',1)
add_table(doc,['顺序','任务','是否需要新数据','完成后解锁'],[
('1','common-support scoring + development-only calibration 重构','否（可先用现有数据做方法重跑）','解决评分正确性与 leakage 结构问题。'),
('2','W1/KS internal construct ablation','否','证明或简化 Q_dist 组成。'),
('3','旧 D4V14 vs 新 D4 版本全量 sensitivity audit','否','决定是否需要实质重写 Results。'),
('4','DO14 episode-duration block-bootstrap','否','把空间持续性推断升级为正式结果。'),
('5','新 D4 release + 新 calibration ID + manifest','否','形成论文可冻结的 scoring core。'),
('6','重新执行 five-dimension composite + uncertainty','否','形成最终 aggregation 主分析。'),
('7','真正未来 terminal test','是：必须是未见未来期','提供无调参污染的 final internal generalization。'),
('8','maintenance/process truth + second plant（优先其一，最好两者）','是','解锁 causal/transportability stronger claims。'),
], [1.5,8.0,4.3,4.6], 7.4)

add_callout(doc,'最关键的取舍','目前不建议再继续“加更多复杂算法”。D4 距顶刊的主要差距已经不是模型复杂度，而是 protocol purity（严格冻结）、support correctness（共同时间支持）、construct necessity（W1/KS 增量）、dependence-aware inference（episode CI）与 external truth。把这几项做严谨，比继续增加新的 detector 更有论文价值。','D9EAD3','548235')

# References

doc.add_heading('12 可直接引用的 SCI / 方法学文献',1)
refs = [
'1. Spindler, A. (2014). Structural redundancy of data from wastewater treatment systems. Determination of individual balance equations. Water Research, 57, 193–201. DOI: 10.1016/j.watres.2014.03.042.',
'2. Villez, K., Vanrolleghem, P. A., & Corominas, L. (2016). Optimal flow sensor placement on wastewater treatment plants. Water Research, 101, 75–83. DOI: 10.1016/j.watres.2016.05.068.',
'3. Schneider, M. Y., Carbajal, J. P., Furrer, V., Sterkele, B., Maurer, M., & Villez, K. (2019). Beyond signal quality: The value of unmaintained pH, dissolved oxygen, and oxidation-reduction potential sensors for remote performance monitoring of on-site sequencing batch reactors. Water Research, 161, 639–651. DOI: 10.1016/j.watres.2019.06.007.',
'4. Künsch, H. R. (1989). The Jackknife and the Bootstrap for General Stationary Observations. The Annals of Statistics, 17(3), 1217–1241. DOI: 10.1214/aos/1176347265.',
'5. Politis, D. N., & White, H. (2004). Automatic Block-Length Selection for the Dependent Bootstrap. Econometric Reviews, 23(1), 53–70. DOI: 10.1081/ETC-120028836.',
'6. Cawley, G. C., & Talbot, N. L. C. (2010). On Over-fitting in Model Selection and Subsequent Selection Bias in Performance Evaluation. Journal of Machine Learning Research, 11, 2079–2107.',
'7. Ramdas, A., García Trillos, N., & Cuturi, M. (2017). On Wasserstein Two-Sample Testing and Related Families of Nonparametric Tests. Entropy, 19(2), 47. DOI: 10.3390/e19020047.',
'8. Schneider, M. Y. et al. (2019). Characterizing long-term wear and tear of ion-selective pH sensors. Water Science and Technology. DOI: 10.2166/wst.2019.301.',
]
for ref in refs:
    doc.add_paragraph(ref)

doc.add_heading('13 最终专家结论',1)
doc.add_paragraph('D4 当前不是“未完成项目”，而是已经完成了一个可重复的 retrospective score-ready release；PR #19 又显著提升了 manuscript figure architecture。但如果目标从“回顾性 DQR 聚合”提升到“顶刊主方法 + 最终锁定 composite”，则必须把本报告四项问题中第 1、2、4 项作为评分内核 P0，把第 3 项作为结论级统计确认，并在其后重新冻结 D4 release 与 aggregation lineage。特别是：不能用 PR #19 的图件升级来替代 development/validation/terminal 隔离，也不能把当前研究期中已经被方法开发使用过的后段数据重新命名为 untouched terminal test。最终最有价值的升级路线应是：common-support primary → development-only calibration → W1/KS incremental ablation → episode-duration inference → new D4 release → composite rerun → genuine future terminal test。')

# Save
for p in doc.paragraphs:
    for r in p.runs:
        if r.font.name is None:
            r.font.name='Arial'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')

doc.save(DOCX)
print(DOCX)
