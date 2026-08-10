from pathlib import Path
import hashlib, struct, subprocess
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
D4 = ROOT / 'Project 1 Data Quality Assessment' / 'D4 Parallel-redundancy Temporal Consistency'
FIG = D4 / 'outputs' / 'figures'
OUT = ROOT / 'deliverables' / '260810_D4_GitHub_outputs八组图件逐面板定量分析_准确性审查_SCI文献对照_仓库原图最终版.docx'
OUT.parent.mkdir(parents=True, exist_ok=True)

FIGS = [
('FigD4_1_paired_residual_consistency.png','FigD4_1 — Paired de-periodised residual consistency'),
('FigD4_2_subscore_contribution.png','FigD4_2 — Subscore contribution and pair-level D4 distribution'),
('FigD4_3_trend_slope_scatter.png','FigD4_3 — Robust trend-slope concordance'),
('FigD4_4_raw_score_heatmap.png','FigD4_4 — Weekly D4_raw score heatmap'),
('FigD4_5_status_evaluability.png','FigD4_5 — Pair status and evaluability'),
('FigD4_6_context_independence.png','FigD4_6 — Context independence and protection of D4_raw'),
('FigD4_7_validation_roc_pr.png','FigD4_7 — Mechanism validation by ROC and precision–recall analysis'),
('FigD4_8_validation_ablation.png','FigD4_8 — Component ablation and false-alarm control'),
]

PAIR = [
('DO11','3.368','31.4%','97.5%','124'),('DO12','3.391','26.5%','98.7%','128'),
('DO13','3.198','40.6%','98.8%','162'),('DO14','2.978','47.8%','98.8%','79'),
('ORP11','3.425','27.7%','92.3%','136'),('ORP12','3.040','49.2%','70.3%','215'),
('ORP13','2.650','67.6%','64.2%','164')]

SUB = [
('DO11','≈3.70','≈4.03','≈4.14','≈2.77','≈3.69','3.368'),
('DO12','≈4.11','≈4.02','≈3.96','≈2.62','≈3.76','3.391'),
('DO13','≈3.86','≈3.82','≈3.56','≈2.76','≈3.57','3.198'),
('DO14','≈2.90','≈4.20','≈4.01','≈2.43','≈3.35','2.978'),
('ORP11','≈3.85','≈4.06','≈3.95','≈3.14','≈3.78','3.425'),
('ORP12','≈2.90','≈3.82','≈2.71','≈3.53','≈3.22','3.040'),
('ORP13','≈2.73','≈3.21','≈2.56','≈2.84','≈2.84','2.650')]

REFS = [
"Spindler A. Structural redundancy of data from wastewater treatment systems. Determination of individual balance equations. Water Research. 2014;57:193–201. doi:10.1016/j.watres.2014.03.042.",
"Villez K, Vanrolleghem PA, Corominas L. Optimal flow sensor placement on wastewater treatment plants. Water Research. 2016;101:75–83. doi:10.1016/j.watres.2016.05.068.",
"Schneider MY, Carbajal JP, Furrer V, Sterkele B, Maurer M, Villez K. Beyond signal quality: The value of unmaintained pH, dissolved oxygen, and oxidation-reduction potential sensors for remote performance monitoring of on-site sequencing batch reactors. Water Research. 2019;161:639–651. doi:10.1016/j.watres.2019.06.007.",
"Ba-Alawi AH, Vilela P, Loy-Benitez J, Heo S, Yoo CK. Intelligent sensor validation for sustainable influent quality monitoring in wastewater treatment plants using stacked denoising autoencoders. Journal of Water Process Engineering. 2021;43:102206. doi:10.1016/j.jwpe.2021.102206.",
"Ba-Alawi AH, Loy-Benitez J, Kim SY, Yoo CK. Missing data imputation and sensor self-validation towards a sustainable operation of wastewater treatment plants via deep variational residual autoencoders. Chemosphere. 2022;288:132647. doi:10.1016/j.chemosphere.2021.132647.",
"Sen PK. Estimates of the Regression Coefficient Based on Kendall's Tau. Journal of the American Statistical Association. 1968;63(324):1379–1389. doi:10.1080/01621459.1968.10480934.",
"Künsch HR. The Jackknife and the Bootstrap for General Stationary Observations. The Annals of Statistics. 1989;17(3):1217–1241. doi:10.1214/aos/1176347265.",
"Killick R, Fearnhead P, Eckley IA. Optimal Detection of Changepoints With a Linear Computational Cost. Journal of the American Statistical Association. 2012;107(500):1590–1598. doi:10.1080/01621459.2012.737745.",
"Corominas L, Garrido-Baserba M, Villez K, Olsson G, Cortés U, Poch M. Transforming data into knowledge for improved wastewater treatment operation. Environmental Modelling & Software. 2018;106:89–103. doi:10.1016/j.envsoft.2017.11.023.",
"Newhart KB, Holloway RW, Hering AS, Cath TY. Data-driven performance analyses of wastewater treatment plants: A review. Water Research. 2019;157:498–513. doi:10.1016/j.watres.2019.03.030.",
]

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def git_blob_sha(path):
    data = path.read_bytes(); h=hashlib.sha1(); h.update(f'blob {len(data)}\0'.encode()); h.update(data); return h.hexdigest()

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def add_table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    hdr=t.rows[0]; set_repeat_table_header(hdr)
    for i,h in enumerate(headers):
        hdr.cells[i].text=h; set_cell_shading(hdr.cells[i],'17365D'); hdr.cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in hdr.cells[i].paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(8.5)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.size=Pt(8.3)
    return t

def add_p(doc, text, bold_lead=None):
    p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0.74); p.paragraph_format.line_spacing=1.32; p.paragraph_format.space_after=Pt(5)
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead); r.bold=True; p.add_run(text[len(bold_lead):])
    else: p.add_run(text)
    return p

def add_bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3); p.add_run(text); return p

def add_fig(doc, fname, caption):
    path=FIG/fname
    if not path.exists(): raise FileNotFoundError(path)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(16.2))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after=Pt(7)
    r=c.add_run(f'图 5-{fname.split("_")[1]}  GitHub 仓库原图：{caption}。规范文件：{fname}；Git blob SHA：{git_blob_sha(path)}；快照：881709ef0e8fe09605fad913d8d0d46dd6bc6870。')
    r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(79,129,189)


doc=Document()
sec=doc.sections[0]; sec.top_margin=Cm(1.7); sec.bottom_margin=Cm(1.6); sec.left_margin=Cm(1.9); sec.right_margin=Cm(1.8)
styles=doc.styles
styles['Normal'].font.name='Arial'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'等线'); styles['Normal'].font.size=Pt(10.5)
for s,sz,col in [('Title',26,'1F4E78'),('Heading 1',16,'17365D'),('Heading 2',13,'2F5597'),('Heading 3',11,'1F7A8C')]:
    st=styles[s]; st.font.name='Arial'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'等线'); st.font.size=Pt(sz); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(col)

# Header/footer
for section in doc.sections:
    hp=section.header.paragraphs[0]; hp.text='D4 Parallel-redundancy Temporal Consistency | GitHub repository-original figure review'; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(72)
r=p.add_run('北岸厂 Class C-minDQR'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor(23,54,93)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('D4 平行冗余时序一致性'); r.bold=True; r.font.size=Pt(27); r.font.color.rgb=RGBColor(31,78,121)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('GitHub outputs 八组图件逐面板定量分析、准确性审查与 SCI 文献对照'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(47,85,151)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run('仓库原图最终版｜直接嵌入 outputs/figures 原始 PNG｜快照 881709e').font.size=Pt(11)
doc.add_paragraph('\n')
add_table(doc,['项目','锁定信息'],[
('GitHub repository','muzi-1998/beian_tree'),('Commit','881709ef0e8fe09605fad913d8d0d46dd6bc6870'),('D4 run','D4V14_20260726_100717'),('Data span','2025-08-01 00:00 – 2026-04-13 23:50'),('Pair-hours','42,847 total; 37,987 finalized/evaluable D4 pair-hours'),('Formal numeric source','D4_raw'),('Figure bundle','8 figures; 19 panels; repository PNG/SVG/PDF all passed QA')])
doc.add_page_break()

# Executive summary
h=doc.add_heading('摘要与专家结论',1)
add_p(doc,'本报告直接采用 GitHub main 快照 881709ef0e8fe09605fad913d8d0d46dd6bc6870 中 D4 `outputs/figures` 的八幅规范原图，对 FigD4_1–FigD4_8 的 19 个面板开展逐面板定量分析、统计实现审查、构念边界审查以及 SCI 文献对照。图件不进行重绘；Word 中插入的均为仓库原始 PNG，图题同步记录规范文件名与 Git blob SHA，以建立“代码—源数据—图件—论文结论”的可追踪证据链。')
add_p(doc,'D4 的科学构念应严格定义为 homologous-pair temporal consistency，即两个平行生物处理线中同工艺位置、同变量传感器在去周期残差层面的时间一致性。D4_raw<3 表示 pair-level temporal inconsistency / parallel-line temporal asymmetry，而不是 sensor fault probability。当前 D4 已达到单厂回顾性数据质量子维度聚合的科学可用水平，但在 sensor-specific attribution、真实事件 PPV/FAR、跨厂外部有效性和严格时序外推方面仍不足以支撑“现场故障诊断器”式表述。')
add_p(doc,'七个 homologous pairs 呈现显著异质性：DO11、DO12 与 ORP11 相对稳定；DO14 在 98.8% 高可评价性的同时 mean D4_raw=2.978、47.8% 的可评价 pair-hours 低于 3，证明其主要问题不是信息不可获得；ORP13 最严重，mean D4_raw=2.650，且 67.6% 的可评价 pair-hours 低于 3，但其 evaluability 仅 64.2%，因此必须将低分负荷与 evidence coverage 联合报告。')
add_p(doc,'内部机制验证显示 D4 对持续 unilateral drift 与 step 的 AUROC 分别为 0.934 和 0.936，对 freeze 为 0.798，而 isolated spike 仅为 0.561；equal common-process perturbation 的 conditional new false-alarm rate 为 0.041。这说明 D4 的优势是识别持续、非对称的跨线变化并抑制同步共模变化，而不是替代 D1 的尖峰/冻结传感器故障检测。')

# Figure traceability
h=doc.add_heading('5 GitHub outputs 八组图件逐面板定量分析、准确性审查与 SCI 文献对照',1)
add_p(doc,'本章仅对仓库正式输出图件作科学解释。机械 QA 中的 passed 仅证明文件存在、非空、Arial 字体声明、面板标签规范以及与锁定源数据 lineage 的新鲜度一致；它不等同于算法阈值、现场因果归属或外部有效性已经通过科学验证。')
rows=[]
for fn,cap in FIGS: rows.append((fn, git_blob_sha(FIG/fn), cap))
add_table(doc,['仓库原图','Git blob SHA','论文角色'],rows)

# 5.1
h=doc.add_heading('5.1 FigD4_1 — Paired de-periodised residual consistency',2)
add_fig(doc,FIGS[0][0],FIGS[0][1])
add_p(doc,'统计对象与估计量。该图包含 7 个纵向面板，依次对应 DO11、DO12、DO13、DO14、ORP11、ORP12 和 ORP13。图中展示的不是原始 DO/ORP 量值，而是去周期化后的 residual trajectories；绘图层采用 6 h median aggregation，并以 pair-specific robust scale 标准化。因此该图用于判断同一 homologous pair 内两条处理线残差轨迹是否同步，不应跨 DO/ORP 直接比较绝对振幅。')
add_table(doc,['Pair','mean D4_raw','D4_raw<3','Evaluable','≥3 h events'],PAIR)
add_p(doc,'面板 (a)–(c)：DO11、DO12 与 DO13。DO11/DO12 的平均 D4_raw 分别为 3.368/3.391，低分率分别为 31.4%/26.5%，表明两个位置多数时间仍维持较好的跨线 residual synchrony。DO13 平均分降至 3.198，低分率升至 40.6%，并形成 162 个持续≥3 h 的事件，是前三个好氧 DO pair 中事件最密集者，提示靠后好氧位置存在更频繁的阶段性 inter-line divergence。')
add_p(doc,'面板 (d)：DO14。DO14 的 evaluability 为 98.8%，但 mean D4_raw=2.978、low-score burden=47.8%。这一组合具有最强的构念意义：低分发生在充分可观测条件下，不能用“后缺氧低 DO 导致不可评价”解释。DO14 仅有 79 个≥3 h events，却覆盖接近一半低分小时；与 DO13 的 162 个 events/40.6% low-score rate 对照，合理推断 DO14 更偏向少数但持续时间更长的 asymmetry episodes。该推断应在投稿前用 episode-duration distribution 与 block-bootstrap CI 正式验证。')
add_p(doc,'面板 (e)–(g)：ORP11、ORP12 与 ORP13。ORP11 mean D4_raw=3.425、low-score rate=27.7%，为 ORP 中相对稳定 pair。ORP12 降至 3.040，低分率 49.2%，同时 evaluability 仅 70.3%。ORP13 是全组最显著的不一致 pair：mean D4_raw=2.650，67.6% 的可评价小时低于 3，但 evaluability 只有 64.2%。因此论文必须表述为“67.6% among evaluable pair-hours”，禁止外推为全时间段的传感器故障比例。')
add_p(doc,'准确性审查。仓库 FigD4_1 中的红色/粉色 shading 是可视化提示层，并非逐 episode 精确标注的 event-window truth；此外 6 h median 会有意抑制 sub-6-h spikes。因此该图适合呈现 persistence 和 inter-line divergence，不适合用于证明 spike detection 或精确事件持续时间。')
add_p(doc,'SCI 对照。Spindler [1] 将 structural redundancy 引入 WWTP 数据验证，Villez 等 [2] 进一步强调 observability/redundancy 对 fault detection and identification 的基础价值。D4 的增量在于把这种冗余思想扩展到 homologous spatial redundancy 的动态残差层，而不是简单比较两个原始测点是否相等。Schneider 等 [3] 也提示 signal quality 与 retained process information 并非同一概念，支持 D4 与 D1 分离解释。')
add_p(doc,'可直接用于 SCI Results：The de-periodised residual trajectories revealed marked heterogeneity in homologous-pair consistency. DO11 and DO12 retained comparatively high temporal agreement, whereas DO14 exhibited a high evaluability of 98.8% but a mean D4 score below the consistency boundary (2.978), indicating that its low pair-consistency was not attributable to a lack of observable information. ORP13 represented the most persistent asymmetric pair, with 67.6% of evaluable pair-hours below D4=3.')

# 5.2
h=doc.add_heading('5.2 FigD4_2 — Subscore contribution and pair-level D4 distribution',2)
add_fig(doc,FIGS[1][0],FIGS[1][1])
add_p(doc,'面板 (a) 将四类 evidence 的平均加权贡献进行堆叠：distribution 0.35、trend 0.25、variability 0.20、change-point timing 0.20。重要的是，该堆叠对应 D4_base，而不是最终 D4_raw；最终分还包含 25% minimum-subscore penalty。因此柱状图总高不能直接当作 D4_raw。')
add_table(doc,['Pair','Q_dist','Q_trend','Q_var','Q_cp','D4_base','D4_raw'],SUB)
add_p(doc,'上述 Q 子分为根据仓库原始 SVG/PNG 的柱高与纵轴标度反演的近似均值，用于机制解释；mean D4_raw 为正式结果值。投稿时应从 D4_main_scores.xlsx 精确导出 Q 子分，不应以图像反演值替代源表。')
add_p(doc,'DO14 的 Q_trend≈4.20、Q_var≈4.01 均较高，而 Q_dist≈2.90、Q_cp≈2.43，说明其低 D4 主要由 residual distribution displacement 与 structural-change timing mismatch 驱动，而不是长期斜率或波动幅度普遍失配。ORP12 的主要弱项为 Q_dist≈2.90 与 Q_var≈2.71，更接近 distribution–variability asymmetry；ORP13 四项均下降，尤其 Q_dist≈2.73、Q_var≈2.56，属于 multi-component temporal divergence。')
add_p(doc,'面板 (b) 的箱线图只包含 usable_for_D4=True 的记录，因而应解释为 evaluable D4_raw distribution。ORP13 与 DO14 的总体分布向阈值以下移动，而 ORP12 存在较重 lower-tail burden。箱线图说明 pair-level score 具有明显低尾结构，资产排序不应只依赖全期 mean D4。')
add_p(doc,'准确性审查。distribution risk 结合 Wasserstein-1 距离与 two-sample KS statistic，trend 使用 hourly median 后的 Theil–Sen slope difference，variability 使用 log-IQR ratio 并带 process-floor deadband。一个需要在投稿前修订的细节是：当前实现对 target/reference 分别删除 NaN 后计算 W1/KS，而不是先限制在两侧共同有效的时间戳。虽然 D2 bilateral observability gate 缓解了风险，但无法严格排除 asymmetric temporal support 人为形成 distribution difference。建议将 common-support W1/KS 设为 primary analysis。')
add_p(doc,'可直接用于 SCI Results：Subscore decomposition demonstrated that low D4 values were mechanistically heterogeneous. The persistently low DO14 score was driven primarily by distributional and change-point-timing discrepancies, while its trend and variability subscores remained comparatively high. In contrast, ORP12 was characterized predominantly by distributional and variability asymmetry, whereas ORP13 exhibited concordant degradation across all four evidence components.')

# 5.3
h=doc.add_heading('5.3 FigD4_3 — Robust trend-slope concordance',2)
add_fig(doc,FIGS[2][0],FIGS[2][1])
add_p(doc,'面板 (a) 汇总所有 DO homologous pairs 的 target/reference Theil–Sen slopes，面板 (b) 对 ORP pairs 作相同处理。每个 24 h residual window 先聚合为 hourly medians，再以 Theil–Sen estimator 计算局部斜率；y=x 为完全趋势一致参照。')
add_p(doc,'DO11、DO12、DO13 与 DO14 的平均 Q_trend 约为 4.03、4.02、3.82 和 4.20。尤其是 DO14：尽管 overall D4 最低于所有 DO pairs，trend 子分却最高之一，这再次证明其问题不能笼统描述为“两条线趋势长期背离”。ORP11、ORP12、ORP13 的 Q_trend 约为 4.06、3.82 和 3.21，其中 ORP13 同时存在 trend、distribution 与 variability divergence，是最符合 multi-scale temporal divergence 的 ORP pair。')
add_p(doc,'统计边界。该图是 density visualization，不是相关性检验；仓库图没有给出 Pearson r、Spearman ρ、concordance correlation coefficient 或 Deming regression，因此不得从视觉散点直接写出相关系数。若进入主文，建议增加 pair-specific robust concordance effect size 和 month/block-bootstrap 95% CI。Theil–Sen 的稳健性可由 Sen [6] 支撑。')

# 5.4
h=doc.add_heading('5.4 FigD4_4 — Weekly D4_raw score heatmap',2)
add_fig(doc,FIGS[3][0],FIGS[3][1])
add_p(doc,'该图将 hourly D4_raw 聚合为 weekly mean，突出 pair asymmetry 的 persistence、temporal clustering 与 pair-specific hierarchy。DO11/DO12/ORP11 构成相对稳定层；DO13/DO14/ORP12 为中重度不一致层；ORP13 的全期 mean D4_raw=2.650 且 67.6% 的可评价小时低于 3，其周尺度低分带说明问题不是由少量孤立极值窗口造成。')
add_p(doc,'趋势解释。若低 D4 主要来自独立随机测量噪声，weekly aggregation 后应明显收敛；DO14、ORP12、ORP13 仍表现持续 lower-score bands，支持 persistent inter-line divergence 或 regime-dependent structural asymmetry 的解释。不过缺少维护记录、便携仪比对和独立 process-state adjudication 时，不能进一步将其归因为 sensor bias、hydraulic imbalance 或真实工艺非对称。')
add_p(doc,'统计边界。热图底层为高度重叠的 24 h windows，周像素之间并不满足独立同分布假设。因此颜色块数量不能直接进入普通 t-test/binomial CI。推荐在补充材料报告 weekly/monthly low-score burden，并使用 block bootstrap 处理时间相关性；Künsch [7] 可作为经典统计依据。')

# 5.5
h=doc.add_heading('5.5 FigD4_5 — Pair status and evaluability',2)
add_fig(doc,FIGS[4][0],FIGS[4][1])
add_p(doc,'面板 (a) 将小时级状态按周取 mode：D4_raw≥3.5 为 paired_consistent，3.0≤D4_raw<3.5 为 borderline，D4_raw<3.0 为 pair_asymmetry，不可评价则为 not_evaluable。由于使用 weekly mode，该面板表达“该周占优势的状态”，不是该周 asymmetry fraction；例如 49% asymmetry + 51% borderline 仍只显示 borderline。因此建议投稿版增加 weekly asymmetry fraction 作为定量补充。')
add_p(doc,'面板 (b) 给出 weekly evaluable fraction。全期 DO11–DO14 的 evaluability 均≥97.5%，ORP11 为 92.3%，ORP12/ORP13 则降至 70.3%/64.2%。这组结果建立了 D2 与 D4 的构念分离：DO14 的高低分负荷发生在高可评价性条件下，而 ORP13 的强低分必须与较窄 evidence coverage 联合解释。')
add_p(doc,'SCI 对照。Ba-Alawi 等 [4,5] 的 WWTP sensor-validation 研究将 missing/faulty intervals 与故障检测/重构作为不同任务处理。D4 更保守地让 D2 决定是否可评价，D4 只在 evidence sufficient 时给出 pair consistency，这种 abstention rather than forced classification 更适合作为多维 DQR 子项。')

# 5.6
h=doc.add_heading('5.6 FigD4_6 — Context independence and protection of D4_raw',2)
add_fig(doc,FIGS[5][0],FIGS[5][1])
add_p(doc,'面板 (a) 比较 D4_raw 与历史兼容接口 D4_forDQR_provisional；它应被理解为 independence diagnostic，而不是当前正式聚合规则。最终仲裁锁定 37,987 个 finalizable pair-hours，最大 |D4_final−D4_raw|=0.0；当前 D5 action gate 无适用 pair-hours，process Guard 与 sensor Veto 均未对 D4_raw 进行数值改写。')
add_p(doc,'面板 (b) 展示 D1 interpretation context。仓库图中的有效计数为 valid_pair=37,940、target_suspect=22、reference_unreliable=13、bilateral_unreliable=0。即在具有完整 provisional D1 context 的 37,975 行中，仅 35 行进入 unilateral D1 warning states。该结果说明 D1 与 D4 的构念隔离在实际数据流中被执行：sensor-health context 可以改变因果解释，但不能删除“pair 是否时间不一致”这一观察事实。')
add_p(doc,'SCI 对照。Schneider 等 [3] 的长期 DO/ORP/pH 研究表明 signal quality 与可保留的过程信息并非等价概念。这与当前 D4 设计一致：D1 回答 sensor health，D4 回答 homologous-pair symmetry；二者可以相关，但不能互相定义。')
add_p(doc,'可直接用于 SCI Discussion：The numerical independence of D4 was preserved throughout the final integration. Although D1 provided sensor-health context for interpretation, the final D4 value remained identical to D4_raw for all finalized pair-hours. This design separates the observation of inter-line temporal asymmetry from subsequent causal attribution and prevents a sensor-health prior from erasing genuine process-level asymmetry.')

# 5.7
h=doc.add_heading('5.7 FigD4_7 — Mechanism validation by ROC and precision–recall analysis',2)
add_fig(doc,FIGS[6][0],FIGS[6][1])
add_p(doc,'面板 (a) 的 mechanism-specific ROC 结果为 drift AUROC=0.934（95% CI 0.905–0.960）、step=0.936（0.903–0.963）、freeze=0.798（0.749–0.849）、spike=0.561（0.546–0.581）。因此 D4 对 sustained unilateral drift/step 具有强区分能力，对 freeze 中等，而对 isolated spike 接近弱区分水平。后者应被写成 D4 的设计边界：D4 使用 24 h temporal-consistency evidence，本来就不应替代 D1 的分钟级 spike detector。')
add_p(doc,'面板 (b) 给出 PR curves。最新 confirmatory 结果进一步显示 target/peer drift 的 AUPRC=0.915/0.892，target/peer step=0.920/0.900，target/peer freeze=0.774/0.795；drift/step direction accuracy 为 0.984–1.000。投稿图注应把 AUPRC 与 95% CI 直接列出，而不是要求读者从曲线估读。')
add_p(doc,'更关键的是 common-process specificity：equal common-process perturbation 的 conditional new FAR=0.041（95% CI 0–0.095），低于预设 0.10 ceiling。这说明 D4 的核心不是“发现所有变化”，而是识别非对称变化；当两条 homologous lines 同步发生相同真实工艺变化时，D4 应保持相对稳定。')
add_p(doc,'验证边界。上述结果属于 internal mechanism validation，而不是 field fault-diagnosis accuracy。当前缺少独立维护记录/便携式比对/专家 episode truth，不能报告现场 PPV、field sensitivity 或 sensor-specific specificity。另一个关键问题是 mapping benchmark 由全研究期高质量窗口构建，而 injection validation 又使用研究后段窗口；这尚未形成严格 development-only temporal calibration。顶刊投稿前应按 development → frozen mapping → terminal test → future/second-plant external test 重新组织。')
add_p(doc,'SCI 对照。Ba-Alawi 等 [4,5] 证明 mechanism-specific faulty/missing interval validation 对 WWTP sensor validation 有价值；Corominas 等 [9] 与 Newhart 等 [10] 强调高质量、可追溯数据对 WWTP 模型与运行分析的重要性。当前 D4 的优势在于显式验证 common-mode rejection，而不是只追求单类 fault AUC。')
add_p(doc,'可直接用于 SCI Results：Mechanism-specific injections demonstrated strong discrimination of sustained unilateral drift and step disturbances (AUROC 0.934 and 0.936), moderate sensitivity to unilateral freezing (AUROC 0.798), and limited sensitivity to isolated spikes (AUROC 0.561). Equal common-process perturbations generated a conditional new false-alarm rate of only 0.041, indicating that D4 preferentially responded to asymmetric rather than synchronous process changes. These estimates represent internal mechanism validation and were not interpreted as field fault-diagnosis accuracy.')

# 5.8
h=doc.add_heading('5.8 FigD4_8 — Component ablation and false-alarm control',2)
add_fig(doc,FIGS[7][0],FIGS[7][1])
add_p(doc,'面板 (a) 的 pooled unilateral-fault AUROC 从原图柱高反演约为：full≈0.807、no_dist≈0.720、no_trend≈0.799、no_var≈0.819、no_cp≈0.812、no_deadband≈0.807。最明确的结果是去除 distribution evidence 后 AUROC 下降约 0.087，表明 Q_dist 是当前 unilateral drift/step/freeze/spike mixture 中最主要的 discrimination contributor。')
add_p(doc,'但 no_var/no_cp 的 pooled AUC 略高于 full 不能被解释为 Q_var/Q_cp“多余”。消融结果依赖 benchmark mechanism mixture：当前注入集合并未专门强化 variance-only mismatch 或 pure timing-lag mismatch。因此，单一 pooled AUROC 只能回答平均区分增量，不能回答构件的理论必要性。')
add_p(doc,'面板 (b) 从仓库原图反演的 synchronous conditional new FAR 约为 full≈0.041、no_dist≈0.086、no_trend≈0.100、no_var≈0.038、no_cp≈0.024、no_deadband≈0.041。这里揭示了一个比单看 AUROC 更重要的结果：trend evidence 对抑制 common-process false asymmetry 具有关键作用。去除 trend 后 FAR 上升到约 0.10，接近预设 acceptance ceiling；因此 distribution 主要贡献 sensitivity，而 trend 的价值之一是维持 synchronous-perturbation invariance。')
add_p(doc,'Deadband 的正确解释。no_deadband 对当前注入 benchmark 几乎不改变 pooled AUROC/FAR，并不意味着 deadband 无价值。其主要职责是处理 low-dynamic process floor：当两侧 IQR 都低于物理分辨率 deadband 时豁免 variability mismatch，避免后缺氧 DO 或低动态 ORP 的分辨率限制被误判为 variance asymmetry。应使用 process-floor challenge 和 low-dynamic negative controls 专门验证，而不是用 drift/step benchmark 判断。')
add_p(doc,'Change-point 限制。confirmatory lag grid 0/10/30/60/180 min 中，lag severity 与 D4 anomaly severity 的 Spearman ρ=0.000；同时生产 CP 层先聚合到 1 h，再进行 adjacent-window distribution comparison，因此 10/30 min 低于其 nominal resolution。当前只能将 Q_cp 表述为 coarse structural-change asymmetry evidence，不能宣称精确 sub-hour lag quantification。若要强化该项，应限制确证性 lag≥1 h，或另加 minute-scale auxiliary CP detector，并按 Killick 等 [8] 的变点定位思想独立验证 detection delay/localization。')
add_p(doc,'可直接用于 SCI Discussion：Ablation analysis revealed complementary rather than interchangeable functions of the four evidence components. Distributional evidence contributed most strongly to unilateral-fault discrimination, whereas trend evidence was particularly important for suppressing false asymmetry under synchronous common-process perturbations. Variability and change-point components showed limited incremental benefit in the pooled injection benchmark; however, this should not be interpreted as redundancy because the benchmark under-represented variance-specific and timing-specific mechanisms.')

# Synthesis
h=doc.add_heading('5.9 八组图件形成的综合科学证据链',2)
for txt in [
'现场现象证据：FigD4_1、FigD4_4 与 FigD4_5 共同表明 homologous pairs 不具有统一 temporal-consistency profile。DO11/DO12/ORP11 相对稳定；DO14、ORP12、ORP13 具有明显更高的持续 pair-asymmetry burden。',
'机制分解证据：FigD4_2 与 FigD4_3 证明相同低 D4 可以由不同机制产生。DO14 主要是 distribution/change-point disparity，ORP12 主要是 distribution/variability disparity，ORP13 为多构件共同下降，因此 D4 不是简单 correlation score。',
'构念独立性证据：FigD4_5 与 FigD4_6 证明 observability、sensor health 和 pair consistency 被分别处理。D2 决定是否可评价，D1 只提供 interpretation context，正式 D4 数值源保持 D4_raw。',
'内部效度证据：FigD4_7 与 FigD4_8 表明 D4 对 sustained unilateral drift/step 灵敏，并可抑制 equal common-process changes；但 isolated spike 与 sub-hour CP lag 不是当前模块强项。']:
    add_bullet(doc,txt)

h=doc.add_heading('5.10 顶刊投稿前必须补充的验证工作包',2)
for txt in [
'P0 — development-only calibration：用早期 development period 构建 variable×regime benchmark quantiles，完全冻结后再运行 internal validation、terminal holdout 与 future-period test，禁止全期 benchmark 参与测试期映射。',
'P0 — common temporal support：W1、KS、IQR 与 slope comparison 在 target/reference 同时有效的时间戳上作为 primary analysis；当前 independent finite filtering 仅保留为 sensitivity analysis。',
'P0 — episode-level field truth：至少对 DO14、ORP12、ORP13 低分 episode 链接维护/清洗/校准记录、便携式比对、曝气/回流/碳源/负荷工艺记录，并由两名专家盲审形成 adjudicated label。',
'P0 — CP resolution contract：正式声明当前 CP 层的最小可验证时间尺度；不能用小时层算法宣称 10–30 min lag quantification。',
'P1 — benchmark expansion：补充 variance-only、timing-only、common-mode unequal、opposite-direction 与 process-floor negative controls，避免消融结论受 mechanism mixture 偏置。',
'P1 — external validity：至少增加未来独立周期；若目标 Water Research/Environmental Modelling & Software 主文，优先增加第二污水厂或第二工艺线的外部验证。',
'P1 — downstream utility：比较保留/剔除低 D4 pair-hours 对 SUMO/ASM 参数校准、软测量误差或控制决策稳定性的影响，证明 D4 是数据质量维度而非仅过程一致性监视器。']:
    add_bullet(doc,txt)

h=doc.add_heading('5.11 可直接用于顶刊论文的综合 Results',2)
add_p(doc,'Parallel-redundancy temporal consistency differed substantially among homologous sensor pairs. Of 42,847 pair-hours, 37,987 were finalized under the predefined bilateral observability contract. DO11, DO12 and ORP11 showed comparatively preserved pair consistency, whereas DO14, ORP12 and ORP13 exhibited markedly higher low-score burdens. ORP13 represented the most pronounced asymmetric pair, with a mean D4_raw of 2.650 and 67.6% of evaluable pair-hours below the predefined asymmetry threshold. Importantly, DO14 remained evaluable for 98.8% of pair-hours despite a mean D4_raw of 2.978, demonstrating that low pair-level consistency could not be attributed to limited temporal availability.')
add_p(doc,'Evidence decomposition revealed distinct mechanisms underlying similar aggregate scores. DO14 was primarily characterized by distributional and change-point-timing discrepancies while retaining high trend and variability consistency. ORP12 was dominated by distributional and variability differences, whereas ORP13 exhibited concurrent degradation across distribution, trend, variability and structural-change evidence. These patterns demonstrate that D4 represents a multidimensional pair-consistency construct rather than a simple correlation measure.')
add_p(doc,'Internal mechanism validation showed strong discrimination of sustained unilateral drift and step perturbations (AUROC=0.934 and 0.936), moderate discrimination of unilateral freezing (AUROC=0.798), and limited sensitivity to isolated spikes (AUROC=0.561). Equal common-process perturbations yielded a conditional new false-alarm rate of 0.041, indicating that the framework preferentially responded to asymmetric rather than synchronous disturbances. The estimates were interpreted as internal mechanism validation rather than field fault-diagnosis accuracy.')

h=doc.add_heading('5.12 参考文献与方法学对照',2)
for i,ref in enumerate(REFS,1):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.hanging_indent=Cm(0.5); p.paragraph_format.space_after=Pt(4); p.add_run(f'{i}. {ref}')

h=doc.add_heading('最终专家结论',1)
add_p(doc,'从 GitHub 原始图件形成的证据链看，D4 已经具备作为 Class C-minDQR 独立“平行冗余时序一致性”子维度的基本科学完整性：正式数值源 D4_raw 被保护；D2、D1、D5 与 D4 的角色边界清晰；pair-level 现场异质性、四类 evidence mechanism、common-mode rejection 与内部注入验证能够相互支撑。当前最重要的科学价值不是宣称“识别出了哪支传感器故障”，而是证明在两个理论上可比的平行处理线上，某些 homologous measurements 在排除主要可评价性问题后仍出现持续、机制可分解的 temporal asymmetry。')
add_p(doc,'若目标为顶刊方法论文，当前版本仍必须补齐 development-only calibration、common-support statistics、episode-level field truth、CP temporal-resolution contract、future/second-site external validation 和 downstream utility。完成这些工作后，D4 才能从“单厂回顾性 pair-consistency 子分”提升为具有可迁移证据的 wastewater process data-quality methodology。')

# Ensure figures were original repo files and list SHA in appendix
h=doc.add_heading('附录 A  仓库原图溯源清单',1)
add_table(doc,['序号','文件','Git blob SHA'],[(i+1,fn,git_blob_sha(FIG/fn)) for i,(fn,_) in enumerate(FIGS)])
add_p(doc,'注：本 Word 中所有图均直接由上述仓库 PNG 文件嵌入，未调用 matplotlib、seaborn 或其他绘图库重绘。图中数据、坐标、面板、配色和注释均保持仓库原始输出。')

doc.save(OUT)
print(OUT)
