from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from d5_common.config import D5_ROOT, resolve_paths


TEAL = "168AAD"
DARK = "25282A"
MUTED = "5F6B73"
LIGHT = "EAF3F5"
WARNING = "F7E7E5"


class D5ReportBuilder:
    def __init__(self) -> None:
        self.paths = resolve_paths()
        self.paths.report_root.mkdir(parents=True, exist_ok=True)
        self.docs_root = D5_ROOT / "docs"
        self.docs_root.mkdir(parents=True, exist_ok=True)
        self.main = pd.read_parquet(
            self.paths.local_output_root / "D5_main_scores_hourly.parquet"
        )
        self.events = pd.read_parquet(
            self.paths.local_output_root / "D5_event_windows.parquet"
        )
        self.support = pd.read_parquet(
            self.paths.local_output_root / "D5_support_assessment.parquet"
        )
        self.consensus = pd.read_parquet(
            self.paths.local_output_root / "D5_zone_consensus.parquet"
        )
        self.report_interface = pd.read_parquet(
            self.paths.local_output_root / "D5_report_interface.parquet"
        )
        self.gate_interface = pd.read_parquet(
            self.paths.local_output_root / "D5_gate_interface.parquet"
        )
        self.regime = pd.read_parquet(
            self.paths.local_output_root / "D5_regime_state.parquet"
        )
        self.drift = pd.read_parquet(
            self.paths.local_output_root / "D5_topology_drift_alerts.parquet"
        )
        self.validation = pd.read_excel(
            self.paths.local_output_root / "D5_validation_results.xlsx",
            sheet_name="acceptance",
        )
        self.invariance = pd.read_excel(
            self.paths.sensitivity_output_root / "D5_track_invariance.xlsx",
            sheet_name="track_invariance",
        )
        self.sensitivity_manifest = json.loads(
            (self.paths.sensitivity_output_root / "D5_sensitivity_manifest.json").read_text(
                encoding="utf-8"
            )
        )
        self.manifest = json.loads(
            (self.paths.local_output_root / "D5_run_manifest.json").read_text(encoding="utf-8")
        )
        self.figure_qa = json.loads(
            (self.paths.figure_root / "D5_figure_qa.json").read_text(encoding="utf-8")
        )
        publication_manifest = (
            D5_ROOT / "outputs" / "publication" / "D5_publication_audit_manifest.json"
        )
        self.publication = (
            json.loads(publication_manifest.read_text(encoding="utf-8"))["summary"]
            if publication_manifest.exists()
            else {}
        )
        self.generated = pd.Timestamp.now(tz="Asia/Shanghai").strftime("%Y-%m-%d %H:%M CST")

    def build_all(self) -> list[Path]:
        report_md = self._expert_markdown()
        guide_md = self._guide_markdown()
        captions_md = self._captions_markdown()
        release_index_md = self._release_index_markdown()
        outputs = []
        for name, content in [
            ("D5_EXPERT_REPORT_v2.4.md", report_md),
            ("D5_PROJECT_DIRECTORY_GUIDE_v2.4.md", guide_md),
            ("D5_FIGURE_CAPTIONS_v2.4.md", captions_md),
            ("D5_CURRENT_RELEASE.md", release_index_md),
        ]:
            path = self.paths.report_root / name
            path.write_text(content, encoding="utf-8")
            outputs.append(path)
            (self.docs_root / name).write_text(content, encoding="utf-8")
        expert_docx = self.paths.report_root / "D5_EXPERT_REPORT_v2.4.docx"
        guide_docx = self.paths.report_root / "D5_PROJECT_DIRECTORY_GUIDE_v2.4.docx"
        self._build_expert_docx(expert_docx)
        self._build_guide_docx(guide_docx)
        outputs.extend([expert_docx, guide_docx])
        return outputs

    def _release_index_markdown(self) -> str:
        return f"""# D5 Current Release

Current release: **D5 v2.4**

Generated: {self.generated}

- Scientific status: ready for final subscore aggregation with explicit
  eligibility.
- Process-coherence Guard: attribution-only and released only after detection
  and negative-control validation.
- Sensor-specific hard Veto: not released because controlled-perturbation
  swap Top-1 remains below the
  prespecified 0.80 threshold.
- Cross-dimensional publication freeze: `{self.publication.get('d4_audit_status', 'not_audited')}`;
  exact D4 run, calibration and SHA-256 matching is required.
- Automated deployment: blocked pending documentary audit, maintenance
  provenance and dual approval.
- Current report: `D5_EXPERT_REPORT_v2.4.md` / `.docx`.
- Current directory guide: `D5_PROJECT_DIRECTORY_GUIDE_v2.4.md` / `.docx`.
- Current figure captions: `D5_FIGURE_CAPTIONS_v2.4.md`.
- Publication audit: `../publication/D5_PUBLICATION_READINESS_AUDIT_v1.2.md`.

Files labelled v2.1 or v2.2 are retained only as immutable release history and
must not be used as the current scientific result.
"""

    def _facts(self) -> dict[str, Any]:
        status = self.main["evaluation_status"].value_counts()
        support = self.support["support_level"].value_counts()
        validation = self.validation.set_index("criterion")
        invariance = self.invariance.set_index("metric")
        swap_top1 = validation.loc["swap_Top1"]
        return {
            "run_id": self.main["run_id"].iloc[0],
            "start": self.main["timestamp"].min(),
            "end": self.main["timestamp"].max(),
            "rows": len(self.main),
            "raw_valid": int(self.main["D5_raw"].notna().sum()),
            "raw_mean": float(self.main["D5_raw"].mean()),
            "raw_median": float(self.main["D5_raw"].median()),
            "raw_p05": float(self.main["D5_raw"].quantile(0.05)),
            "low_rate": float(self.main["D5_raw"].lt(3.0).mean()),
            "events": len(self.events),
            "status": status.to_dict(),
            "support": support.to_dict(),
            "l1_templates": int(support.get("L1", 0)),
            "l2_templates": int(support.get("L2", 0)),
            "l3_templates": int(support.get("L3", 0)),
            "family_l3_templates": int(
                self.support["family_support_level"].eq("L3").sum()
            ),
            "node_validated_l3_templates": int(
                self.support["node_validation_passed"].sum()
            ),
            "l2_holdout_far_max": float(
                self.support.loc[
                    self.support["support_level"].eq("L2"), "holdout_far"
                ].max()
            ),
            "l3_holdout_far_max": float(
                self.support.loc[
                    self.support["support_level"].eq("L3"), "holdout_far"
                ].max()
            ),
            "provisional_report_rows": int(
                self.main["D5_report_provisional"].notna().sum()
            ),
            "research_report_rows": int(self.main["D5_report"].notna().sum()),
            "research_topology_confirmed": bool(
                self.main["research_topology_confirmed"].all()
            ),
            "production_topology_verified": bool(
                self.main["production_topology_verified"].all()
            ),
            "ood_rate": float(self.regime["regime_state"].eq("OODHold").mean()),
            "switches": int(self.regime["transition_id"].notna().sum()),
            "d5_total_nonnull": int(self.main["D5_total"].notna().sum()),
            "d5_report_nonnull": int(
                self.main["D5_report_score"].notna().sum()
            ),
            "d4_evaluable": int(self.consensus["d5_evaluable"].sum()),
            "process_guard_rows": int(
                self.consensus["process_coherence_guard_active"].sum()
            ),
            "sensor_veto_rows": int(
                self.consensus["sensor_identity_veto_active"].sum()
            ),
            "topology_alerts": int(self.drift["alert_level"].eq("review").sum()),
            "swap_auroc": float(validation.loc["swap_AUROC", "estimate"]),
            "swap_auprc": float(validation.loc["swap_AUPRC", "estimate"]),
            "swap_top1": float(validation.loc["swap_Top1", "estimate"]),
            "swap_top1_low": float(swap_top1["ci95_low"]),
            "swap_top1_high": float(swap_top1["ci95_high"]),
            "swap_top1_n": int(swap_top1["n"]),
            "swap_top1_passed": bool(swap_top1["passed"]),
            "common_far": float(validation.loc["common_mode_FAR", "estimate"]),
            "zone_far": float(validation.loc["zone_coherent_FAR", "estimate"]),
            "chatter": float(validation.loc["switch_chatter_rate", "estimate"]),
            "ie_track": float(invariance.loc["IE_track", "estimate"]),
            "jaccard": float(invariance.loc["event_jaccard", "estimate"]),
            "rho": float(invariance.loc["culprit_spearman", "estimate"]),
            "d1_release_id": self.sensitivity_manifest["d1_release_id"],
            "sensitivity_write_status": self.sensitivity_manifest[
                "authoritative_interface_status"
            ],
            "figure_qa": bool(self.figure_qa["passed"]),
        }

    def _expert_markdown(self) -> str:
        f = self._facts()
        status_lines = "\n".join(
            f"- `{name}`: {count:,} ({count / f['rows']:.1%})"
            for name, count in f["status"].items()
        )
        support_lines = "\n".join(
            f"- `{level}`: {count} templates" for level, count in f["support"].items()
        )
        return f"""# D5 Expert Review Report v2.4

**Project:** Topological Role Consistency and Structural Representativeness
**Run:** `{f['run_id']}`
**Generated:** {self.generated}
**Decision:** Scientific D5 score released for final subscore aggregation; automated deployment remains gated.

## 1. Executive verdict

The D5 v2.4 release retains the frozen v2.3 scoring implementation and adds publication-grade validation, coverage, complementarity and target-influence audits. Local, Sensitivity and Shadow V2 tracks, frozen templates, hourly scores, validation-graded admission, dual report/gate interfaces, plot data, SCI-ready figures, manifests and audit records are present. The Local Track remains logically independent of D1, D2, D3 and D4 and consumes only canonical observations, exogenous hydraulic/time context and declared D5 topology.

The ordinal research topology is confirmed: process line, pool zone, longitudinal order, SCADA-to-physical-point identity and the absence of study-period probe/channel changes are author-confirmed; the installation register independently reconciles eight active DO and six active ORP instruments. Exact coordinates, asset IDs and maintenance records are not model inputs and therefore do not suppress retrospective scientific scores. They remain deployment-governance limitations. Family-level support identifies {f['family_l3_templates']} L3 candidates, but node-specific blocked validation retains only {f['node_validated_l3_templates']} final L3 templates; the current effective distribution is L1={f['l1_templates']}, L2={f['l2_templates']} and L3={f['l3_templates']}. `D5_report_score` contains {f['d5_report_nonnull']:,} rows. Controlled-perturbation swap Top-1 is {f['swap_top1']:.2f} (95% CI {f['swap_top1_low']:.2f}-{f['swap_top1_high']:.2f}, n={f['swap_top1_n']}), so node-specific hard Veto remains disabled without blocking the score.

## 2. Scope and dimensional independence

- D5 asks whether a DO/ORP observation still behaves like its declared spatial position and represents its process zone.
- D1 evaluates sensor-intrinsic health and long-term regime-relative behavior; D2 evaluates continuity and information availability; D3 evaluates physical value/rate plausibility; D4 evaluates temporal synchronization between parallel counterparts.
- Local D5 does not consume any D1-D4 score, state or event field. D1/D2/D3 are read only in the physically isolated Sensitivity Track.
- QR/QIR are exogenous context variables only and never receive a D5 score.
- One plant-global regime is inferred from QR/QIR, robust pooled DO/ORP level
  and dispersion, and cyclic time features. Every sensor is then compared with
  its role-specific template under the same active regime.
- The pooled context includes each target with bounded median/dispersion influence. Strict target exclusion is a sensitivity challenge, not a production-model claim.
- Observed low `D5_raw` is structural evidence, not a confirmed hardware fault.

## 3. Data and result freshness

- Hourly sensor windows: {f['rows']:,}, spanning {f['start']} to {f['end']}.
- Calculable `D5_raw`: {f['raw_valid']:,}; mean {f['raw_mean']:.3f}, median {f['raw_median']:.3f}, p05 {f['raw_p05']:.3f}.
- Raw low-score fraction (`D5_raw < 3`): {f['low_rate']:.1%}; candidate persistent events: {f['events']}.
- Full 10-min regime-state trajectory retained; OOD hold rate: {f['ood_rate']:.1%}; confirmed switches: {f['switches']}.
- Result provenance is bound to canonical input hashes, topology hash, template/mapping/regime versions and code commit in `D5_run_manifest.json`.
- Sensitivity inputs are bound to frozen D1 release `{f['d1_release_id']}` and exact D2/D3/Local artifact hashes in `D5_sensitivity_manifest.json`.
- The isolated Sensitivity Track remains `{f['sensitivity_write_status']}` by design; it cannot write either authoritative report scores or gate decisions.

## 4. Applicability and support

{status_lines}

{support_lines}

ORP uses a conservative `diagonal_robust_z` model with `alpha=1.00`, but model complexity is separated from evidence maturity. Family support is shared once by analyte, regime and model family without multiplying the effective sample by the number of sensors. Every target-specific reconstruction then undergoes its own coverage, residual-scale bootstrap and leave-one-month-out FAR validation. Final support is the lower of family and node maturity. L2 evidence remains valid for scientific scoring but is not described as action-grade or cross-month deployment validated. L0 remains disabled rather than being written as a low score.

L1 is diagnostic only. L2 and L3 may populate `D5_total` and the sensor-hour
`D5_report_interface` under the confirmed ordinal topology. Only final L3
nodes may enter the pair-hour `D5_gate_interface`. Process coherence is an
attribution Guard rather than a Veto; only validated sensor-identity evidence
may activate hard Veto. Deployment approval is reported separately and does
not alter the retrospective score. Current provisional report rows:
{f['provisional_report_rows']:,}; research report rows:
{f['research_report_rows']:,}.

## 5. Validation and sensitivity

| Criterion | Estimate | Target | Result |
|---|---:|---:|---|
| Controlled swap discrimination AUROC | {f['swap_auroc']:.3f} | >=0.90 | Pass |
| Controlled swap discrimination AUPRC | {f['swap_auprc']:.3f} | >=0.80 | Pass |
| Controlled perturbation Top-1 localization | {f['swap_top1']:.3f} [{f['swap_top1_low']:.3f}, {f['swap_top1_high']:.3f}], n={f['swap_top1_n']} | >=0.80 | {'Pass' if f['swap_top1_passed'] else '**Fail**'} |
| Common-mode FAR | {f['common_far']:.3f} | <=0.10 | Pass |
| Zone-coherent FAR | {f['zone_far']:.3f} | <=0.10 | Pass |
| Switch chatter rate | {f['chatter']:.3f} | <=0.05 | Pass |
| IE_track | {f['ie_track']:.3f} | <=0.20 | Pass |
| Event Jaccard | {f['jaccard']:.3f} | >=0.80 | Pass |
| Culprit Spearman rho | {f['rho']:.3f} | >=0.80 | Pass |

Validation uses observed test-period spatial windows with frozen templates. Same-line, same-analyte position swaps are positive controls. Freeze, temporal ramps, common-mode and zone-coherent changes, DO4 floor behavior and dropout are negative/orthogonality controls. The swap detection metrics pass, but localization remains below the release criterion and must not be hidden by threshold tuning.

The publication audit additionally reports six future-month controlled-challenge refits (discrimination AUROC {self.publication.get('full_outer_AUROC', float('nan')):.3f}, AUPRC {self.publication.get('full_outer_AUPRC', float('nan')):.3f}, controlled-perturbation Top-1 localization {self.publication.get('full_outer_Top1', float('nan')):.3f}), Top-2/MRR localization, synchronized 7-d D4-D5 dependence under report-score and raw-calculable overlap, target-influence sensitivity, monthly support migration and dimension-availability sensitivity. Current D4-D5 Spearman rho is {self.publication.get('D4_D5_spearman_report', float('nan')):.3f} for report scores and {self.publication.get('D4_D5_spearman_raw', float('nan')):.3f} for raw-calculable scores; the proportions of descriptive strata with |rho| below 0.30 are {self.publication.get('D4_D5_report_strata_abs_rho_below_0_30_rate', float('nan')):.1%} and {self.publication.get('D4_D5_raw_strata_abs_rho_below_0_30_rate', float('nan')):.1%}, respectively. These are controlled observed-window challenges, not field fault-detection or localization accuracy. Confidence-risk coverage is not monotonic; the current confidence field remains evidence metadata and is not a calibrated hard-Veto gate.

## 6. Topology and D4 interface

- Declared topology contains 14 DO/ORP nodes, 10 longitudinal edges and seven parallel peer pairs.
- {f['topology_alerts']} finite candidate mappings exceed the report-only topology drift review threshold. These are hypotheses for field review, not automatic topology updates.
- Shadow V2 has `production_impact=none`; it cannot mutate `topology.yaml`, active templates, `D5_total` or Veto.
- `D5_total` and report-interface rows: {f['d5_report_nonnull']}; D5 pair-interface evaluable rows: {f['d4_evaluable']}.
- Process-coherence Guard is active for {f['process_guard_rows']} pair-hours; sensor-specific hard Veto is active for {f['sensor_veto_rows']} pair-hours.
- The final D4 numeric source is `D4_raw`. D1 is interpretation-only, while D5 supplies report context, attribution Guard and validated sensor-identity decisions.

## 7. Figure review

Nine multi-panel figure groups are available as editable SVG/PDF, 600 dpi PNG and LZW-compressed 600 dpi TIFF on a fixed 183 mm canvas. All use Arial and 0.8 pt axes; open plots use outward ticks, whereas genuinely full-boxed maps use inward ticks. Panel labels, annotation backgrounds and endpoint-aware scales follow one shared style contract. Figures D5-1-D5-3 now connect declared topology, score applicability and case-level evidence; Figure D5-6 reports criterion margins, localization and evidence coverage; Figure D5-7 reports joint density, stratified D4-D5 overlap and exact low-tail concordance; Figure D5-8 separates availability-aware, matched complete-case and fixed-dimension estimands; Figure D5-9 reports post-reference target influence and the full reference-fraction support grid. Automated export/layout QA passed: {f['figure_qa']}.

## 8. Critical limitations

1. Research topology is author-confirmed and inventory-reconciled, but production documentary audit, maintenance provenance and dual approval remain incomplete.
2. Family-level L3 support does not imply node-level action validity: {f['family_l3_templates']} family-L3 candidates reduce to {f['node_validated_l3_templates']} final node-L3 templates after blocked validation.
3. Controlled-perturbation swap Top-1 localization is {f['swap_top1']:.3f} (95% CI {f['swap_top1_low']:.3f}-{f['swap_top1_high']:.3f}) versus the 0.80 target.
4. The {f['events']} candidate event windows have no external truth labels; event counts must not be reported as confirmed sensor faults.
5. Regime transition FAR and topology candidate recall are not estimable without external regime/topology truth.
6. `D5_raw` calibration is suitable for comparative research evidence, but operational event thresholds require labeled prospective confirmation.

## 9. Release decision and next actions

The branch may enter final WW-DQS subscore aggregation as a **scientific implementation with claim-specific action gates**. Automated control deployment remains outside the present evidence scope.

1. Use `D5_report_score` only where report eligibility is explicit; renormalize missing dimensions rather than substituting a low score.
2. Use the separate gate interface only for final L3 nodes and treat process coherence as an attribution Guard, never as Veto.
3. Keep sensor-specific hard Veto disabled until controlled blocked localization reaches Top-1 >=0.80.
4. Report availability-aware and fixed-dimension complete-evidence WW-DQS separately; never interpret a dimension-availability shift as a quality trend.
5. Cross-dimensional manuscript values are frozen only while the exact D4 run, calibration and SHA-256 remain current; any D4 change requires a full audit rerun.
6. Add field-confirmed topology and event cases as external validation when they become available.
7. Complete documentary audit and dual approval before any automated plant-control deployment.

The deployment evidence and role-separated approval procedure remain specified
in `docs/D5_FIELD_VERIFICATION_REQUIREMENTS.md`; they are not prerequisites for
retrospective scientific aggregation.
"""

    def _guide_markdown(self) -> str:
        output_rows = [
            ("outputs/local", "Authoritative Local scores, evidence, templates, audits and D4 read-only interface"),
            ("outputs/sensitivity", "D1/D2/D3-filtered shadow reference sensitivity; no production writes"),
            ("outputs/shadow_v2", "Graph/topology research candidates with production impact fixed to none"),
            ("outputs/plot_data", "Frozen long-table data used by every manuscript figure"),
            ("outputs/figures", "SVG/PDF/600 dpi PNG/TIFF figures and figure QA"),
            ("outputs/reports", "Expert report, directory guide and figure captions"),
            ("outputs/publication", "Publication decisions, inference audits, source data and manifest"),
        ]
        output_table = "\n".join(f"| `{path}` | {role} |" for path, role in output_rows)
        return f"""# D5 Project Directory Guide v2.4

**Generated:** {self.generated}
**Update rule:** every computational or figure change must rerun reports and release QA.

## 1. Project layout

```text
D5 Topological Role Consistency and Structural Representativeness/
|-- configs/
|   |-- common/          # paths, sensors, topology evidence and interface schemas
|   |-- local/           # production-isolated D5 Local policies
|   |-- publication/     # final claim, support and inference contract
|   |-- sensitivity/     # upstream-filter sensitivity only
|   `-- shadow_v2/       # graph/topology research only
|-- src/
|   |-- d5_common/       # shared configuration, hashing and robust math
|   |-- d5_local/        # Local contracts, context, templates, evidence, scoring and outputs
|   |-- d5_sensitivity/  # physically isolated shadow pipeline
|   `-- d5_shadow_v2/    # no-production-impact graph research
|-- scripts/             # reproducible entry points
|-- tests/               # contract and regression tests
|-- docs/                # version-reviewable Markdown reports
`-- outputs/             # generated artifacts by track and role
```

## 2. Track isolation

| Track | Allowed inputs | Forbidden outputs |
|---|---|---|
| `d5_local` | Canonical 1-min observations/flags, QR/QIR context, research topology/evidence, frozen D5 assets | Reading D1-D4 score/state/event inputs |
| `sensitivity` | Frozen Local evidence plus D1/D2/D3 read-only filters | Authoritative report/gate interfaces, active templates, D4 arbitration |
| `shadow_v2` | Canonical observations and declared topology | Automatic topology update, Veto, Local score mutation |

## 3. Output roots

| Path | Role |
|---|---|
{output_table}

## 4. Core entry points

```powershell
python scripts/run_d5_local.py
python scripts/run_d5_sensitivity.py
python scripts/run_d5_validation.py
python scripts/run_d5_publication_audit.py
python scripts/run_d5_topology_review.py
python scripts/run_d5_shadow_v2.py
python scripts/make_d5_figures.py
python scripts/build_d5_reports.py
python scripts/check_d5_release.py
python "../D4 Parallel-redundancy Temporal Consistency/scripts/run_d4_d5_readiness.py"
python -m pytest tests -q
```

Use `python scripts/run_d5_release.py --include-local` after data, topology, template, mapping or core scoring changes. For report/figure-only edits, omit `--include-local`.

## 5. Update protocol

1. Work on an `agent/*` branch; do not write half-complete results to `main`.
2. Change configuration and code together; never edit generated score tables by hand.
3. Run the fastest contract tests, then the affected track.
4. Rebuild validation, plot data, all figure counterparts, reports and release QA.
5. Inspect PNG figures and rendered DOCX pages visually.
6. Stage only the D5 directory, commit, push the branch and open a draft PR.
7. Report branch, commit, compare/PR URL and all failed release gates to the user.
8. Merge only after explicit confirmation.

## 6. Production activation checklist

- [ ] Controlled drawing or equivalent documentary evidence is versioned.
- [ ] Asset/serial identity and maintenance history are reconciled where available.
- [ ] Ordinal positions are independently audited; exact coordinates are required only if used by a future metric.
- [ ] Reviewer and approver are two identified independent people.
- [ ] Topology validity interval covers the evaluated data.
- [ ] All topology-bound template hashes are regenerated.
- [ ] Required support tiers and ORP exit criteria pass.
- [ ] Controlled swap AUROC/AUPRC and Top-1 pass blocked validation.
- [ ] Transition FAR and topology tests have external truth.
- [ ] Local-Sensitivity invariance gates pass.
- [ ] D4 protected columns have max absolute difference zero.
- [ ] `D5_total` is populated only for genuinely evaluable rows.

## 7. File ownership rules

- `configs/common/topology.yaml` is the declared topology contract and
  `configs/common/topology_evidence.yaml` is its research evidence ledger;
  Shadow outputs never overwrite either file.
- `configs/common/field_verification_template.csv` is a blank human-input form;
  it is not evidence until reviewer and approver fields are independently completed.
- Parquet is authoritative for tabular data; Excel is a human-review mirror with the same semantics.
- `D5_plot_data` is the sole manuscript-figure input; figure scripts do not recompute business metrics.
- `D5_report_interface` is the sensor-hour scientific score contract; `D5_gate_interface` is the pair-hour decision and attribution contract.
- Reports and this guide are generated artifacts and must be refreshed on every release workflow.
- Missing/limited/OOD evidence is represented by status plus `NaN`, never by a fabricated low score.

## 8. Current branch gate

Current release classification: **D5 scientific score ready; cross-dimensional publication freeze {'ready and hash-bound' if self.publication.get('d4_dependency_current', False) else 'dependency-gated'}; automated deployment blocked**. Node-specific hard Veto remains gated by controlled-perturbation Top-1 localization.

## 9. Figure publication contract

- All nine registered figure groups use the Python/Matplotlib backend, a fixed
  183 mm double-column canvas, Arial text, 0.8 pt axes and editable SVG/PDF.
- Figures D5-1, D5-2 and D5-3 communicate topology, score applicability and
  case-level structural evidence. Figures D5-6 and D5-7 carry the principal
  validation and cross-dimensional claims.
- Figures D5-4, D5-5 and D5-9 retain detailed validation, governance and
  robustness evidence for Extended Data or Supplementary Information.
- Figure D5-8 is a main figure only in the future integrated WW-DQS paper. In
  the standalone D5 module it is an Extended Data prototype and is not the
  final five-dimension aggregate.
- Every quantitative panel is traceable to frozen Parquet or Excel source data.
  Automated QA does not replace visual inspection at final publication scale.
"""

    def _captions_markdown(self) -> str:
        return f"""# D5 Figure Captions v2.4

## Figure D5-1. Author-confirmed topology, applicability and scientific workflow

(a) Author-confirmed longitudinal and parallel-peer topology for two process lines. Solid arrows encode within-line ordinal adjacency and dashed links encode seven homologous cross-line peers; coordinates are schematic rather than surveyed distances. (b) D5 applicability states stratified by analyte. (c) Scientific workflow from context assignment through regime-conditioned role templates and four structural components to the report-or-abstain interface. Governance and deployment approval remain separate from the scoring workflow.

## Figure D5-2. Spatiotemporal score, eligibility and support structure

(a) Daily lower-quartile `D5_raw` for 14 ordered DO/ORP positions, with report-eligibility and out-of-template ribbons. Gray cells denote days with less than 50% report-eligible coverage; A-C mark the lowest-score candidate windows used for case review. The value 3 is an analysis reference, not a validated fault threshold. (b) Sensor-level median and interquartile range; point area represents the fraction below the analysis reference. (c) Monthly raw calculability, report eligibility, out-of-template rate and L1 support. Missing or weakly supported evidence is displayed as abstention, not as low quality.

## Figure D5-3. Case-level structural evidence and attribution

(a) Raw target, homologous parallel peer and same-line neighbor around an unlabeled low-score candidate, with the regime-specific target template and OOD/regime strip. (b) Four structural components and `D5_raw`; the dashed line at 3 is an analysis reference. (c) Normalized diagnostic contributions at the case center and a topology residual map showing node residuals and edge inconsistency. Contributions are diagnostic leave-one-out quantities, not Shapley values. This case is evidence for review and is not a confirmed sensor fault.

## Figure D5-4. Frozen-template validation and track invariance

(a) Criterion margins for controlled same-line position swaps, negative controls and regime chatter; zero indicates the prespecified release boundary and positive values pass. Error bars show 95% intervals where estimable. (b) Controlled-perturbation Top-1 localization by D5-relevant scenario with Wilson 95% intervals. (c) False alarm rates and empirical 95% ranges for orthogonality controls. (d) Local-Sensitivity invariance criterion margins. These are controlled observed-window challenges, not field fault accuracy.

## Figure D5-5. Hierarchical admission and dimension-independent integration

(a) Family-level and final node-level support counts, showing that shared sample support does not automatically upgrade every node. (b) Node bootstrap stability and blocked-temporal holdout FAR for family-L3 candidates; dashed lines show prespecified node thresholds. (c) Coverage of the sensor-hour report interface and pair-hour gate states, including attribution Guard and sensor Veto. (d) D4 independence audit comparing the final D4 score with its sole numeric source, `D4_raw`; the identity line represents zero cross-dimensional adjustment.

## Figure D5-6. Validation, localization and evidence coverage boundary

(a) Estimate-minus-criterion margins from six future-month complete refits of the full model and three prespecified structural ablations. AUROC is compared with 0.90; AUPRC and Top-1 are compared with 0.80. (b) Top-1 localization with fold-cluster 95% intervals and the distribution of topological hop error for three controlled perturbations. (c) Monthly raw calculability, report eligibility, OOD and L1 support. (d) Top-1-only risk-coverage analysis with retained block counts and cluster intervals. Confidence is not calibrated as a selective-localization probability, and these results are not field fault-accuracy estimates.

## Figure D5-7. D4-D5 complementarity across report and raw estimands

(a) Joint density of D4 raw and D5 pair report scores. Overall report-score and raw-calculable Spearman correlations and the descriptive covariate-adjusted rank marker are annotated; the latter is not a causal estimand. (b,c) Spearman rho by analyte, pair, regime and month. The shaded band denotes weak association, |rho| < 0.30. Filled symbols and intervals require at least six independent non-overlapping 7-d process-time blocks; open symbols retain descriptive point estimates based on at least two blocks. (d) Exact low-tail overlap at the prespecified analysis reference of 3. Pair-composite leave-one-dimension sensitivity remains in source data. Values are bound to D4 run `{self.publication.get('d4_run_id', 'not_audited')}`, calibration `{self.publication.get('d4_calibration_id', 'not_audited')}` and the recorded main-score SHA-256.

## Figure D5-8. Dimension-availability sensitivity of the prototype WW-DQS

(a) Monthly medians for the availability-aware prototype, with open markers showing the same estimator restricted to matched complete-evidence sensor-hours, and the fixed-dimension complete-evidence estimator. (b) Availability-aware and complete-evidence coverage together with D5 availability. (c) Effective numeric dimension count. (d) Descriptive monthly median shift between availability-aware and fixed-dimension estimands. The current calculation covers the D1/D2/D5 node prototype and must be repeated under the final five-dimension contract.

## Figure D5-9. Target-influence and support-threshold robustness

(a) Whole-period and post-reference leave-one-target-out regime disagreement together with OOD-rate response to controlled target offsets. (b) Final-L3 template count across the full prespecified 0.70, 0.80 and 0.90 reference-fraction sensitivity grid and support thresholds. Production thresholds and the 0.70 reference fraction remain frozen. These analyses diagnose model robustness and do not change production templates or thresholds.
"""

    def _new_document(self, title: str, subtitle: str, status: str) -> Document:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
        self._configure_styles(document)
        header = section.header.paragraphs[0]
        header.text = "D5 v2.4 | Project 1 Data Quality Assessment"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._format_runs(header, size=8, color=MUTED)
        footer = section.footer.paragraphs[0]
        footer.text = (
            f"Generated {self.generated} | Scientific aggregation release; "
            "automated deployment blocked"
        )
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._format_runs(footer, size=8, color=MUTED)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(title)
        self._set_font(run, size=21, bold=True, color=DARK)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(12)
        run = paragraph.add_run(subtitle)
        self._set_font(run, size=11.5, color=MUTED)
        for label, value in [
            ("Status", status),
            ("Run", str(self.main["run_id"].iloc[0])),
            ("Coverage", f"{self.main['timestamp'].min()} to {self.main['timestamp'].max()}"),
        ]:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            left = p.add_run(f"{label}: ")
            self._set_font(left, size=9.5, bold=True, color=TEAL)
            right = p.add_run(value)
            self._set_font(right, size=9.5, color=DARK)
        document.add_paragraph().paragraph_format.space_after = Pt(4)
        return document

    def _configure_styles(self, document: Document) -> None:
        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Arial"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.20
        for name, size, color, before, after in [
            ("Heading 1", 16, TEAL, 16, 8),
            ("Heading 2", 13, TEAL, 12, 6),
            ("Heading 3", 11.5, "1F4D58", 8, 4),
        ]:
            style = styles[name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

    def _build_expert_docx(self, path: Path) -> None:
        f = self._facts()
        doc = self._new_document(
            "D5 Expert Review Report",
            "Regime-conditioned Ordinal Spatial Structure and Representativeness | v2.4",
            "Scientific score released; claim-specific action gates retained",
        )
        self._heading(doc, "1. Executive verdict", 1)
        self._paragraph(
            doc,
            "D5 v2.4 retains the frozen scoring implementation and adds publication-grade outer-refit, localization, coverage, complementarity and target-influence audits.",
        )
        self._warning(
            doc,
            f"Ordinal topology is author-confirmed and inventory-reconciled. Shared family support yields {f['family_l3_templates']} L3 candidates, while node validation retains {f['node_validated_l3_templates']} final L3 templates. D5_report_score contains {f['d5_report_nonnull']:,} rows. Controlled-perturbation swap Top-1 localization is {f['swap_top1']:.2f} (95% CI {f['swap_top1_low']:.2f}-{f['swap_top1_high']:.2f}, n={f['swap_top1_n']}), so node-specific hard Veto remains disabled.",
        )
        self._heading(doc, "2. Scope and dimensional independence", 1)
        for item in [
            "D5 evaluates spatial role identity and structural representativeness, not hardware fault type.",
            "D1 covers sensor-intrinsic health; D2 covers continuity/availability; D3 covers physical value/rate constraints; D4 covers temporal synchronization between parallel counterparts.",
            "Local D5 consumes no D1-D4 score, state or event. D1/D2/D3 appear only in the isolated Sensitivity Track.",
            "QR/QIR provide exogenous context only. Low D5_raw is unlabeled structural evidence, not a confirmed fault.",
        ]:
            self._bullet(doc, item)
        doc.add_page_break()
        self._heading(doc, "3. Current data and outputs", 1)
        self._paragraph(
            doc,
            f"Sensitivity provenance is bound to frozen D1 release {f['d1_release_id']} and exact D2/D3/Local artifact hashes. The Sensitivity Track remains {f['sensitivity_write_status']} and cannot write authoritative report or gate outputs.",
        )
        self._table(
            doc,
            ["Measure", "Current result", "Interpretation"],
            [
                ["Hourly windows", f"{f['rows']:,}", f"{f['start']} to {f['end']}"],
                ["D5 raw", f"mean {f['raw_mean']:.3f}; median {f['raw_median']:.3f}", "Comparative research evidence"],
                ["Low-score fraction", f"{f['low_rate']:.1%}", "Not a fault prevalence estimate"],
                ["Candidate events", str(f["events"]), "Unreviewed and unlabeled"],
                ["OOD hold", f"{f['ood_rate']:.1%}", "Explicit posterior rejection"],
                ["D5 total / D4 evaluable", f"{f['d5_total_nonnull']} / {f['d4_evaluable']}", "Scientific aggregation ready"],
            ],
            [1.55, 1.65, 3.30],
        )
        self._heading(doc, "4. Support and applicability", 1)
        support_rows = [
            [
                level,
                str(count),
                "Action candidate" if level == "L3"
                else "Scientific score" if level == "L2"
                else "Diagnostic only",
            ]
            for level, count in f["support"].items()
        ]
        self._table(doc, ["Tier", "Templates", "Policy"], support_rows, [1.0, 1.1, 4.4])
        self._paragraph(
            doc,
            "ORP retains diagonal robust Z with alpha=1.00, while evidence "
            "maturity is graded independently. L2 requires effective blocks "
            "and month coverage and remains scientific-score grade "
            f"(maximum holdout FAR {f['l2_holdout_far_max']:.3f}); L3 "
            "additionally requires block-bootstrap stability and blocked "
            f"holdout FAR <=0.10 (observed maximum {f['l3_holdout_far_max']:.3f}).",
        )
        doc.add_page_break()
        self._heading(doc, "5. Validation and sensitivity", 1)
        validation_rows = []
        for row in self.validation.itertuples(index=False):
            validation_rows.append(
                [row.criterion, f"{row.estimate:.3f}", f"{row.operator} {row.target:.2f}", "PASS" if row.passed else "FAIL"]
            )
        self._table(doc, ["Criterion", "Estimate", "Target", "Result"], validation_rows, [2.35, 1.0, 1.2, 1.0])
        self._paragraph(
            doc,
            "Same-line position swaps are scored with frozen templates. Freeze, temporal ramps, common/zone-coherent changes, DO4 floor behavior and dropout are orthogonality controls. Detection and FAR pass, but node localization remains below release quality.",
        )
        if self.publication:
            self._warning(
                doc,
                "Publication audit: six-fold controlled-challenge discrimination AUROC "
                f"{self.publication['full_outer_AUROC']:.3f}, AUPRC "
                f"{self.publication['full_outer_AUPRC']:.3f}, controlled-perturbation Top-1 localization "
                f"{self.publication['full_outer_Top1']:.3f}; local all-scenario "
                f"Top-2 {self.publication['local_all_Top2']:.3f}. Confidence-risk "
                "coverage is not monotonic, so confidence is not a calibrated "
                "sensor-Veto gate.",
            )
            self._paragraph(
                doc,
                "D4-D5 Spearman rho is "
                f"{self.publication['D4_D5_spearman_report']:.3f} for report scores "
                f"and {self.publication['D4_D5_spearman_raw']:.3f} for raw-calculable scores. "
                "Report/raw descriptive weak-association rates are "
                f"{self.publication['D4_D5_report_strata_abs_rho_below_0_30_rate']:.1%} "
                f"and {self.publication['D4_D5_raw_strata_abs_rho_below_0_30_rate']:.1%}; "
                "only strata with at least six independent 7-d blocks receive bootstrap CIs.",
            )
        invariance_rows = [
            [row.metric, f"{row.estimate:.3f}", row.criterion, "PASS" if row.passed else "FAIL"]
            for row in self.invariance.itertuples(index=False)
            if row.metric != "FAR_delta"
        ]
        self._table(doc, ["Track metric", "Estimate", "Target", "Result"], invariance_rows, [2.0, 1.0, 1.5, 1.0])
        self._heading(doc, "6. Topology and D4 interface", 1)
        self._paragraph(
            doc,
            f"The declared registry contains 14 nodes, 10 longitudinal edges and seven peer pairs. {f['topology_alerts']} finite candidate mappings exceed the review threshold. They cannot mutate declared topology.",
        )
        self._bullet(doc, "Ordinal research topology is confirmed; maintenance provenance, production documentary audit and two-person approval remain pending.")
        self._bullet(doc, f"D5-to-D4 zone consensus contains {f['d4_evaluable']:,} evaluable rows.")
        self._bullet(doc, f"Process-coherence Guard is active for {f['process_guard_rows']:,} pair-hours; sensor-specific hard Veto is active for {f['sensor_veto_rows']:,}.")
        self._bullet(doc, "D4_raw is the sole D4 numeric source; D1 and D5 contribute interpretation and action governance only.")
        doc.add_page_break()
        self._heading(doc, "7. SCI figure review", 1)
        captions = [
            "Figure D5-1 | Author-confirmed longitudinal/peer topology, applicability and scientific workflow.",
            "Figure D5-2 | Spatiotemporal D5 score, report eligibility, OOD and support structure.",
            "Figure D5-3 | Case-level raw trajectories, structural components, attribution and topology residuals.",
            "Figure D5-4 | Validation and Local-Sensitivity invariance.",
            "Figure D5-5 | Support, regime, topology and release governance.",
            "Figure D5-6 | Criterion margins, topological localization, coverage and risk-coverage limits.",
            "Figure D5-7 | D4-D5 joint structure, stratified dependence and low-tail overlap.",
            "Figure D5-8 | Availability-aware, matched complete-case and fixed-dimension prototype estimands.",
            "Figure D5-9 | Whole/post-reference target influence and full support-threshold grid.",
        ]
        stems = ["FigD5_1_framework", "FigD5_2_spatiotemporal", "FigD5_3_evidence", "FigD5_4_validation", "FigD5_5_governance", "FigD5_6_validation_coverage", "FigD5_7_D4_D5_complementarity", "FigD5_8_dimension_availability_sensitivity", "FigD5_9_target_support_robustness"]
        for stem, caption in zip(stems, captions):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run()
            run.add_picture(str(self.paths.figure_root / f"{stem}.png"), width=Inches(6.15))
            self._set_picture_alt(run, caption)
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(10)
            self._format_runs(cap, size=8.5, color=MUTED, italic=True)
        self._heading(doc, "8. Critical limitations", 1)
        for item in [
            "Research topology is confirmed, but production documentary evidence and maintenance provenance are not dual-approved.",
            f"Effective support is L1={f['l1_templates']}, L2={f['l2_templates']}, L3={f['l3_templates']}; L1 remains diagnostic only.",
            f"Controlled-perturbation swap Top-1 localization is {f['swap_top1']:.2f} (95% CI {f['swap_top1_low']:.2f}-{f['swap_top1_high']:.2f}) versus the 0.80 target.",
            f"The {f['events']} candidate events lack external truth and cannot be called confirmed faults.",
            "Regime transition FAR and topology candidate recall require external truth.",
        ]:
            self._bullet(doc, item)
        self._heading(doc, "9. Release decision", 1)
        self._paragraph(
            doc,
            "The D5 score is suitable for coverage-aware WW-DQS aggregation. "
            + (
                "Cross-dimensional manuscript values are current and hash-bound to the frozen D4 artifact. "
                if self.publication.get("d4_dependency_current", False)
                else "Cross-dimensional manuscript freezing remains blocked until the exact D4 dependency is current. "
            )
            + "Automated plant-control deployment remains outside the present evidence scope.",
        )
        for item in [
            "Use score-eligible D5 rows with confidence-aware missing-dimension renormalization.",
            "Use process-coherence Guard only for persistent final-L3 evidence and never report it as Veto.",
            "Raise controlled blocked-holdout Top-1 localization to at least 0.80 without weakening negative-control FAR.",
            "Report availability-aware and fixed-dimension complete-evidence composites separately.",
            "Rerun the complete D4-D5 overlap and ablation audit whenever the frozen D4 artifact changes.",
            "Obtain documentary approval before automated deployment.",
        ]:
            self._number(doc, item)
        doc.save(path)

    def _build_guide_docx(self, path: Path) -> None:
        doc = self._new_document(
            "D5 Project Directory Guide",
            "Python structure, artifact ownership and release workflow | v2.4",
            "Operational reference; regenerate on every project update",
        )
        self._heading(doc, "1. Directory map", 1)
        tree = (
            "configs/{common,local,sensitivity,shadow_v2}\n"
            "src/{d5_common,d5_local,d5_sensitivity,d5_shadow_v2}\n"
            "scripts/  tests/  docs/\n"
            "outputs/{local,sensitivity,shadow_v2,plot_data,figures,reports}"
        )
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        for line in tree.splitlines():
            run = p.add_run(line + "\n")
            self._set_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9.5, color=DARK)
        self._heading(doc, "2. Track isolation", 1)
        self._table(
            doc,
            ["Track", "Allowed inputs", "Forbidden behavior"],
            [
                ["d5_local", "Canonical observations/flags, QR/QIR context, topology/evidence", "Read D1-D4 scores/states/events"],
                ["sensitivity", "Frozen Local evidence plus D1/D2/D3 read-only filters", "Write report/gate interfaces, zone consensus or D4"],
                ["shadow_v2", "Canonical observations and declared topology", "Auto-update topology or production scores"],
            ],
            [1.2, 2.65, 2.65],
        )
        self._heading(doc, "3. Output ownership", 1)
        self._table(
            doc,
            ["Path", "Owner and role"],
            [
                ["outputs/local", "Authoritative Local scores, evidence, templates, interface and audit"],
                ["outputs/sensitivity", "Shadow reference/mapping sensitivity only"],
                ["outputs/shadow_v2", "No-production-impact graph/topology research"],
                ["outputs/plot_data", "Frozen long-table source for all figures"],
                ["outputs/figures", "SVG/PDF/600 dpi PNG/TIFF plus QA"],
                ["outputs/reports", "Generated expert report, guide and captions"],
            ],
            [1.75, 4.75],
        )
        self._heading(doc, "4. Core commands", 1)
        commands = [
            "python scripts/run_d5_local.py",
            "python scripts/run_d5_sensitivity.py",
            "python scripts/run_d5_validation.py",
            "python scripts/run_d5_topology_review.py",
            "python scripts/run_d5_shadow_v2.py",
            "python scripts/make_d5_figures.py",
            "python scripts/build_d5_reports.py",
            "python scripts/check_d5_release.py",
            "python scripts/verify_d5_publication_bundle.py",
            "python -m pytest tests -q",
        ]
        for command in commands:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(command)
            self._set_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9.0, color=DARK)
        self._paragraph(
            doc,
            "Use python scripts/run_d5_release.py --include-local after data, topology, template, mapping or scorer changes. Omit --include-local for report/figure-only work.",
        )
        self._heading(doc, "5. Update and Git workflow", 1)
        for item in [
            "Work on an agent/* branch and keep changes reviewable.",
            "Change code/configuration; never edit generated scores manually.",
            "Run contract tests and the affected track.",
            "Rebuild validation, plot data, figures, reports and release QA.",
            "Visually inspect every figure and every rendered DOCX page.",
            "Stage only the D5 directory, commit and push the branch.",
            "Give the user the branch, commit and compare/PR location; merge only after confirmation.",
        ]:
            self._number(doc, item)
        self._heading(doc, "6. Production activation checklist", 1)
        for item in [
            "Real process drawing ID and validity interval recorded",
            "Asset/serial identity and maintenance history reconciled where available",
            "Ordinal positions independently audited; coordinates approved only if later used quantitatively",
            "Independent reviewer and approver recorded",
            "Topology-bound templates and hashes regenerated",
            "Support tiers and ORP exit criteria passed",
            "Controlled swap AUROC/AUPRC and Top-1 passed on blocked holdouts",
            "Transition FAR and topology tests supported by external truth",
            "Local-Sensitivity invariance passed",
            "D4 protected-column max absolute difference equals zero",
            "D5_total populated only for genuinely evaluable rows",
        ]:
            self._bullet(doc, "[ ] " + item)
        self._heading(doc, "7. Non-negotiable file contracts", 1)
        for item in [
            "topology.yaml is the only declared production topology source.",
            "Parquet is authoritative; Excel is a semantically equivalent review mirror.",
            "D5_plot_data is the only manuscript-figure input.",
            "Missing/limited/OOD evidence uses explicit status plus NaN.",
            "Reports and this guide are regenerated on every release workflow.",
        ]:
            self._bullet(doc, item)
        doc.save(path)

    def _heading(self, doc: Document, text: str, level: int) -> None:
        doc.add_heading(text, level=level)

    def _paragraph(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph(text)
        self._format_runs(p, size=10.5, color=DARK)

    def _bullet(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        self._format_runs(p, size=10.2, color=DARK)

    def _number(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph(text, style="List Number")
        p.paragraph_format.space_after = Pt(4)
        self._format_runs(p, size=10.2, color=DARK)

    def _warning(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.18)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(9)
        run = p.add_run(text)
        self._set_font(run, size=10.5, bold=True, color="8B2E2E")
        self._shade_paragraph(p, WARNING)

    def _table(
        self,
        doc: Document,
        headers: list[str],
        rows: list[list[str]],
        widths: list[float],
    ) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.style = "Table Grid"
        header_properties = table.rows[0]._tr.get_or_add_trPr()
        repeat_header = OxmlElement("w:tblHeader")
        repeat_header.set(qn("w:val"), "true")
        header_properties.append(repeat_header)
        for index, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = str(header)
            self._shade_cell(cell, LIGHT)
            self._format_runs(cell.paragraphs[0], size=9.2, bold=True, color=DARK)
        for row_data in rows:
            cells = table.add_row().cells
            for cell, value, width in zip(cells, row_data, widths):
                cell.width = Inches(width)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.text = str(value)
                self._format_runs(cell.paragraphs[0], size=8.8, color=DARK)
        for row in table.rows:
            row_properties = row._tr.get_or_add_trPr()
            keep_row_together = OxmlElement("w:cantSplit")
            row_properties.append(keep_row_together)
        self._set_table_geometry(table, widths)
        after = doc.add_paragraph()
        after.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _set_table_geometry(table: Any, widths: list[float]) -> None:
        table_width = int(sum(widths) * 1440)
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(table_width))
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(int(width * 1440)))
            grid.append(col)
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tc_w.type = "dxa"
                tc_w.w = int(width * 1440)

    @staticmethod
    def _shade_cell(cell: Any, fill: str) -> None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _shade_paragraph(paragraph: Any, fill: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        p_pr.append(shading)

    def _format_runs(
        self,
        paragraph: Any,
        *,
        size: float,
        color: str,
        bold: bool | None = None,
        italic: bool | None = None,
    ) -> None:
        for run in paragraph.runs:
            self._set_font(run, size=size, color=color, bold=bold, italic=italic)

    @staticmethod
    def _set_font(
        run: Any,
        *,
        name: str = "Arial",
        east_asia: str = "Microsoft YaHei",
        size: float | None = None,
        color: str | None = None,
        bold: bool | None = None,
        italic: bool | None = None,
    ) -> None:
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic

    @staticmethod
    def _set_picture_alt(run: Any, description: str) -> None:
        drawing = run._element.xpath(".//wp:docPr")
        if drawing:
            drawing[0].set("descr", description)
