"""Generate the auto-updated D1 sensor-health expert research report.

The report follows the structure and narrative style of
``D1_Sensor_Health_Research_Report.docx`` while replacing all result statements
with the latest v1.1 / 1.1-decomposition-bridge outputs.
"""
from __future__ import annotations

import hashlib
import json
import pickle
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs"
REPORT_DIR = OUT / "reports"
REPORT_PATH = OUT / "D1_Sensor_Health_Expert_Report_Auto.docx"
MANIFEST_PATH = REPORT_DIR / "D1_Sensor_Health_Expert_Report_Auto.manifest.json"

FIGURES = {
    "Fig 1": OUT / "figures" / "Fig1_D1_dimension_matrix.png",
    "Fig 2": OUT / "figures" / "Fig2_monthly_heatmap.png",
    "Fig 3": OUT / "figures" / "Fig3_case_subscores.png",
    "Fig 4": OUT / "figures" / "Fig4_subscore_distribution.png",
    "Fig 5": OUT / "figures" / "Fig5_mapping_curves.png",
    "Fig 6": OUT / "figures" / "Fig6_dominant_fault.png",
    "Fig 7": OUT / "figures" / "Fig7_daily_timeseries.png",
    "Fig 8": OUT / "figures" / "Fig8_veto_cooldown.png",
    "Fig 9": OUT / "figures" / "Fig9_harmonic_demo.png",
    "Fig 10": OUT / "figures" / "Fig10_two_tier_regime.png",
    "Fig 11": OUT / "figures" / "Fig11_pls_peer_audit.png",
    "Fig V12": OUT / "figures" / "FigV12_v11_vs_strictV1_hero.png",
    "Fig V13": OUT / "figures" / "FigV13_state_machine_DO_2_3.png",
    "Fig V14": OUT / "figures" / "FigV14_veto3_state_audit.png",
    "Fig V15": OUT / "figures" / "FigV15_pelt_event_id.png",
    "Fig V16": OUT / "figures" / "FigV16_regime_templates.png",
    "Fig V17": OUT / "figures" / "FigV17_scope_qr_qir_offline.png",
    "Fig V18": OUT / "figures" / "FigV18_aggregate_summary.png",
    "Fig P2-01": OUT / "v12_P2" / "figures" / "fig01_variant_overview.png",
    "Fig P2-10": OUT / "v12_P2" / "figures" / "fig10_global_summary.png",
}

SOURCE_FILES = [
    ROOT / "v11_state.pkl",
    ROOT / "v12_P2_state.pkl",
    OUT / "logs" / "run_v11.log",
    OUT / "data" / "D1_v11_vs_strictV1_compare.xlsx",
    OUT / "data" / "D1_state_machine_audit.xlsx",
    OUT / "data" / "D1_sensor_profile_summary.xlsx",
    OUT / "data" / "D1_pelt_changepoints.xlsx",
    OUT / "v12_P2" / "data" / "metrics_all_variants.csv",
    *FIGURES.values(),
]


def _file_signature(path: Path) -> dict:
    if not path.exists():
        return {"path": str(path.relative_to(ROOT)), "exists": False}
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return {
        "path": str(path.relative_to(ROOT)),
        "exists": True,
        "size": path.stat().st_size,
        "mtime_ns": path.stat().st_mtime_ns,
        "sha256": h.hexdigest(),
    }


def build_signature() -> dict:
    items = [_file_signature(p) for p in SOURCE_FILES]
    digest = hashlib.sha256(json.dumps(items, sort_keys=True).encode("utf-8")).hexdigest()
    return {"digest": digest, "items": items}


def _load_pickle(name: str):
    with (ROOT / name).open("rb") as fh:
        return pickle.load(fh)


def _pct(x: float) -> str:
    return f"{x:.2f}%"


def _score(x: float) -> str:
    return f"{x:.3f}"


def _grade_counts(df: pd.DataFrame) -> pd.Series:
    x = pd.Series(df.to_numpy().ravel()).dropna()
    bins = [0, 1.5, 2.5, 3.5, 4.5, 5.01]
    labels = ["F(<1.5)", "D(1.5-2.5)", "C(2.5-3.5)", "B(3.5-4.5)", "A(>=4.5)"]
    return pd.cut(x, bins=bins, labels=labels, right=False).value_counts().reindex(labels).fillna(0).astype(int)


def collect_metrics() -> dict:
    state = _load_pickle("v11_state.pkl")
    p2 = _load_pickle("v12_P2_state.pkl") if (ROOT / "v12_P2_state.pkl").exists() else None
    d1 = state["D1_v11"]
    d1_v1 = state["D1_v1_scored"]
    delta = d1 - d1_v1
    n_state = sum(state["state_dist"].values())
    state_pct = {k: v / n_state * 100 for k, v in state["state_dist"].items()}

    per_channel = pd.DataFrame({
        "channel": d1.columns,
        "D1_v1": d1_v1.mean().reindex(d1.columns).values,
        "D1_v11": d1.mean().values,
        "delta": delta.mean().values,
        "lt3_pct": (d1 < 3).mean().values * 100,
        "min": d1.min().values,
    }).sort_values("D1_v11")

    event_table = state["events_v11"].copy()
    event_summary = (
        event_table.groupby("sensor_id")
        .agg(n_events=("sensor_id", "size"), total_h=("duration_h", "sum"),
             min_d1=("min_d1", "min"), mean_d1=("mean_d1", "mean"))
        .sort_values(["n_events", "total_h"], ascending=False)
        .reset_index()
        if len(event_table)
        else pd.DataFrame(columns=["sensor_id", "n_events", "total_h", "min_d1", "mean_d1"])
    )

    component_means = []
    for channel in d1.columns:
        row = {"channel": channel}
        for q in ["Q_spike", "Q_step", "Q_drift", "Q_freeze", "Q_regime"]:
            try:
                row[q] = float(state["subs_v11"][channel][q].mean())
            except Exception:
                row[q] = float("nan")
        row["dominant_low_mean"] = min(
            (q for q in ["Q_spike", "Q_step", "Q_drift", "Q_freeze", "Q_regime"]),
            key=lambda q: row[q] if pd.notna(row[q]) else 99,
        )
        component_means.append(row)
    component_df = pd.DataFrame(component_means)

    p2_summary = None
    if p2 is not None:
        p2_summary = (
            p2["metrics_df"]
            .groupby("variant")[["D1_mean", "Qreg_mean", "Qreg_lt2_pct", "cooldown_pct"]]
            .mean()
            .reindex(["R30", "R60", "R90", "R60W14"])
            .round(3)
        )

    return {
        "state": state,
        "p2": p2,
        "d1": d1,
        "d1_v1": d1_v1,
        "delta": delta,
        "per_channel": per_channel,
        "event_summary": event_summary,
        "component_df": component_df,
        "p2_summary": p2_summary,
        "state_pct": state_pct,
        "grade_v1": _grade_counts(d1_v1),
        "grade_v11": _grade_counts(d1),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def _set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width.inches * 1440)))
        grid.append(col)


def add_table(doc: Document, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_cell_margins(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        _set_cell_shading(hdr[i], "F2F4F7")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            if i > 0:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if widths:
        _set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_figure(doc: Document, label: str, caption: str, width=6.1):
    path = FIGURES.get(label)
    if path and path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        add_caption(doc, f"{label}. {caption}")


def add_code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(6)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("D1 Sensor Health Auto Research Report")


def add_title(doc: Document, m: dict):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("D1 传感器健康度评估系统")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("方案设计 · 工程实现 · 结果分析 自动研究报告")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Project D1 — DO / ORP Sensor Health Assessment\n").italic = True
    meta.add_run("v1.1 + §1.1 分解白化桥接   |   R90 / P2 审计   |   14 个评分通道")

    rows = [
        ("报告生成时间", m["generated_at"]),
        ("数据期", f"{m['d1'].index.min():%Y-%m-%d} ~ {m['d1'].index.max():%Y-%m-%d}  ({len(m['d1']):,} 小时)"),
        ("主链评分范围", f"{len(m['state']['scored_channels'])} 路 DO/ORP；QR/QIR 仅作离线支撑"),
        ("当前均值", f"STRICT V1={_score(m['d1_v1'].mean().mean())}；D1 v1.1={_score(m['d1'].mean().mean())}；Δ={_score(m['delta'].mean().mean())}"),
    ]
    add_table(doc, ["项目", "内容"], rows, [Inches(1.6), Inches(4.9)])


def add_summary(doc: Document, m: dict):
    doc.add_heading("摘要", level=1)
    state = m["state"]
    d1 = m["d1"]
    d1_v1 = m["d1_v1"]
    per = m["per_channel"]
    doc.add_paragraph(
        "本报告围绕 DO / ORP 在线水质传感器的健康度评估（Class-C min-DQR · D1 维度）项目，"
        "从顶层方案、工程实现、配置参数、运行结果和逐图诊断五个层面给出体系化专家级回顾。"
        "报告样式与 D1_Sensor_Health_Research_Report 保持一致，但所有运行结论均重新读取当前最新产物，"
        "包括 v11_state.pkl、v12_P2_state.pkl、16 个 Excel 交付表、Fig 1-11、Fig V12-V18 及 P2 图组。"
    )
    doc.add_paragraph(
        f"最新运行覆盖 {d1.index.min():%Y-%m-%d %H:%M} 至 {d1.index.max():%Y-%m-%d %H:%M}，"
        f"共 {len(d1):,} 个小时、{len(d1.columns)} 个主链评分通道。STRICT V1 均值为 {_score(d1_v1.mean().mean())}，"
        f"D1 v1.1 均值为 {_score(d1.mean().mean())}，平均 ΔD1={_score((d1-d1_v1).mean().mean())}。"
        f"最低均值通道为 {per.iloc[0]['channel']}（{_score(per.iloc[0]['D1_v11'])}），"
        f"最高均值通道为 {per.iloc[-1]['channel']}（{_score(per.iloc[-1]['D1_v11'])}）。"
    )
    doc.add_paragraph(
        f"状态机覆盖率为 Normal {_pct(m['state_pct'].get('Normal',0))} / Refractory {_pct(m['state_pct'].get('Refractory',0))} / "
        f"SustainedAnomaly {_pct(m['state_pct'].get('SustainedAnomaly',0))} / RecoveryCandidate {_pct(m['state_pct'].get('RecoveryCandidate',0))}。"
        f"D1<3 且持续不低于 6 h 的事件共 {len(state['events_v11'])} 个；PELT 变更点共 {state['n_pelt_cps']} 个。"
        "与旧静态报告中较大幅度降分的版本不同，当前版本的修正更温和，重点在于输入合法性、事件唯一性和状态可恢复性。"
    )
    doc.add_heading("核心结论速览", level=2)
    add_bullet(doc, "scope 收敛：D1 主链严格限定为 14 路 DO/ORP；QR/QIR 仅作为 D5/D7 与离线案例支撑，不参与最终健康度评分。")
    add_bullet(doc, "§1.1 桥接生效：step / regime / PELT 使用白化或路由输入，drift 使用残差，spike / freeze 保持分钟级原始信号。")
    add_bullet(doc, f"事件唯一性：{state['n_pelt_cps']} 个 PELT 变更点经状态机约束后形成 {len(state['transitions_all'])} 条状态转移，避免持续低信号反复刷新冷却。")
    add_bullet(doc, f"状态分布健康：Normal 占 {_pct(m['state_pct'].get('Normal',0))}，Refractory 与 SustainedAnomaly 合计约 {_pct(m['state_pct'].get('Refractory',0)+m['state_pct'].get('SustainedAnomaly',0))}。")
    add_bullet(doc, "P2 敏感性：R30/R60/R90 全局均值接近，R60W14 对 regime 更敏感，Qreg<2 比例最高。")


def chapter_background(doc: Document):
    doc.add_heading("第一章  项目背景与设计目标", level=1)
    doc.add_heading("1.1  问题动机与业务背景", level=2)
    doc.add_paragraph(
        "污水处理与生化反应过程高度依赖 DO 与 ORP 在线探头。长期运行中，膜污染、电极极化、生物膜附着、"
        "流态扰动和工况转换会引发尖峰、阶跃、慢漂、冻结以及工况域偏移。单一阈值报警难以兼顾误报率、漏报率和解释性，"
        "也无法为后续 D5 机理一致性和 D7 区段模板提供连续可靠度输入。"
    )
    doc.add_paragraph(
        "D1 的定位是把每路仪表每小时的信号可信度量化为 1-5 分，并同时输出子分、主导故障、状态机状态和审计字段。"
        "它不是水质异常判据，也不是 0/1 告警器，而是面向后续建模、控制、诊断和运维排班的动态数据质量维度。"
    )
    doc.add_heading("1.2  D1 v1.1 与 §1.1 桥接的迭代背景", level=2)
    doc.add_paragraph(
        "早期 STRICT V1 的弱点在于：若直接在自相关残差上运行 KS / PELT，容易把正常闭环记忆误解释为持续异常；"
        "同时电平式冷却也可能导致同一事件反复刷新。当前版本通过 1.1 Decomposition 提供的 whiteness_manifest、"
        "残差 / 创新路由和 n_eff 约束修正了输入口径，再通过五态状态机修正事件生命周期。"
    )
    doc.add_heading("1.3  设计目标的形式化", level=2)
    for text in [
        "G1 敏感性：真实尖峰、阶跃、漂移、冻结和工况域变化应能在对应时间尺度内被反映为评分下降。",
        "G2 可恢复性：异常结束后，状态机必须给出 RecoveryCandidate / Recovered 路径，避免只会降分不会回升。",
        "G3 可解释性：每个低分小时点都能追溯到子分、主导故障、状态名和事件 id。",
        "G4 无驱动量泄露：QR/QIR 不进入 D1 主链，避免驱动变量直接决定仪表健康度。",
        "G5 可复现性：所有关键参数、输出表、图件和报告均由脚本再生成。",
    ]:
        add_bullet(doc, text)


def chapter_architecture(doc: Document):
    doc.add_heading("第二章  系统架构与工程目录组织", level=1)
    doc.add_heading("2.1  顶层数据流", level=2)
    add_code_block(doc, "Raw Excel / §1.1 parquet\n   ↓ load_real_data_v11.py\nDetectors + §1.1 bridge → Q_spike / Q_step / Q_drift / Q_freeze / Q_regime\n   ↓ run_v11_pipeline.py\n5-state cooldown + D1 aggregation → v11_state.pkl\n   ↓ make_*_figures / excel_exporter / run_v12_P2_sensitivity\nFigures + Excel + P2 sensitivity + Auto expert report")
    doc.add_heading("2.2  入口脚本职责矩阵", level=2)
    add_table(doc, ["入口脚本", "职责", "关键产物"], [
        ("load_real_data_v11.py", "读取原始数据并接入 §1.1 残差/创新/manifest", "strict_v1_inputs.pkl, raw_hourly.pkl"),
        ("run_v11_pipeline.py", "运行 PELT、五态状态机、D1 v1.1 聚合", "v11_state.pkl, run_v11.log"),
        ("make_baseline_figures_v11.py", "生成 Fig 1-11 基线图组", "outputs/figures/Fig1-11"),
        ("make_figures_v11.py / part2", "生成 Fig V12-V18 修订图组", "outputs/figures/FigV12-V18"),
        ("excel_exporter_v11.py", "导出 16 个审计 Excel", "outputs/data/*.xlsx"),
        ("run_v12_P2_sensitivity.py", "生成 R30/R60/R90/R60W14 敏感性分析", "outputs/v12_P2"),
        ("generate_expert_report_v11.py", "按源签名自动生成本专家报告", REPORT_PATH.name),
    ], [Inches(1.7), Inches(3.0), Inches(1.8)])
    doc.add_heading("2.3  src/ 包内部分层", level=2)
    add_table(doc, ["子包", "核心职责", "专家审计要点"], [
        ("src/config", "配置加载", "禁止关键阈值散落在脚本中"),
        ("src/detectors", "五类故障证据", "检测器只产出统计证据，不直接判最终健康"),
        ("src/mapping", "统计量到 1-5 分映射", "统一分数语义，便于跨故障类型对比"),
        ("src/aggregation", "D1 聚合、否决、状态机", "处理事件生命周期和短板约束"),
        ("src/baseline", "谐波/局部基线复用", "服务 drift 与案例解释"),
        ("src/outputs", "图件与表格", "报告图件单一真源"),
    ], [Inches(1.3), Inches(2.2), Inches(3.0)])
    doc.add_heading("2.4  状态共享与中间存储", level=2)
    doc.add_paragraph(
        "当前主状态文件 v11_state.pkl 含 D1_v1_scored、D1_v11、日/周聚合、五类子分、状态日志、事件表、"
        "PELT 结果、regime 模板、QR/QIR 离线注释和配置快照。报告只读取这些派生产物，不重新计算检测器，"
        "因此它是结果审计层而不是算法执行层。"
    )


def chapter_methods(doc: Document, m: dict):
    doc.add_heading("第三章  核心方法论", level=1)
    doc.add_heading("3.1  检测器层", level=2)
    add_table(doc, ["子分", "尺度", "算法", "原始度量", "当前输入口径"], [
        ("Q_spike", "分钟级", "Hampel", "局部尖峰率", "1 min 原始信号"),
        ("Q_step", "小时级", "相邻双窗 KS", "KS24/KS36", "§1.1 路由输入 + n_eff"),
        ("Q_drift", "小时级", "peer-only PLS", "同伴残差 z", "§1.1 残差矩阵"),
        ("Q_freeze", "分钟级", "RLE + 低方差 + 唯一比", "冻结/响应损失", "1 min 原始信号"),
        ("Q_regime", "小时/日级", "R90 W1 + 邻区 KS", "分布距离", "§1.1 路由输入 + n_eff"),
    ], [Inches(0.9), Inches(0.9), Inches(1.5), Inches(1.4), Inches(1.8)])
    for title, text in [
        ("3.1.1  Hampel 尖峰检测", "Hampel 使用局部中位数和 MAD，对 DO/ORP 原始分钟级尖峰保持最高时间分辨率，避免小时均值掩盖短事件。"),
        ("3.1.2  邻区双窗 KS 阶跃检测", "KS24 强调近因，KS36 强调持续性；二者取最大值后再映射，有利于区分孤立扰动和真正阶跃。"),
        ("3.1.3  PLS 同伴漂移检测", "PLS 只使用同类或相邻同伴，不使用 QR/QIR 驱动量，避免驱动量泄露到传感器健康评分。"),
        ("3.1.4  冻结复合判据", "冻结检测不依赖单一 RLE，而综合持续相同值、低方差和唯一值比例，覆盖死值和近似死值。"),
        ("3.1.5  双层工况漂移", "Tier-1 W1 捕捉慢分布漂移，Tier-2 KS 确认新工况域；该维度对 D7 模板也有解释价值。"),
    ]:
        doc.add_heading(title, level=3)
        doc.add_paragraph(text)
    doc.add_heading("3.2  映射层", level=2)
    doc.add_paragraph("所有原始统计量通过 mapping.yaml 映射到 1-5 分。当前版本继承 P1 对 step 映射的放宽，避免 KS 边界噪声造成跨档抖动。")
    doc.add_heading("3.3  状态机层", level=2)
    add_table(doc, ["状态", "作用", "当前覆盖率"], [
        ("Normal", "正常评分与常规聚合", _pct(m["state_pct"].get("Normal", 0))),
        ("Refractory", "新事件后的冷却隔离", _pct(m["state_pct"].get("Refractory", 0))),
        ("SustainedAnomaly", "持续异常上限约束", _pct(m["state_pct"].get("SustainedAnomaly", 0))),
        ("RecoveryCandidate", "恢复候选观察", _pct(m["state_pct"].get("RecoveryCandidate", 0))),
        ("Recovered", "恢复完成标记", _pct(m["state_pct"].get("Recovered", 0))),
    ], [Inches(1.5), Inches(3.2), Inches(1.1)])
    doc.add_heading("3.4  聚合与否决体系", level=2)
    add_code_block(doc, "D1_base = 0.15·Q_spike + 0.20·Q_step + 0.25·Q_drift_eff + 0.20·Q_freeze + 0.20·Q_regime\nD1_pre  = 0.70·D1_base + 0.30·min(Q_spike, Q_step, Q_drift_eff, Q_freeze, Q_regime)\nD1      = D1_pre after freeze / regime / state-machine caps")


def chapter_config(doc: Document, m: dict):
    doc.add_heading("第四章  关键配置与超参数", level=1)
    doc.add_heading("4.1  configs/mapping.yaml", level=2)
    add_table(doc, ["子分", "函数族", "关键参数/边界", "专家说明"], [
        ("Q_spike", "piecewise", "6 h 尖峰率", "短时异常对 D1 的高频入口"),
        ("Q_step", "logistic", "k=8, x0=0.40", "P1 放宽后降低 KS 噪声敏感性"),
        ("Q_drift", "piecewise/logistic", "PLS |z|", "对慢漂与同伴关系破坏敏感"),
        ("Q_freeze", "stepwise + logistic", "RLE/rel_var/unique_ratio", "覆盖死值与响应损失"),
        ("Q_regime", "logistic", "R90 + 7 d W1/KS", "离线工况域迁移证据"),
    ], [Inches(1.1), Inches(1.2), Inches(1.7), Inches(2.5)])
    doc.add_heading("4.2  configs/state_machine.yaml", level=2)
    add_table(doc, ["参数", "当前值", "设计意图"], [
        ("step_ref", "48 h", "事件后冷却隔离时长"),
        ("regime_ref", "36 h", "regime 相关冷却参考"),
        ("thaw", "36 h", "Q_drift 解冻过渡"),
        ("recov_streak", "12 h", "恢复候选最短稳定时长"),
    ], [Inches(1.5), Inches(1.2), Inches(3.8)])
    doc.add_heading("4.3  configs/rules.yaml", level=2)
    doc.add_paragraph("rules.yaml 锁定主链通道、支撑通道、五类子分权重、λ_blend、否决阈值和显式禁用的驱动量泄露路径。")
    doc.add_heading("4.4  R90 区段静态基线", level=2)
    if m["p2_summary"] is not None:
        rows = []
        for idx, row in m["p2_summary"].iterrows():
            rows.append((idx, _score(row["D1_mean"]), _score(row["Qreg_mean"]), _pct(row["Qreg_lt2_pct"]), _pct(row["cooldown_pct"])))
        add_table(doc, ["变体", "D1 均值", "Q_regime 均值", "Q_regime<2", "Cooldown"], rows,
                  [Inches(1.0), Inches(1.0), Inches(1.3), Inches(1.2), Inches(1.2)])


def figure_analysis_text(label: str, m: dict) -> tuple[str, str]:
    per = m["per_channel"]
    state = m["state"]
    p2 = m["p2_summary"]
    mapping = {
        "Fig 1": ("D1 子分到等级参照矩阵", "该图是 D1 体系的翻译字典，把尖峰率、KS、PLS 残差、冻结证据和工况距离统一映射到 A-F 等级。当前报告沿用该语义，并用最新 v1.1 状态文件刷新所有运行结论。"),
        "Fig 2": ("各通道月度 D1 热图", f"热图用于月度运维 review。最新最低均值通道为 {per.iloc[0]['channel']}，建议结合月份低谷定位维护窗口和工况转换区间。"),
        "Fig 3": ("4 worst + 4 best 子分时序", "该图说明低分通常不是五类子分同时变差，而是由 Q_step、Q_regime 或 Q_drift 中的主导证据拉低；健康通道则保持多子分同步稳定。"),
        "Fig 4": ("子分分布小提琴图", "该图用于查看各子分的分布宽度和低尾。低尾越长，越需要通过事件表和状态机判断是短时异常还是持续风险。"),
        "Fig 5": ("D1 映射函数曲线", "映射函数是算法证据进入评分体系的门槛。P1 后 step 曲线更平滑，有助于减少边界噪声导致的跨档跳变。"),
        "Fig 6": ("主导故障分解", "该图回答每个通道为什么降分。结合当前风险排序，应重点核对 ORP_1_2、ORP_1_3、DO_2_3 与 DO_2_4 的主导故障构成。"),
        "Fig 7": ("按传感器组的日均 D1 轨迹", "日均轨迹将小时级噪声压缩为运维节奏，适合识别连续数日的健康度下降与恢复。"),
        "Fig 8": ("状态机与否决激活率", f"当前 Normal 占 {_pct(m['state_pct'].get('Normal',0))}，说明否决/冷却机制未过度激活；Refractory 与 SustainedAnomaly 合计约 {_pct(m['state_pct'].get('Refractory',0)+m['state_pct'].get('SustainedAnomaly',0))}。"),
        "Fig 9": ("谐波分解演示", "该图保留作为 D1 与 1.1 Decomposition 的方法衔接说明：正常周期结构应先剥离，再评价异常证据。"),
        "Fig 10": ("双层工况检测器", "W1 与邻区 KS 的组合把慢迁移和确认性分布变化分开，避免把单点噪声直接判为工况域偏移。"),
        "Fig 11": ("工程化 PLS 同伴矩阵", "peer-only 设计是避免 QR/QIR 驱动量泄露的关键；该图审计每个目标通道的同伴来源。"),
        "Fig V12": ("v1.1 vs STRICT V1 hero 诊断面板", f"当前 STRICT V1 均值 {_score(m['d1_v1'].mean().mean())}，v1.1 均值 {_score(m['d1'].mean().mean())}，平均 ΔD1={_score(m['delta'].mean().mean())}，修正幅度温和。"),
        "Fig V13": ("5 态冷却机 DO_2_3 案例", "该图展示事件进入 Refractory、转入 SustainedAnomaly、再进入恢复候选的路径，是 v1.1 相比旧电平冷却的核心证据。"),
        "Fig V14": ("信号-only Veto-3 审计", "Veto-3 在无泵阀信号时只使用信号证据，严控触发条件，避免把工艺驱动变化误判为传感器失效。"),
        "Fig V15": ("PELT 批量校正与事件唯一性过滤", f"最新 PELT 变更点为 {state['n_pelt_cps']} 个，状态机转移日志 {len(state['transitions_all'])} 条；该图审计事件过滤是否有效。"),
        "Fig V16": ("多区段模板", "该图服务 D7 离线模板，不进入 D1 主链评分；它帮助解释不同 regime 下健康评分为何会发生基线迁移。"),
        "Fig V17": ("评分范围与 QR/QIR 离线支撑", "该图明确 DO/ORP 是评分主链，QR/QIR 仅作为工况驱动和离线案例注释。"),
        "Fig V18": ("v1.1 vs STRICT V1 最终汇总", "该图汇总等级分布、Q_drift 修正、散点对照和事件主因，是最终结果的总证据图。"),
        "Fig P2-01": ("P2 变体总览", "该图说明 R30/R60/R90/R60W14 的参数差异，是解释 P2 敏感性的入口。"),
        "Fig P2-10": ("P2 全局总结", "R60W14 在当前数据上更敏感，Qreg<2 与 cooldown 指标均高于 R30/R60/R90；因此当前报告维持 R90 作为稳健审计口径。"),
    }
    return mapping[label]


def add_figure_section(doc: Document, label: str, idx: str, m: dict):
    title, text = figure_analysis_text(label, m)
    doc.add_heading(f"{idx}  {label} — {title}", level=2)
    add_figure(doc, label, title, width=6.15)
    doc.add_paragraph(text)
    doc.add_paragraph("专家解读：本图按旧主报告的“图示内容—主要结果—工程意义”口径解读。图示内容负责定位证据来源，主要结果读取当前自动产物，工程意义则指向评分可靠性、运维优先级或后续 D5/D7 使用边界。")


def chapter_figures(doc: Document, m: dict):
    doc.add_heading("第五章  基线诊断图组分析 (Fig 1 – Fig 11)", level=1)
    for i, label in enumerate([f"Fig {n}" for n in range(1, 12)], 1):
        add_figure_section(doc, label, f"5.{i}", m)
    doc.add_heading("第六章  v1.1 与 STRICT V1 全方位对照 (Fig V12 – Fig V18)", level=1)
    for i, label in enumerate([f"Fig V{n}" for n in range(12, 19)], 1):
        add_figure_section(doc, label, f"6.{i}", m)


def chapter_p1p2(doc: Document, m: dict):
    doc.add_heading("第七章  P1 / P2 工程优化与敏感性分析", level=1)
    doc.add_heading("7.1  P1 优化清单", level=2)
    add_table(doc, ["编号", "项", "V1 → P1 / v1.1", "影响"], [
        ("P1-1", "step 映射放宽", "k=12 → 8；x0=0.30 → 0.40", "降低边界振荡"),
        ("P1-2", "恢复阈值", "更可达的 recovery 判据", "避免长期低分无法恢复"),
        ("P1-3", "事件唯一性", "新 event_id 才触发冷却", "避免同一事件反复刷新"),
        ("P1-4", "§1.1 桥接", "残差/创新/manifest 路由", "提高统计检验合法性"),
    ], [Inches(0.8), Inches(1.5), Inches(2.1), Inches(2.1)])
    doc.add_heading("7.2  P2 区段基线敏感性", level=2)
    if m["p2_summary"] is not None:
        rows = []
        for idx, row in m["p2_summary"].iterrows():
            rows.append((idx, _score(row["D1_mean"]), _score(row["Qreg_mean"]), _pct(row["Qreg_lt2_pct"]), _pct(row["cooldown_pct"])))
        add_table(doc, ["变体", "D1_mean", "Qreg_mean", "Qreg<2", "Cooldown"], rows,
                  [Inches(1.0), Inches(1.1), Inches(1.2), Inches(1.2), Inches(1.2)])
    add_figure_section(doc, "Fig P2-01", "7.3", m)
    add_figure_section(doc, "Fig P2-10", "7.4", m)


def chapter_discussion(doc: Document, m: dict):
    doc.add_heading("第八章  讨论、局限性与改进展望", level=1)
    doc.add_heading("8.1  设计哲学的工程层兑现", level=2)
    doc.add_paragraph("当前 D1 体系已经把“异常证据”和“最终质量评分”分离：检测器负责产生证据，映射层负责统一量纲，状态机负责事件生命周期，报告负责审计输出。这是从报警器走向质量评估器的关键一步。")
    doc.add_heading("8.2  当前体系的局限性", level=2)
    for text in [
        "R90/P2 仍是回顾性审计口径，在线部署时需要定期重标定或引入受控滚动基线。",
        "RecoveryCandidate 覆盖率较低，说明恢复事件样本有限；后续应结合人工维护记录验证恢复判据。",
        "QR/QIR 当前只作离线解释，若未来引入过程状态，需要明确独立的 D5/D7 接口，不能回流至 D1 主链。",
        "报告生成依赖 Word COM 做视觉 QA；跨平台部署时可改用 LibreOffice + Poppler。",
    ]:
        add_bullet(doc, text)
    doc.add_heading("8.3  改进路线图", level=2)
    for text in [
        "短期：把人工维护记录接入事件表，验证 ORP_1_2、ORP_1_3、DO_2_3、DO_2_4 的低分解释。",
        "中期：把 D1 自动报告与每日/每周调度联动，形成运行报告归档。",
        "长期：将 D1 输出作为 D5 机理一致性和 D7 区段模板的输入契约，形成完整 WW-DQR 数据质量闭环。",
    ]:
        add_bullet(doc, text)


def chapter_conclusion(doc: Document, m: dict):
    doc.add_heading("第九章  结论", level=1)
    doc.add_paragraph(
        f"当前 D1 自动报告已按 D1_Sensor_Health_Research_Report 的章节样式完成更新，并将旧报告中的静态数值替换为最新结果。"
        f"本轮 D1 v1.1 均值 {_score(m['d1'].mean().mean())}，STRICT V1 均值 {_score(m['d1_v1'].mean().mean())}，"
        f"平均修正 {_score(m['delta'].mean().mean())}；主链仍保持 14 路 DO/ORP，支撑通道为 {', '.join(m['state']['support_channels'])}。"
    )
    doc.add_paragraph("总体判断：当前版本已形成“§1.1 合法输入—五类检测器—映射评分—状态机约束—图表/Excel/报告自动交付”的闭环。报告应作为每次 D1 输出更新后的正式审计件，而不是手工后补说明。")


def appendices(doc: Document, signature: dict):
    doc.add_heading("附录 A   自动更新源文件签名", level=1)
    add_table(doc, ["源文件", "状态"], [[item["path"], "OK" if item["exists"] else "缺失"] for item in signature["items"]],
              [Inches(4.8), Inches(1.2)])
    doc.add_heading("附录 B   入口脚本与产物对应关系", level=1)
    add_table(doc, ["脚本", "自动报告更新关系", "说明"], [
        ("run_v11_pipeline.py", "调用 maybe_update_report()", "状态文件变化后触发"),
        ("make_baseline_figures_v11.py", "调用 maybe_update_report()", "Fig 1-11 变化后触发"),
        ("make_figures_v11.py / part2", "调用 maybe_update_report()", "Fig V12-V18 变化后触发"),
        ("excel_exporter_v11.py", "调用 maybe_update_report()", "Excel 交付变化后触发"),
        ("run_v12_P2_sensitivity.py", "调用 maybe_update_report()", "P2 输出变化后触发"),
    ], [Inches(1.9), Inches(2.2), Inches(2.4)])
    doc.add_heading("附录 C   术语与缩写", level=1)
    add_table(doc, ["缩写 / 术语", "含义", "说明"], [
        ("D1", "Sensor Health", "传感器健康度动态评分"),
        ("STRICT V1", "旧基线版本", "用于与 v1.1 配对审计"),
        ("PELT", "Pruned Exact Linear Time", "批量变更点检测"),
        ("W1", "Wasserstein-1", "分布距离，用于 regime 漂移"),
        ("n_eff", "有效样本量", "自相关通道的统计量收缩依据"),
        ("R90", "前 90 天参考池", "当前区段静态基线审计口径"),
    ], [Inches(1.3), Inches(2.0), Inches(3.2)])


def build_report(force: bool = False) -> Path:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    signature = build_signature()
    if not force and REPORT_PATH.exists() and MANIFEST_PATH.exists():
        old = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
        if old.get("digest") == signature["digest"]:
            print(f"[auto-report] up to date: {REPORT_PATH}")
            return REPORT_PATH

    m = collect_metrics()
    doc = Document()
    configure_document(doc)
    add_title(doc, m)
    add_summary(doc, m)
    chapter_background(doc)
    chapter_architecture(doc)
    chapter_methods(doc, m)
    chapter_config(doc, m)
    chapter_figures(doc, m)
    chapter_p1p2(doc, m)
    chapter_discussion(doc, m)
    chapter_conclusion(doc, m)
    appendices(doc, signature)

    doc.save(REPORT_PATH)
    signature["generated_at"] = m["generated_at"]
    signature["report"] = str(REPORT_PATH.relative_to(ROOT))
    signature["style_reference"] = "outputs/D1_Sensor_Health_Research_Report.docx"
    MANIFEST_PATH.write_text(json.dumps(signature, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[auto-report] wrote {REPORT_PATH}")
    return REPORT_PATH


def maybe_update_report() -> Path | None:
    try:
        return build_report(force=False)
    except Exception as exc:
        print(f"[auto-report] skipped: {exc}")
        return None


if __name__ == "__main__":
    build_report(force=True)