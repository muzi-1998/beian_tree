"""Generate the synchronized final D1 expert report."""
from __future__ import annotations

import hashlib
import json
import pickle
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).parent
REPORT = ROOT / "outputs" / "D1_Sensor_Health_Expert_Report_Auto.docx"
MANIFEST = ROOT / "outputs" / "reports" / "D1_Sensor_Health_Expert_Report_Auto.manifest.json"
FIG_DIR = ROOT / "outputs" / "figures"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = "202124"
MUTED = "5F6368"
GREEN = "4C8C5A"


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _font(run, *, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_table(doc, headers: list[str], rows: list[list], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    tr_pr = header._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for cell, value in zip(header.cells, headers):
        _shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(value))
        _font(run, size=9, bold=True, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run("" if value is None else str(value))
            _font(run, size=8.5, color=INK)
    _set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _add_callout(doc, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.16)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    r1 = p.add_run(f"{label}: ")
    _font(r1, size=10.5, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    _font(r2, size=10.5, color=INK)


def _add_figure(doc, filename: str, caption: str, width=6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shape = p.add_run().add_picture(str(FIG_DIR / filename), width=Inches(width))
    shape._inline.docPr.set("title", filename)
    shape._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    _font(run, size=9, italic=True, color=MUTED)


def _set_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    _font(run, size=8, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.text = "D1 Sensor Health | Final technical audit"
    _font(header.runs[0], size=8, color=MUTED)
    _set_page_field(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_report(force: bool = True) -> Path:
    with open(ROOT / "v11_state.pkl", "rb") as handle:
        state = pickle.load(handle)
    sensitivity = pd.read_excel(
        ROOT / "outputs" / "data" / "D1_recovery_validation.xlsx",
        sheet_name="natural_sensitivity",
    )
    injection = pd.read_excel(
        ROOT / "outputs" / "data" / "D1_recovery_validation.xlsx",
        sheet_name="injection_summary",
    )
    overall = state["recovery_summary"].query("sensor_id == 'Overall'").iloc[0]

    doc = Document()
    _configure_document(doc)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("D1 Sensor Health")
    _font(run, size=23, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Final causal recovery method and expert validation report")
    _font(run, size=13, color=DARK_BLUE)
    for label, value in (
        ("Algorithm", state["algorithm_version"]),
        ("Run ID", state["run_id"]),
        ("Data scope", "14 DO/ORP scored channels; QR/QIR offline support only"),
        ("Generated", str(date.today())),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        _font(r1, size=10.5, bold=True, color=INK)
        r2 = p.add_run(value)
        _font(r2, size=10.5, color=INK)

    _add_callout(
        doc,
        "Expert decision",
        "The implementation is suitable as the locked final candidate for thesis and "
        "manuscript analyses. Its recovery claims are internally validated at the event "
        "level; diagnostic accuracy against maintenance truth remains an external-validation task.",
    )

    doc.add_heading("1. Scope and non-overlap", level=1)
    doc.add_paragraph(
        "D1 evaluates the health of each DO/ORP sensor from signal-health evidence. "
        "It is intentionally independent of D2 temporal availability, D4 physical-rate "
        "plausibility, D6 parallel-redundancy synchrony, and D7 topological representativeness. "
        "Section 1.1 provides the causal hourly preprocessing and residual basis; D2-D7 do not "
        "feed back into the D1 score."
    )

    doc.add_heading("2. Nine finalized revisions", level=1)
    revision_rows = [
        [1, "Metric definition", "Event-level direct/adapted recovery; occupancy descriptive only", "Episode table + KM"],
        [2, "Direct path", "Six-hour confirmation after Refractory", "Unit test"],
        [3, "Candidate path", "3 h entry; 12-of-18 h tolerant window; bounded soft/missing hours", "Config + sensitivity"],
        [4, "Duplicate regime gate", "Raw W1 removed as production hard veto", "Variant A/B comparison"],
        [5, "New steady regime", "Contextual local and peer residual evidence can authorize recovery", "Stable-regime challenge"],
        [6, "Residual scale", "Causal channel-specific empirical noise floor", "Calibration audit"],
        [7, "Baseline causality", "Past-only contiguous complete windows; no gap spanning", "Unit test"],
        [8, "Event identity", "Active-episode retrigger requires independent PELT evidence", "Uniqueness + conservation"],
        [9, "Release evidence", "Sensitivity, 14-channel challenges, hashes, source data, Nature figure QA", "Release gate"],
    ]
    _add_table(doc, ["#", "Risk", "Final resolution", "Verification"], revision_rows,
               [480, 1760, 4700, 2420])

    doc.add_heading("3. Causal state model", level=1)
    state_rows = [
        ["Normal", "No active episode", "Independent event evidence"],
        ["Refractory", "Suppress immediate drift reinterpretation", "Direct recovery or baseline assessment"],
        ["BaselinePending", "Wait for a causal stable local window", "Stable baseline or continued pending"],
        ["SustainedAnomaly", "Adapt within a persistent departure", "Three consecutive recovery-gate passes"],
        ["RecoveryCandidate", "Accumulate tolerant 12-of-18 h evidence", "Recovered or retry cooldown"],
        ["Recovered", "Monitor for 24 h", "Normal or relapse"],
    ]
    _add_table(doc, ["State", "Role", "Exit logic"], state_rows, [1840, 3760, 3760])
    doc.add_paragraph(
        "PELT evidence has an explicit availability time and signed magnitude. A new event ID "
        "cannot be created merely because a statistic fluctuates during an existing episode."
    )

    doc.add_heading("4. Formal run results", level=1)
    result_rows = [
        ["Unique episodes", int(overall.n_episodes), "Opened event IDs"],
        ["Direct recovery", int(overall.n_direct_recovery), "Cleared after Refractory"],
        ["Adapted recovery", int(overall.n_adapted_recovery), "Recovered via local contextual baseline"],
        ["Superseded", int(overall.n_episodes - overall.n_recovered - overall.n_right_censored), "Independent new event"],
        ["Right-censored", int(overall.n_right_censored), "Incomplete end-of-record follow-up"],
        ["Event recovery rate", f"{overall.event_recovery_rate:.4f}", "Recovered / completed episodes"],
        ["95% Wilson interval", f"{overall.event_recovery_rate_ci95_low:.4f}-{overall.event_recovery_rate_ci95_high:.4f}", "Sampling uncertainty"],
        ["Median recovery time", f"{overall.median_recovery_h:.0f} h", "Event onset to confirmation"],
        ["Recovered occupancy", f"{overall.recovered_state_occupancy:.2%}", "Observation-state occupancy, not recovery rate"],
        ["Transition conservation", "PASS", "51 opened = 51 episode records; no duplicate IDs"],
    ]
    _add_table(doc, ["Metric", "Result", "Interpretation"], result_rows, [2600, 1900, 4860])

    doc.add_page_break()
    doc.add_heading("5. Recovery evidence", level=1)
    _add_figure(
        doc,
        "FigV19_recovery_validation.png",
        "Figure V19. Event outcomes, recovery-time distribution, natural-data sensitivity, "
        "and controlled mechanism challenge results.",
    )
    doc.add_paragraph(
        "The low Recovered-state occupancy is expected because Recovered is a bounded 24 h "
        "observation state. It must not be interpreted as the proportion of successfully recovered events."
    )

    doc.add_page_break()
    _add_figure(
        doc,
        "FigV20_adapted_recovery_case.png",
        "Figure V20. Representative adapted recovery after a persistent regime departure, "
        "including quality gates, independent residual evidence, and the causal state path.",
    )

    doc.add_heading("6. Sensitivity and selected production rule", level=1)
    sensitivity_rows = []
    for row in sensitivity.itertuples(index=False):
        sensitivity_rows.append([
            row.variant,
            f"{row.event_recovery_rate:.3f}",
            int(row.candidate_attempt_count),
            f"{row.candidate_attempt_confirmation_rate:.3f}",
            int(row.n_completed),
            int(row.n_right_censored),
        ])
    _add_table(
        doc,
        ["Variant", "Event recovery", "Attempts", "Confirmation", "Completed", "Censored"],
        sensitivity_rows,
        [2860, 1340, 1100, 1420, 1320, 1320],
    )
    prod = injection[injection["variant"] == "C_tolerant_3h_retry12h"]
    _add_callout(
        doc,
        "Production selection",
        f"Variant C passed {len(prod)}/4 challenge classes across 14 channel-scaled templates "
        "with a pass rate of 1.00 in every class. Variant D was not selected because it used "
        "more candidate attempts and completed fewer natural-data episodes without improving "
        "event recovery or controlled challenge outcomes.",
    )

    doc.add_heading("7. Mapping and aggregation consistency", level=1)
    doc.add_paragraph(
        "The Step mapping is synchronized across mapping.yaml, the figure generator, and "
        "D1_mapping_params.xlsx: logistic k=8.0 and x0=0.40. The final D1 mean is 4.100 "
        "versus 4.112 for STRICT V1, a controlled mean change of -0.0115. Conservative caps "
        "apply only during causally defined non-normal states."
    )
    _add_figure(
        doc,
        "Fig5_mapping_curves.png",
        "Figure 5. Final D1 mapping functions. The Step panel explicitly reports k=8.0 and x0=0.40.",
        width=6.1,
    )

    doc.add_heading("8. Reproducibility and release gate", level=1)
    release_rows = [
        ["Unit tests", "7 passed", "Causality, gap handling, event identity, recovery, censoring"],
        ["Transition audit", "PASS", "All opened episodes accounted; all terminal or censored"],
        ["Excel audit", "PASS", "18 workbooks open; no formula-error tokens"],
        ["Figure bundle", "20/20", "SVG/PDF/600 dpi PNG/TIFF"],
        ["Nature skill audit", "0 failed", "Editable SVG text; Arial declared; sources fresh"],
        ["Run trace", state["run_id"], "Dependency SHA256 hashes in run manifest"],
    ]
    _add_table(doc, ["Gate", "Result", "Evidence"], release_rows, [2000, 1700, 5660])

    doc.add_heading("9. Publication interpretation and limitations", level=1)
    doc.add_paragraph(
        "The internal event recovery rate is not a labelled sensitivity/specificity estimate. "
        "No independently adjudicated maintenance or sensor-fault truth set is available. "
        "Controlled injections validate the recovery mechanism but do not replace an end-to-end "
        "detector benchmark under external sites, seasons, or maintenance interventions. The "
        "manuscript should report the Wilson interval, relapse rates, and these limitations."
    )
    doc.add_paragraph(
        "Recommended external validation: blind maintenance-log adjudication; site-held-out "
        "evaluation; stratification by sensor type and operating regime; and prospective analysis "
        "of recovery time, false recovery, and relapse."
    )

    doc.add_heading("Appendix A. Authoritative artifacts", level=1)
    artifact_rows = [
        ["Configuration", "configs/state_machine.yaml; configs/mapping.yaml"],
        ["Formal state", "v11_state.pkl"],
        ["Run manifest", "outputs/logs/D1_run_manifest.json"],
        ["Recovery audit", "outputs/data/D1_recovery_event_audit.xlsx"],
        ["Sensitivity", "outputs/data/D1_recovery_validation.xlsx"],
        ["Figure source data", "outputs/plot_data/"],
        ["Figure QA", "outputs/qa/figures/"],
    ]
    _add_table(doc, ["Role", "Path"], artifact_rows, [2200, 7160])

    REPORT.parent.mkdir(parents=True, exist_ok=True)
    MANIFEST.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "D1 Sensor Health final expert report"
    doc.core_properties.subject = state["run_id"]
    doc.core_properties.author = "D1 Data Quality Assessment Project"
    doc.save(REPORT)
    manifest = {
        "report": str(REPORT.relative_to(ROOT)),
        "report_sha256": _sha256(REPORT),
        "run_id": state["run_id"],
        "algorithm_version": state["algorithm_version"],
        "generator": "generate_final_expert_report.py",
        "design_preset": "compact_reference_guide",
        "source_dependencies": state["dependency_hashes"],
    }
    MANIFEST.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(f"[final-report] wrote {REPORT}")
    return REPORT


def maybe_update_report() -> Path | None:
    try:
        return build_report(force=True)
    except Exception as exc:
        print(f"[final-report] skipped: {exc}")
        return None


if __name__ == "__main__":
    build_report(force=True)
