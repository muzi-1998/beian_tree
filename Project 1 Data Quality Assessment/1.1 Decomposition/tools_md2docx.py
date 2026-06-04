"""tools_md2docx.py — convert 结果分析报告.md to a Word .docx (CJK-safe).

Tailored to this report's markdown subset: # / ## / ### headings, **bold** and
`code` inline, pipe tables, ![alt](path) images (embedded with caption),
blockquotes, ordered/unordered lists, --- rules. Run from the project root:
    python tools_md2docx.py 结果分析报告.md 结果分析报告.docx
"""
from __future__ import annotations
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
IMG_W = 6.3   # inches


def set_cjk_font(doc, latin="Microsoft YaHei", cjk="微软雅黑"):
    style = doc.styles["Normal"]
    style.font.name = latin
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cjk)


_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(p, text):
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        else:
            p.add_run(tok)


def add_table(doc, rows):
    # rows: list of list[str]; row0 = header, drop the |---| separator already
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            txt = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], txt)
            if ri == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


def add_image(doc, alt, rel):
    path = (ROOT / rel) if not Path(rel).is_absolute() else Path(rel)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        run = p.add_run()
        run.add_picture(str(path), width=Inches(IMG_W))
    else:
        r = p.add_run(f"【缺图: {rel}】"); r.italic = True
    if alt:
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(alt); cr.italic = True; cr.font.size = Pt(8.5)
        cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def split_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def convert(md_path, out_path):
    doc = Document()
    set_cjk_font(doc)
    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    i, n = 0, len(lines)
    img_re = re.compile(r"!\[(.*?)\]\((.*?)\)")
    while i < n:
        ln = lines[i]
        s = ln.strip()
        # tables
        if s.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|?\s*$", lines[i + 1].strip()):
            rows = [split_row(lines[i])]
            i += 2  # skip header + separator
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i])); i += 1
            add_table(doc, rows)
            continue
        # images
        m = img_re.match(s)
        if m:
            add_image(doc, m.group(1), m.group(2)); i += 1; continue
        # headings
        if s.startswith("# "):
            doc.add_heading(s[2:], level=0); i += 1; continue
        if s.startswith("## "):
            doc.add_heading(s[3:], level=1); i += 1; continue
        if s.startswith("### "):
            doc.add_heading(s[4:], level=2); i += 1; continue
        if s.startswith("#### "):
            doc.add_heading(s[5:], level=3); i += 1; continue
        # horizontal rule
        if s == "---":
            i += 1; continue
        # blockquote
        if s.startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, s.lstrip("> ").rstrip())
            for r in p.runs:
                r.italic = True; r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1; continue
        # ordered list
        mo = re.match(r"^(\d+)\.\s+(.*)", s)
        if mo:
            p = doc.add_paragraph(style="List Number"); add_inline(p, mo.group(2)); i += 1; continue
        # unordered list
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, s[2:]); i += 1; continue
        # blank
        if s == "":
            i += 1; continue
        # normal paragraph
        p = doc.add_paragraph(); add_inline(p, s); i += 1

    doc.save(out_path)
    print("wrote", out_path)


if __name__ == "__main__":
    md = sys.argv[1] if len(sys.argv) > 1 else "结果分析报告.md"
    out = sys.argv[2] if len(sys.argv) > 2 else "结果分析报告.docx"
    convert(ROOT / md, ROOT / out)
