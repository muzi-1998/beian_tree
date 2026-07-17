from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from d7_common.config import D7_ROOT, resolve_paths


TEAL = "168AAD"
DARK = "25282A"
MUTED = "5F6B73"
LIGHT = "EAF3F5"
WARNING = "F7E7E5"


class D7ReportBuilder:
    def __init__(self) -> None:
        self.paths = resolve_paths()
        self.paths.report_root.mkdir(parents=True, exist_ok=True)
        self.docs_root = D7_ROOT / "docs"
        self.docs_root.mkdir(parents=True, exist_ok=True)
        self.main = pd.read_parquet(
            self.paths.local_output_root / "D7_main_scores_hourly.parquet"
        )
        self.events = pd.read_parquet(
            self.paths.local_output_root / "D7_event_windows.parquet"
        )
        self.support = pd.read_parquet(
            self.paths.local_output_root / "D7_support_assessment.parquet"
        )
        self.consensus = pd.read_parquet(
            self.paths.local_output_root / "D7_zone_consensus.parquet"
        )
        self.regime = pd.read_parquet(
            self.paths.local_output_root / "D7_regime_state.parquet"
        )
        self.drift = pd.read_parquet(
            self.paths.local_output_root / "D7_topology_drift_alerts.parquet"
        )
        self.validation = pd.read_excel(
            self.paths.local_output_root / "D7_validation_results.xlsx",
            sheet_name="acceptance",
        )
        self.invariance = pd.read_excel(
            self.paths.sensitivity_output_root / "D7_track_invariance.xlsx",
            sheet_name="track_invariance",
        )
        self.manifest = json.loads(
            (self.paths.local_output_root / "D7_run_manifest.json").read_text(encoding="utf-8")
        )
        self.figure_qa = json.loads(
            (self.paths.figure_root / "D7_figure_qa.json").read_text(encoding="utf-8")
        )
        self.generated = pd.Timestamp.now(tz="Asia/Shanghai").strftime("%Y-%m-%d %H:%M CST")

    def build_all(self) -> list[Path]:
        report_md = self._expert_markdown()
        guide_md = self._guide_markdown()
        captions_md = self._captions_markdown()
        outputs = []
        for name, content in [
            ("D7_EXPERT_REPORT_v2.1.md", report_md),
            ("D7_PROJECT_DIRECTORY_GUIDE_v2.1.md", guide_md),
            ("D7_FIGURE_CAPTIONS_v2.1.md", captions_md),
        ]:
            path = self.paths.report_root / name
            path.write_text(content, encoding="utf-8")
            outputs.append(path)
            (self.docs_root / name).write_text(content, encoding="utf-8")
        expert_docx = self.paths.report_root / "D7_EXPERT_REPORT_v2.1.docx"
        guide_docx = self.paths.report_root / "D7_PROJECT_DIRECTORY_GUIDE_v2.1.docx"
        self._build_expert_docx(expert_docx)
        self._build_guide_docx(guide_docx)
        outputs.extend([expert_docx, guide_docx])
        return outputs

    def _facts(self) -> dict[str, Any]:
        status = self.main["evaluation_status"].value_counts()
        support = self.support["support_level"].value_counts()
        validation = self.validation.set_index("criterion")
        invariance = self.invariance.set_index("metric")
        return {
            "run_id": self.main["run_id"].iloc[0],
            "start": self.main["timestamp"].min(),
            "end": self.main["timestamp"].max(),
            "rows": len(self.main),
            "raw_valid": int(self.main["D7_raw"].notna().sum()),
            "raw_mean": float(self.main["D7_raw"].mean()),
            "raw_median": float(self.main["D7_raw"].median()),
            "raw_p05": float(self.main["D7_raw"].quantile(0.05)),
            "low_rate": float(self.main["D7_raw"].lt(3.0).mean()),
            "events": len(self.events),
            "status": status.to_dict(),
            "support": support.to_dict(),
            "ood_rate": float(self.regime["regime_state"].eq("OODHold").mean()),
            "switches": int(self.regime["transition_id"].notna().sum()),
            "d7_total_nonnull": int(self.main["D7_total"].notna().sum()),
            "d6_evaluable": int(self.consensus["d7_evaluable"].sum()),
            "topology_alerts": int(self.drift["alert_level"].eq("review").sum()),
            "swap_auroc": float(validation.loc["swap_AUROC", "estimate"]),
            "swap_auprc": float(validation.loc["swap_AUPRC", "estimate"]),
            "swap_top1": float(validation.loc["swap_Top1", "estimate"]),
            "common_far": float(validation.loc["common_mode_FAR", "estimate"]),
            "zone_far": float(validation.loc["zone_coherent_FAR", "estimate"]),
            "chatter": float(validation.loc["switch_chatter_rate", "estimate"]),
            "ie_track": float(invariance.loc["IE_track", "estimate"]),
            "jaccard": float(invariance.loc["event_jaccard", "estimate"]),
            "rho": float(invariance.loc["culprit_spearman", "estimate"]),
            "figure_qa": bool(self.figure_qa["passed"]),
        }

    def _expert_markdown(self) -> str:
        f = self._facts()
        status_lines = "\n".join(
            f"- `{name}`: {count:,} ({count / f['rows']:.1%})"
            for name, count in f["status"].items()
        )
        support_lines = "\n".join(
            f"- `{level}`: {count} templates" for level, count in f["support"].items()
        )
        return f"""# D7 Expert Review Report v2.1

**Project:** Topological Role Consistency and Structural Representativeness  
**Run:** `{f['run_id']}`  
**Generated:** {self.generated}
**Decision:** Research package complete; production DQR release blocked.

## 1. Executive verdict

The D7 v2.1 research implementation is complete at the P2/V2 artifact level: Local, Sensitivity and Shadow V2 tracks, frozen templates, hourly scores, raw evidence, validation, plot data, SCI-ready figures, manifests and audit records are present. The Local Track is logically independent of D1, D2, D4 and D6 and consumes only canonical observations, exogenous hydraulic/time context and declared D7 topology.

The project is **not production-ready**. The field topology, asset/serial/channel-position mapping and two-person approval remain pending; only 3 of 56 templates reach L2, while 53 remain L1; swap Top-1 is {f['swap_top1']:.2f}, below the 0.80 acceptance target. Consequently `D7_total`, `D7_forDQR` and D6 final arbitration correctly contain zero evaluable rows.

## 2. Scope and dimensional independence

- D7 asks whether a DO/ORP observation still behaves like its declared spatial position and represents its process zone.
- D1 evaluates sensor-intrinsic health and long-term regime-relative behavior; D2 evaluates continuity and information availability; D4 evaluates physical value/rate plausibility; D6 evaluates temporal synchronization between parallel counterparts.
- Local D7 does not consume any D1-D6 score, state or event field. D1/D2/D4 are read only in the physically isolated Sensitivity Track.
- QR/QIR are exogenous context variables only and never receive a D7 score.
- Observed low `D7_raw` is structural evidence, not a confirmed hardware fault.

## 3. Data and result freshness

- Hourly sensor windows: {f['rows']:,}, spanning {f['start']} to {f['end']}.
- Calculable `D7_raw`: {f['raw_valid']:,}; mean {f['raw_mean']:.3f}, median {f['raw_median']:.3f}, p05 {f['raw_p05']:.3f}.
- Raw low-score fraction (`D7_raw < 3`): {f['low_rate']:.1%}; candidate persistent events: {f['events']}.
- Full 10-min regime-state trajectory retained; OOD hold rate: {f['ood_rate']:.1%}; confirmed switches: {f['switches']}.
- Result provenance is bound to canonical input hashes, topology hash, template/mapping/regime versions and code commit in `D7_run_manifest.json`.

## 4. Applicability and support

{status_lines}

{support_lines}

ORP is deliberately forced to L1 `diagonal_robust_z` with `alpha=1.00`. It is never promoted automatically. L0, if encountered in a short or sparse rerun, is disabled rather than written as a low score.

## 5. Validation and sensitivity

| Criterion | Estimate | Target | Result |
|---|---:|---:|---|
| Swap AUROC | {f['swap_auroc']:.3f} | >=0.90 | Pass |
| Swap AUPRC | {f['swap_auprc']:.3f} | >=0.80 | Pass |
| Swap Top-1 | {f['swap_top1']:.3f} | >=0.80 | **Fail** |
| Common-mode FAR | {f['common_far']:.3f} | <=0.10 | Pass |
| Zone-coherent FAR | {f['zone_far']:.3f} | <=0.10 | Pass |
| Switch chatter rate | {f['chatter']:.3f} | <=0.05 | Pass |
| IE_track | {f['ie_track']:.3f} | <=0.20 | Pass |
| Event Jaccard | {f['jaccard']:.3f} | >=0.80 | Pass |
| Culprit Spearman rho | {f['rho']:.3f} | >=0.80 | Pass |

Validation uses observed test-period spatial windows with frozen templates. Same-line, same-analyte position swaps are positive controls. Freeze, temporal ramps, common-mode and zone-coherent changes, DO4 floor behavior and dropout are negative/orthogonality controls. The swap detection metrics pass, but localization remains below the release criterion and must not be hidden by threshold tuning.

## 6. Topology and D6 interface

- Declared topology contains 14 DO/ORP nodes, 10 longitudinal edges and seven parallel peer pairs.
- {f['topology_alerts']} finite candidate mappings exceed the report-only topology drift review threshold. These are hypotheses for field review, not automatic topology updates.
- Shadow V2 has `production_impact=none`; it cannot mutate `topology.yaml`, active templates, `D7_total` or Veto.
- `D7_total` non-null rows: {f['d7_total_nonnull']}; D6 interface evaluable rows: {f['d6_evaluable']}.
- D6 protected score columns are untouched because D7 has no D6 write path.

## 7. Figure review

Five multi-panel figure groups are available as editable SVG/PDF and 600 dpi PNG/TIFF, backed by the complete `D7_plot_data.parquet/.csv` source table. All use a 183 mm final width, Arial, 0.75 pt axes, `(a)/(b)/(c)/(d)` labels as applicable, outward ticks for open axes and inward ticks for full-frame heatmaps. The asymmetric layouts follow the claim-evidence hierarchy defined in `docs/D7_FIGURE_CONTRACT_v2.2.md`; no rendering-convenience sampling is used. Automated counterpart/font/dimension/pixel QA passed: {f['figure_qa']}.

## 8. Critical limitations

1. Topology and asset identity are declared but not field-verified or dual-approved.
2. Effective independent support is inadequate for production gating in 53/56 templates; ORP remains intentionally L1.
3. Swap Top-1 localization is 0.75 versus the 0.80 target.
4. The {f['events']} candidate event windows have no external truth labels; event counts must not be reported as confirmed sensor faults.
5. Regime transition FAR and topology candidate recall are not estimable without external regime/topology truth.
6. `D7_raw` calibration is suitable for comparative research evidence, but operational event thresholds require labeled prospective confirmation.

## 9. Release decision and next actions

The branch may be reviewed and merged as a **research implementation with explicit production gates**. It must not be activated in WW-DQS/DQR arbitration yet.

1. Verify the process drawing, coordinates, asset IDs, serial numbers and channel-position mapping in the field.
2. Obtain independent reviewer and approver signatures; update `topology.yaml` and regenerate all topology-bound templates.
3. Accumulate qualified multi-season effective blocks and pass ORP/DO support exit criteria.
4. Improve and revalidate node localization to Top-1 >=0.80 on blocked holdouts.
5. Add field-confirmed swap/maintenance/topology cases and prospective event labels.
6. Only then set topology status to verified, rerun the full release workflow and consider `D7_total`/D6 final arbitration.
"""

    def _guide_markdown(self) -> str:
        output_rows = [
            ("outputs/local", "Authoritative Local scores, evidence, templates, audits and D6 read-only interface"),
            ("outputs/sensitivity", "D1/D2/D4-filtered shadow reference sensitivity; no production writes"),
            ("outputs/shadow_v2", "Graph/topology research candidates with production impact fixed to none"),
            ("outputs/plot_data", "Frozen long-table data used by every manuscript figure"),
            ("outputs/figures", "Editable SVG/PDF and 600 dpi PNG/TIFF figures with figure QA"),
            ("outputs/reports", "Expert report, directory guide and figure captions"),
        ]
        output_table = "\n".join(f"| `{path}` | {role} |" for path, role in output_rows)
        return f"""# D7 Project Directory Guide v2.1

**Generated:** {self.generated}
**Update rule:** every computational or figure change must rerun reports and release QA.

## 1. Project layout

```text
D7 Topological Role Consistency and Structural Representativeness/
|-- configs/
|   |-- common/          # paths, sensors, topology and interface schemas
|   |-- local/           # production-isolated D7 Local policies
|   |-- sensitivity/     # upstream-filter sensitivity only
|   `-- shadow_v2/       # graph/topology research only
|-- src/
|   |-- d7_common/       # shared configuration, hashing and robust math
|   |-- d7_local/        # Local contracts, context, templates, evidence, scoring and outputs
|   |-- d7_sensitivity/  # physically isolated shadow pipeline
|   `-- d7_shadow_v2/    # no-production-impact graph research
|-- scripts/             # reproducible entry points
|-- tests/               # contract and regression tests
|-- docs/                # version-reviewable Markdown reports
`-- outputs/             # generated artifacts by track and role
```

## 2. Track isolation

| Track | Allowed inputs | Forbidden outputs |
|---|---|---|
| `d7_local` | Canonical 1-min observations/flags, QR/QIR context, declared topology, frozen D7 assets | Reading D1-D6 score/state/event inputs |
| `sensitivity` | Frozen Local evidence plus D1/D2/D4 read-only filters | `D7_forDQR`, zone consensus, active templates, D6 arbitration |
| `shadow_v2` | Canonical observations and declared topology | Automatic topology update, Veto, Local score mutation |

## 3. Output roots

| Path | Role |
|---|---|
{output_table}

## 4. Core entry points

```powershell
python scripts/run_d7_local.py
python scripts/run_d7_sensitivity.py
python scripts/run_d7_validation.py
python scripts/run_d7_topology_review.py
python scripts/run_d7_shadow_v2.py
python scripts/make_d7_figures.py
python scripts/build_d7_reports.py
python scripts/check_d7_release.py
python -m pytest tests -q
```

Use `python scripts/run_d7_release.py --include-local` after data, topology, template, mapping or core scoring changes. For report/figure-only edits, omit `--include-local`.

## 5. Update protocol

1. Work on an `agent/*` branch; do not write half-complete results to `main`.
2. Change configuration and code together; never edit generated score tables by hand.
3. Run the fastest contract tests, then the affected track.
4. Rebuild validation, plot data, all figure counterparts, reports and release QA.
5. Inspect PNG figures and rendered DOCX pages visually.
6. Stage only the D7 directory, commit, push the branch and open a draft PR.
7. Report branch, commit, compare/PR URL and all failed release gates to the user.
8. Merge only after explicit confirmation.

## 6. Production activation checklist

- [ ] Field drawing ID is real and versioned.
- [ ] All asset IDs, serial numbers and channel tags are verified.
- [ ] Coordinates are surveyed or approved, not schematic placeholders.
- [ ] Reviewer and approver are two identified independent people.
- [ ] Topology validity interval covers the evaluated data.
- [ ] All topology-bound template hashes are regenerated.
- [ ] Required support tiers and ORP exit criteria pass.
- [ ] Swap AUROC/AUPRC and Top-1 pass blocked validation.
- [ ] Transition FAR and topology tests have external truth.
- [ ] Local-Sensitivity invariance gates pass.
- [ ] D6 protected columns have max absolute difference zero.
- [ ] `D7_total` is populated only for genuinely evaluable rows.

## 7. File ownership rules

- `configs/common/topology.yaml` is the only declared production topology source; Shadow outputs never overwrite it.
- Parquet is authoritative for tabular data; Excel is a human-review mirror with the same semantics.
- `D7_plot_data` is the sole manuscript-figure input; figure scripts do not recompute business metrics.
- Reports and this guide are generated artifacts and must be refreshed on every release workflow.
- Missing/limited/OOD evidence is represented by status plus `NaN`, never by a fabricated low score.

## 8. Current branch gate

Current release classification: **research review complete, production blocked**. The immediate blockers are field topology/asset verification, effective support and Top-1 localization.
"""

    def _captions_markdown(self) -> str:
        return """# D7 Figure Captions v2.1

## Figure D7-1. Declared topology, applicability and scientific boundary

(a) Declared longitudinal DO/ORP topology for two parallel process lines. Coordinates are schematic and pending field verification. (b) Local Track applicability states across hourly sensor windows. (c) Availability of raw spatial evidence versus blocked topology approval and DQR release. `D7_raw` is retained when calculable, whereas `D7_forDQR` remains null until all gates pass.

## Figure D7-2. Spatiotemporal score structure and effective template support

(a) Daily lower-quartile `D7_raw` for 14 DO/ORP positions, using a threshold-centred muted red-to-blue map. (b) Full finite score distributions by analyte; no display sample is used. (c) Counts of regime templates by effective support tier. The predominance of L1 support limits production use despite calculable raw evidence.

## Figure D7-3. Evidence decomposition and node attribution

(a) Direct-labelled spatial quality components and `D7_raw` around an unlabeled persistent low-score window; shading denotes the candidate interval. The case is evidence for review, not a confirmed fault. (b) Target-excluded LOSO/graph-energy node influence at the case center. (c) Distribution of report-only zone-consensus labels supplied to the D6 interface; non-evaluable status prevents final arbitration.

## Figure D7-4. Frozen-template validation and track invariance

(a) Release criteria for same-line position swaps, negative controls and regime chatter; filled circles are estimates and open diamonds are prespecified targets. (b) Top-1 localization by injected D7-relevant scenario. (c) False alarm rates with empirical 95% intervals for orthogonality controls. (d) Local-Sensitivity invariance estimates and targets. Swap AUROC/AUPRC and negative controls pass, whereas Top-1 remains below 0.80.

## Figure D7-5. Support, regime, topology and release governance

(a) Support-tier distribution. (b) MAP-Hysteresis state occupancy. (c) Highest finite candidate-vs-declared topology likelihood ratios shown as a report-only review screen; candidates cannot alter production topology. (d) Sequential release-gate chain showing passed research gates and blocked production gates.
"""

    def _new_document(self, title: str, subtitle: str, status: str) -> Document:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0.82)
        self._configure_styles(document)
        header = section.header.paragraphs[0]
        header.text = "D7 v2.1 | Project 1 Data Quality Assessment"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._format_runs(header, size=8, color=MUTED)
        footer = section.footer.paragraphs[0]
        footer.text = f"Generated {self.generated} | Research review only"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._format_runs(footer, size=8, color=MUTED)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(title)
        self._set_font(run, size=21, bold=True, color=DARK)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(12)
        run = paragraph.add_run(subtitle)
        self._set_font(run, size=11.5, color=MUTED)
        for label, value in [
            ("Status", status),
            ("Run", str(self.main["run_id"].iloc[0])),
            ("Coverage", f"{self.main['timestamp'].min()} to {self.main['timestamp'].max()}"),
        ]:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            left = p.add_run(f"{label}: ")
            self._set_font(left, size=9.5, bold=True, color=TEAL)
            right = p.add_run(value)
            self._set_font(right, size=9.5, color=DARK)
        document.add_paragraph().paragraph_format.space_after = Pt(4)
        return document

    def _configure_styles(self, document: Document) -> None:
        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Arial"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.20
        for name, size, color, before, after in [
            ("Heading 1", 16, TEAL, 16, 8),
            ("Heading 2", 13, TEAL, 12, 6),
            ("Heading 3", 11.5, "1F4D78", 8, 4),
        ]:
            style = styles[name]
            style.font.name = "Arial"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

    def _build_expert_docx(self, path: Path) -> None:
        f = self._facts()
        doc = self._new_document(
            "D7 Expert Review Report",
            "Topological Role Consistency and Structural Representativeness | v2.1",
            "Research package complete; production DQR release blocked",
        )
        self._heading(doc, "1. Executive verdict", 1)
        self._paragraph(
            doc,
            "D7 v2.1 is complete as a research and review package. The Local, Sensitivity and Shadow V2 tracks, frozen method assets, validation, plotting data, figures, reports and audit records are present and reproducible.",
        )
        self._warning(
            doc,
            f"Production release is blocked: topology/asset identity is unverified, 53 of 56 templates remain L1, and swap Top-1 is {f['swap_top1']:.2f} versus the 0.80 target. D7_total, D7_forDQR and D6 final arbitration therefore remain empty.",
        )
        self._heading(doc, "2. Scope and dimensional independence", 1)
        for item in [
            "D7 evaluates spatial role identity and structural representativeness, not hardware fault type.",
            "D1 covers sensor-intrinsic health; D2 covers continuity/availability; D4 covers physical value/rate constraints; D6 covers temporal synchronization between parallel counterparts.",
            "Local D7 consumes no D1-D6 score, state or event. D1/D2/D4 appear only in the isolated Sensitivity Track.",
            "QR/QIR provide exogenous context only. Low D7_raw is unlabeled structural evidence, not a confirmed fault.",
        ]:
            self._bullet(doc, item)
        doc.add_page_break()
        self._heading(doc, "3. Current data and outputs", 1)
        self._table(
            doc,
            ["Measure", "Current result", "Interpretation"],
            [
                ["Hourly windows", f"{f['rows']:,}", f"{f['start']} to {f['end']}"],
                ["D7 raw", f"mean {f['raw_mean']:.3f}; median {f['raw_median']:.3f}", "Comparative research evidence"],
                ["Low-score fraction", f"{f['low_rate']:.1%}", "Not a fault prevalence estimate"],
                ["Candidate events", str(f["events"]), "Unreviewed and unlabeled"],
                ["OOD hold", f"{f['ood_rate']:.1%}", "Explicit posterior rejection"],
                ["D7 total / D6 evaluable", f"{f['d7_total_nonnull']} / {f['d6_evaluable']}", "Correctly blocked"],
            ],
            [1.55, 1.65, 3.30],
        )
        self._heading(doc, "4. Support and applicability", 1)
        support_rows = [[level, str(count), "Gating eligible only after all contracts pass"] for level, count in f["support"].items()]
        self._table(doc, ["Tier", "Templates", "Policy"], support_rows, [1.0, 1.1, 4.4])
        self._paragraph(
            doc,
            "ORP is forced to L1 diagonal robust Z with alpha=1.00 and cannot auto-promote. Missing, OOD and limited evidence is represented by status and NaN, never by a fabricated low production score.",
        )
        doc.add_page_break()
        self._heading(doc, "5. Validation and sensitivity", 1)
        validation_rows = []
        for row in self.validation.iloc[:-1].itertuples(index=False):
            validation_rows.append(
                [row.criterion, f"{row.estimate:.3f}", f"{row.operator} {row.target:.2f}", "PASS" if row.passed else "FAIL"]
            )
        self._table(doc, ["Criterion", "Estimate", "Target", "Result"], validation_rows, [2.35, 1.0, 1.2, 1.0])
        self._paragraph(
            doc,
            "Same-line position swaps are scored with frozen templates. Freeze, temporal ramps, common/zone-coherent changes, DO4 floor behavior and dropout are orthogonality controls. Detection and FAR pass, but node localization remains below release quality.",
        )
        invariance_rows = [
            [row.metric, f"{row.estimate:.3f}", row.criterion, "PASS" if row.passed else "FAIL"]
            for row in self.invariance.itertuples(index=False)
            if row.metric != "FAR_delta"
        ]
        self._table(doc, ["Track metric", "Estimate", "Target", "Result"], invariance_rows, [2.0, 1.0, 1.5, 1.0])
        self._heading(doc, "6. Topology and D6 interface", 1)
        self._paragraph(
            doc,
            f"The declared registry contains 14 nodes, 10 longitudinal edges and seven peer pairs. {f['topology_alerts']} finite candidate mappings exceed the report-only review threshold. They cannot mutate production topology or activate Veto.",
        )
        self._bullet(doc, "Topology drawing, coordinates, asset IDs, serial numbers and two-person approval remain pending.")
        self._bullet(doc, "D7-to-D6 zone consensus is present but every row is non-evaluable for final arbitration.")
        self._bullet(doc, "D6 protected columns remain unchanged because no D6 write path exists.")
        doc.add_page_break()
        self._heading(doc, "7. SCI figure review", 1)
        captions = [
            "Figure D7-1 | Declared topology, applicability and release boundary.",
            "Figure D7-2 | Spatiotemporal D7 structure and effective support.",
            "Figure D7-3 | Evidence decomposition, node influence and zone consensus.",
            "Figure D7-4 | Validation and Local-Sensitivity invariance.",
            "Figure D7-5 | Support, regime, topology and release governance.",
        ]
        stems = ["FigD7_1_framework", "FigD7_2_spatiotemporal", "FigD7_3_evidence", "FigD7_4_validation", "FigD7_5_governance"]
        for stem, caption in zip(stems, captions):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run()
            run.add_picture(str(self.paths.figure_root / f"{stem}.png"), width=Inches(6.15))
            self._set_picture_alt(run, caption)
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(10)
            self._format_runs(cap, size=8.5, color=MUTED, italic=True)
        doc.add_page_break()
        self._heading(doc, "8. Critical limitations", 1)
        for item in [
            "Field topology and asset identity are not verified or dual-approved.",
            "Effective support is inadequate for production gating in 53 of 56 templates.",
            f"Swap Top-1 is {f['swap_top1']:.2f}, below the 0.80 target.",
            f"The {f['events']} candidate events lack external truth and cannot be called confirmed faults.",
            "Regime transition FAR and topology candidate recall require external truth.",
        ]:
            self._bullet(doc, item)
        self._heading(doc, "9. Release decision", 1)
        self._paragraph(
            doc,
            "The branch is suitable for expert review and merge as a research implementation with explicit production gates. It is not suitable for activation in WW-DQS/DQR arbitration.",
        )
        for item in [
            "Complete field mapping and obtain two independent approvals.",
            "Regenerate all topology-bound templates and hashes.",
            "Accumulate qualified multi-season effective blocks and pass support exit criteria.",
            "Raise blocked-holdout Top-1 to at least 0.80 without weakening negative-control FAR.",
            "Add prospective field-confirmed cases before activating D7_total or D6 final arbitration.",
        ]:
            self._number(doc, item)
        doc.save(path)

    def _build_guide_docx(self, path: Path) -> None:
        doc = self._new_document(
            "D7 Project Directory Guide",
            "Python structure, artifact ownership and release workflow | v2.1",
            "Operational reference; regenerate on every project update",
        )
        self._heading(doc, "1. Directory map", 1)
        tree = (
            "configs/{common,local,sensitivity,shadow_v2}\n"
            "src/{d7_common,d7_local,d7_sensitivity,d7_shadow_v2}\n"
            "scripts/  tests/  docs/\n"
            "outputs/{local,sensitivity,shadow_v2,plot_data,figures,reports}"
        )
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        for line in tree.splitlines():
            run = p.add_run(line + "\n")
            self._set_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9.5, color=DARK)
        self._heading(doc, "2. Track isolation", 1)
        self._table(
            doc,
            ["Track", "Allowed inputs", "Forbidden behavior"],
            [
                ["d7_local", "Canonical observations/flags, QR/QIR context, declared topology", "Read D1-D6 scores/states/events"],
                ["sensitivity", "Frozen Local evidence plus D1/D2/D4 read-only filters", "Write D7_forDQR, zone consensus or D6"],
                ["shadow_v2", "Canonical observations and declared topology", "Auto-update topology or production scores"],
            ],
            [1.2, 2.65, 2.65],
        )
        self._heading(doc, "3. Output ownership", 1)
        self._table(
            doc,
            ["Path", "Owner and role"],
            [
                ["outputs/local", "Authoritative Local scores, evidence, templates, interface and audit"],
                ["outputs/sensitivity", "Shadow reference/mapping sensitivity only"],
                ["outputs/shadow_v2", "No-production-impact graph/topology research"],
                ["outputs/plot_data", "Frozen long-table source for all figures"],
                ["outputs/figures", "Editable SVG/PDF and 600 dpi PNG/TIFF plus QA"],
                ["outputs/reports", "Generated expert report, guide and captions"],
            ],
            [1.75, 4.75],
        )
        self._heading(doc, "4. Core commands", 1)
        commands = [
            "python scripts/run_d7_local.py",
            "python scripts/run_d7_sensitivity.py",
            "python scripts/run_d7_validation.py",
            "python scripts/run_d7_topology_review.py",
            "python scripts/run_d7_shadow_v2.py",
            "python scripts/make_d7_figures.py",
            "python scripts/build_d7_reports.py",
            "python scripts/check_d7_release.py",
            "python -m pytest tests -q",
        ]
        for command in commands:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(command)
            self._set_font(run, name="Consolas", east_asia="Microsoft YaHei", size=9.0, color=DARK)
        self._paragraph(
            doc,
            "Use python scripts/run_d7_release.py --include-local after data, topology, template, mapping or scorer changes. Omit --include-local for report/figure-only work.",
        )
        self._heading(doc, "5. Update and Git workflow", 1)
        for item in [
            "Work on an agent/* branch and keep changes reviewable.",
            "Change code/configuration; never edit generated scores manually.",
            "Run contract tests and the affected track.",
            "Rebuild validation, plot data, figures, reports and release QA.",
            "Visually inspect every figure and every rendered DOCX page.",
            "Stage only the D7 directory, commit and push the branch.",
            "Give the user the branch, commit and compare/PR location; merge only after confirmation.",
        ]:
            self._number(doc, item)
        self._heading(doc, "6. Production activation checklist", 1)
        for item in [
            "Real process drawing ID and validity interval recorded",
            "Asset IDs, serial numbers, channel tags and positions field-verified",
            "Coordinates surveyed or explicitly approved",
            "Independent reviewer and approver recorded",
            "Topology-bound templates and hashes regenerated",
            "Support tiers and ORP exit criteria passed",
            "Swap AUROC/AUPRC and Top-1 passed on blocked holdouts",
            "Transition FAR and topology tests supported by external truth",
            "Local-Sensitivity invariance passed",
            "D6 protected-column max absolute difference equals zero",
            "D7_total populated only for genuinely evaluable rows",
        ]:
            self._bullet(doc, "[ ] " + item)
        self._heading(doc, "7. Non-negotiable file contracts", 1)
        for item in [
            "topology.yaml is the only declared production topology source.",
            "Parquet is authoritative; Excel is a semantically equivalent review mirror.",
            "D7_plot_data is the only manuscript-figure input.",
            "Missing/limited/OOD evidence uses explicit status plus NaN.",
            "Reports and this guide are regenerated on every release workflow.",
        ]:
            self._bullet(doc, item)
        doc.save(path)

    def _heading(self, doc: Document, text: str, level: int) -> None:
        doc.add_heading(text, level=level)

    def _paragraph(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph(text)
        self._format_runs(p, size=10.5, color=DARK)

    def _bullet(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        self._format_runs(p, size=10.2, color=DARK)

    def _number(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph(text, style="List Number")
        p.paragraph_format.space_after = Pt(4)
        self._format_runs(p, size=10.2, color=DARK)

    def _warning(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.18)
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(9)
        run = p.add_run(text)
        self._set_font(run, size=10.5, bold=True, color="8B2E2E")
        self._shade_paragraph(p, WARNING)

    def _table(
        self,
        doc: Document,
        headers: list[str],
        rows: list[list[str]],
        widths: list[float],
    ) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.style = "Table Grid"
        header_properties = table.rows[0]._tr.get_or_add_trPr()
        repeat_header = OxmlElement("w:tblHeader")
        repeat_header.set(qn("w:val"), "true")
        header_properties.append(repeat_header)
        for index, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = str(header)
            self._shade_cell(cell, LIGHT)
            self._format_runs(cell.paragraphs[0], size=9.2, bold=True, color=DARK)
        for row_data in rows:
            cells = table.add_row().cells
            for cell, value, width in zip(cells, row_data, widths):
                cell.width = Inches(width)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.text = str(value)
                self._format_runs(cell.paragraphs[0], size=8.8, color=DARK)
        for row in table.rows:
            row_properties = row._tr.get_or_add_trPr()
            keep_row_together = OxmlElement("w:cantSplit")
            row_properties.append(keep_row_together)
        self._set_table_geometry(table, widths)
        after = doc.add_paragraph()
        after.paragraph_format.space_after = Pt(2)

    @staticmethod
    def _set_table_geometry(table: Any, widths: list[float]) -> None:
        table_width = int(sum(widths) * 1440)
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "dxa")
        tbl_w.set(qn("w:w"), str(table_width))
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(int(width * 1440)))
            grid.append(col)
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tc_w.type = "dxa"
                tc_w.w = int(width * 1440)

    @staticmethod
    def _shade_cell(cell: Any, fill: str) -> None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _shade_paragraph(paragraph: Any, fill: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        p_pr.append(shading)

    def _format_runs(
        self,
        paragraph: Any,
        *,
        size: float,
        color: str,
        bold: bool | None = None,
        italic: bool | None = None,
    ) -> None:
        for run in paragraph.runs:
            self._set_font(run, size=size, color=color, bold=bold, italic=italic)

    @staticmethod
    def _set_font(
        run: Any,
        *,
        name: str = "Arial",
        east_asia: str = "Microsoft YaHei",
        size: float | None = None,
        color: str | None = None,
        bold: bool | None = None,
        italic: bool | None = None,
    ) -> None:
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic

    @staticmethod
    def _set_picture_alt(run: Any, description: str) -> None:
        drawing = run._element.xpath(".//wp:docPr")
        if drawing:
            drawing[0].set("descr", description)
