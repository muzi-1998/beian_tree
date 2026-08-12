from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
D4 = ROOT / 'Project 1 Data Quality Assessment' / 'D4 Parallel-redundancy Temporal Consistency'
FIG = D4 / 'outputs' / 'figures'
OUT = ROOT / 'deliverables'
OUT.mkdir(exist_ok=True)
DOCX = OUT / '260812_D4_绘图组件_顶刊SCI发表水平_逐图审查_修改与新增图件建议_专家版.docx'

FIGURES = [
    ('FigD4_1_paired_residual_consistency.png', 'FigD4_1  Paired de-periodised residual consistency'),
    ('FigD4_2_subscore_contribution.png', 'FigD4_2  Subscore contribution and D4 distribution'),
    ('FigD4_3_trend_slope_scatter.png', 'FigD4_3  Robust trend-slope concordance'),
    ('FigD4_4_raw_score_heatmap.png', 'FigD4_4  Weekly D4_raw heatmap'),
    ('FigD4_5_status_evaluability.png', 'FigD4_5  Pair status and evaluability'),
    ('FigD4_6_context_independence.png', 'FigD4_6  Context independence diagnostic'),
    ('FigD4_7_validation_roc_pr.png', 'FigD4_7  Mechanism validation: ROC and PR'),
    ('FigD4_8_validation_ablation.png', 'FigD4_8  Component ablation and false-alarm control'),
]

for fn, _ in FIGURES:
    if not (FIG / fn).exists():
        raise FileNotFoundError(FIG / fn)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def borders(cell, color='D9E2F3', sz='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = 'w:' + edge
        el = tcBorders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcBorders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),sz); el.set(qn('w:color'),color)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold; r.font.size = Pt(size); r.font.name = 'Arial'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans CJK SC')
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    borders(cell)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j,h in enumerate(headers):
        set_cell_text(t.rows[0].cells[j], h, True, 'FFFFFF', 8.5)
        shade(t.rows[0].cells[j], '17365D')
        if widths: t.rows[0].cells[j].width = Cm(widths[j])
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,val in enumerate(row):
            set_cell_text(cells[j], val, False, None, 8.2)
            if widths: cells[j].width = Cm(widths[j])
            if i%2: shade(cells[j], 'F7F9FC')
    doc.add_paragraph()
    return t


def add_bullets(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet')
        p.add_run(item)


def add_num(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Number')
        p.add_run(item)


def add_caption(doc, text):
    p=doc.add_paragraph(style='Caption')
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True


def add_original_figure(doc, filename, caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIG/filename), width=Cm(16.6))
    add_caption(doc, caption + '（GitHub main 分支仓库原图，未重绘）')


def add_callout(doc, title, text, fill='EAF2F8'):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,fill); borders(c,'B4C6E7')
    c.text=''
    p=c.paragraphs[0]
    r=p.add_run(title+'：'); r.bold=True; r.font.name='Arial'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
    r2=p.add_run(text); r2.font.name='Arial'; r2._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
    doc.add_paragraph()


doc=Document()
sec=doc.sections[0]
sec.top_margin=Cm(1.7); sec.bottom_margin=Cm(1.6); sec.left_margin=Cm(1.9); sec.right_margin=Cm(1.8)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(10.2)
normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
normal.paragraph_format.line_spacing=1.30; normal.paragraph_format.space_after=Pt(5)
for sty,size,color in [('Title',25,'1F4E78'),('Heading 1',16,'17365D'),('Heading 2',13,'2F5597'),('Heading 3',11,'1F7A8C')]:
    s=styles[sty]; s.font.name='Arial'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
    s._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
styles['Caption'].font.name='Arial'; styles['Caption'].font.size=Pt(8.5); styles['Caption'].font.color.rgb=RGBColor(79,129,189)
styles['Caption']._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(90)
r=p.add_run('北岸厂 Class C-minDQR'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor(23,54,93); r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('D4 平行冗余时序一致性\n绘图组件顶刊 SCI 发表水平审查'); r.bold=True; r.font.size=Pt(25); r.font.color.rgb=RGBColor(31,78,121); r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('逐图科学审查｜图形语义与分母合同｜主文/补充材料重构｜新增高价值图件方案'); r.font.size=Pt(12); r.font.color.rgb=RGBColor(79,129,189); r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(120)
p.add_run('审查基线：GitHub main（2026-08-12）\n目标期刊：Water Research / Environmental Modelling & Software / Journal of Process Control')
doc.add_page_break()

# Executive

doc.add_heading('执行摘要',1)
doc.add_paragraph('本审查直接针对 GitHub main 分支 D4 `outputs/figures` 的 8 幅正式图件、`src/d4/figures.py`、`figure_style.py`、D4 专家审计和 figure-bundle QA。结论是：当前图件在出版工程规范层面已经较成熟，但作为顶刊主文图组仍存在“科学语义尚未完全锁定、可评价分母不一致、部分图与最终 D4 科学合同脱节、缺少校准证据强度与现场真实性”四类核心问题。换言之，下一轮优化的重点不应继续停留在字体、dpi 或配色微调，而应转向 figure-to-claim contract。')

add_table(doc,['维度','专家评分','判断'],[
('出版技术规范','8.8/10','PDF/SVG/600 dpi PNG、Arial、统一 panel label 和尺寸控制已经达到较高水平。'),
('视觉一致性与可读性','8.0/10','版式统一，但部分 6–7 pt 文字和 0.55 pt 轨迹在印刷缩放后偏细；状态颜色需增加色盲安全冗余编码。'),
('统计信息表达','6.8/10','当前图能够描述结果，但均值、周均值和 mode 聚合对低尾事件、依赖结构和不确定性表达不足。'),
('科学语义准确性','6.3/10','FigD4_2(a)、FigD4_3、FigD4_4 存在可评价分母/原始证据口径不一致；FigD4_1 阴影不是正式 episode window。'),
('机制验证完整性','7.0/10','ROC/PR 与 ablation 具有价值，但缺少 per-pair、方向、common-mode、lag-resolution 和外部现场真值的完整闭环。'),
('顶刊主文叙事成熟度','6.5/10','现有 8 图更像“完整项目输出包”，尚未重构成 5–7 幅主文核心证据图。'),
], [3.8,2.4,10.0])

add_callout(doc,'总判断','当前图组“可以用于项目验收和方法学报告”，但不建议原样全部作为顶刊主文图。最优策略是：保留并重构 FigD4_2、FigD4_5、FigD4_7、FigD4_8；将 FigD4_1、FigD4_3、FigD4_6 下沉至补充材料或重构；FigD4_4 与 FigD4_5 合并；新增“方法框架图、校准证据/基线支持图、现场 case-study 图”，并在具备数据后增加“时间外推/外部验证图”。','FFF2CC')

# Standards

doc.add_heading('1 审查依据与顶刊图件标准',1)
doc.add_heading('1.1 仓库当前图件工程质量',2)
doc.add_paragraph('仓库 QA 记录显示 FigD4_1–FigD4_8 均同时具有 PNG、SVG 和 PDF，8/8 非空、Arial 声明和 panel label 规则通过，并通过 locked-SHA lineage 新鲜度检查。`figure_style.py` 使用 183 mm 双栏宽度、Arial/Helvetica 字体栈、7 pt 基准字号、6.5 pt tick、9 pt panel label、0.8 pt 轴线，并以 PDF/SVG/600 dpi PNG 三格式输出。该基础显著优于通常的“论文后期再手工美化”流程。')
doc.add_heading('1.2 Elsevier/目标期刊层面的关键要求',2)
add_bullets(doc,[
'Elsevier artwork sizing 建议成刊正文文字约 7 pt，脚标/上标不小于约 6 pt；当前 D4 大体符合，但 FigD4_7 legend 使用 6 pt 已处于下限，建议统一提升至 7 pt。',
'Elsevier 对显著曲线建议约 1 pt，绝对最小线宽 0.1 pt；当前 FigD4_1 主轨迹约 0.55 pt，技术上可印刷但作为主证据曲线偏细，建议关键曲线提高至 0.8–1.0 pt。',
'Elsevier 更推荐 PDF/EPS 等矢量格式用于图表；当前已有 PDF/SVG，投稿时应以 PDF（或 EPS）作为主提交矢量文件，PNG 作为审阅/预览，不应仅依赖 600 dpi PNG。',
'Elsevier明确建议色盲安全配色并避免仅以颜色编码类别。当前 FigD4_5 的 green–amber–red 状态图以及 FigD4_2 的多色堆叠应增加纹理、边界、文字或色盲安全调色板。',
'Water Research 强调与水科学/工程问题的应用连接；因此仅有算法性能图不足，必须加入现场 process context 与真实案例证据。Environmental Modelling & Software 更强调开发逻辑、严格测试、比较与可重复性；D4 的 lineage/QA 是优势，但需要更强 calibration provenance、holdout 与 uncertainty 图。Journal of Process Control 更关注动态监测、检测延迟、false-alarm、common-mode rejection 与时序定位，因此当前 change-point lag 能力必须进一步收口。'
])

# Semantic issues overview

doc.add_heading('2 现有绘图组件的跨图共性问题',1)
add_table(doc,['问题','代码层事实','顶刊风险','建议'],[
('可评价分母不统一','FigD4_2(a) 对全部 Q 行求均值；FigD4_2(b) 仅使用 usable_for_D4。FigD4_4 直接对 D4_raw 做周均值而未 mask unusable。','同一文章不同图可能对“D4结果”使用不同 denominator，容易被审稿人质疑。','主文所有结果图默认以 usable_for_D4 为正式分母；原始证据图必须显式标注 raw descriptive evidence。'),
('重叠窗口造成伪精密','24 h window、1 h step，相邻窗口共享约 95.8% 原始时间支持。','常规 box/ROC/CI 若把 pair-hour 当独立样本，会高估信息量。','统计推断使用 day/week block bootstrap；图注报告 block unit，而非仅 n windows。'),
('事件视觉标注与正式事件不一致','FigD4_1 通过 `raw_status_label` 每隔 6 行绘制 6 h 阴影，而不是读取 `D4_event_windows.xlsx`。','阴影可能被误读为正式事件 onset/duration。','直接使用正式 episode start/end；代表性事件增加 onset/recovery/duration。'),
('周聚合语义不严','FigD4_4 使用 `W-MON` 并标“ISO week ending”；FigD4_5 用 `W-MON` period start 并标“ISO week starting”。','ISO 周为 Monday–Sunday；当前标签与 pandas 周期定义不完全一致。','统一改为 W-SUN（周一开始、周日结束）或直接写 Week ending Monday，避免“ISO”误用。'),
('校准可信度没有进入图层','当前图未展示 variable×regime exact/fallback/global mapping scope。','ORP fallback 导致跨 regime score 可比性受 calibration evidence 限制。','新增 calibration support/provenance 图，并在所有 ORP 结果图旁给 fallback/quality 标识。'),
('构念与因果归属易混淆','D4检测 pair asymmetry，但当前结果图缺少 sensor-vs-process truth/context。','Water Research 类期刊会追问低 D4 是否真传感器故障。','增加现场 case-study：residual + D4 + D1/D2 + process log/maintenance truth。')
], [3.0,5.2,4.6,5.0])

# Per figure reviews
reviews = [
('3 FigD4_1 — Paired residual consistency：工程合格，但不宜原样作为主文核心图','FigD4_1_paired_residual_consistency.png',
['7 个 pair 纵向排列、6 h residual median、按 pair robust scale 归一化，适合展示长期双线轨迹的宏观同步性。','同一图直接覆盖 DO11–DO14 和 ORP11–ORP13，能够让读者快速识别 ORP13、DO14 等长期偏离位置。'],
['当前红色阴影不是正式 D4 episode：代码对 `pair_asymmetry` 小时序列 `iloc[::6]` 采样后绘制 6 h span；其视觉面积不能解释为正式 event duration。','6 h median 会压制 spike/短时事件，因此该图不能承担“异常检测灵敏度”证据。','各 pair 使用自己的全期 robust scale，纵轴都是 robust-z，但不应将不同 pair 的视觉幅度直接比较为严重程度。','7 行全期轨迹在印刷尺寸下信息过密，关键异常 onset、recovery、D4 分数与机制不可见。'],
['主文不建议保留当前 7-panel 全期版；转为 Supplementary Fig. S1。','主文重新制作 2–3 个代表性 pair 的“全期概览 + 局部事件 inset”，优先 DO14、ORP13，再选一个稳定 reference pair（DO12/ORP11）。','阴影必须直接来自 `D4_event_windows.xlsx` 的 start_ts/end_ts；同时标注 episode minimum/mean D4、dominant evidence。','局部 inset 使用原始 10-min residual 或 1-h median，避免 6 h median 隐去关键结构；必要时叠加 D1/D2/process context。']),
('4 FigD4_2 — Subscore contribution：有核心论文价值，但需要重构统计口径','FigD4_2_subscore_contribution.png',
['能够揭示低 D4 的机制异质性，是 D4 相比单一 correlation score 最具论文价值的图之一。','Panel (b) 使用 usable_for_D4 的 D4_total 箱线图，正式分母逻辑正确。'],
['Panel (a) `main.groupby(pair_id)[Q_cols].mean()` 未过滤 usable_for_D4，而 panel (b) 过滤了 usable_for_D4；两面板 denominator 不一致。','Panel (a) 展示的是加权贡献到 D4_base，不是最终 D4_raw；最终 raw 还包含 25% minimum-subscore penalty，当前视觉容易让读者误以为堆叠总和就是最终 D4。','权重在绘图函数中硬编码为 [0.35,0.25,0.20,0.20]，虽然当前与 config 一致，但不是单一配置真源。','均值柱图不展示重叠时间窗下的不确定性，也看不到 calibration_quality。'],
['将 panel (a) 限定为 usable_for_D4，权重从 cfg.weights 读取。','优先将“加权堆叠均值”改为 pair×subscore 的 point-range / raincloud：中心为 median/mean，CI 为 week-block bootstrap；单独用黑点显示 mean D4_raw。','如果保留堆叠，标题必须写“weighted contribution to D4_base”，并增加 `minimum-subscore penalty` 单独符号/旁图。','ORP12/ORP13 用边框或上方 strip 标注 calibration_quality / fallback share。','建议保留为主文核心图（重构后）。']),
('5 FigD4_3 — Trend-slope scatter：适合作为补充验证，不足以承担主结论','FigD4_3_trend_slope_scatter.png',
['DO/ORP 分面、等比例 1:1 参考线和 log-density hexbin 的视觉规范较好。','能够证明 trend inconsistency 并非所有低 D4 pair 的主导机制。'],
['数据来源为 raw detector outputs，未与 usable_for_D4 显式联结；因此包含的 slope window 不一定都是正式 D4 可评价窗口。','把所有 DO pairs 和所有 ORP pairs 各自汇总到一张 hexbin，会掩盖 pair-specific heterogeneity。','图中没有 concordance coefficient、robust regression、median |Δslope| 或 block-bootstrap CI，视觉接近 1:1 不能替代定量一致性统计。','24 h/1 h 重叠窗口导致密度非常高，bin count 不应被视为独立样本量。'],
['下沉至 Supplementary Fig. S2。','若希望进入主文，改为 7 个 pair 的 concordance forest plot：CCC/Spearman、median normalized slope difference、block-bootstrap 95% CI；当前 hexbin作为补充原始分布。','明确仅使用 usable_for_D4 或在标题写 raw descriptive windows。','色条上限不建议固定 1500，应使用数据驱动或标注 clipped bins。']),
('6 FigD4_4 — Weekly D4_raw heatmap：当前存在重要分母语义问题，应重做','FigD4_4_raw_score_heatmap.png',
['热图非常适合呈现 pair-level temporal persistence，是长时序 D4 结果中最直观的“时空负担图”。'],
['代码直接对 `D4_raw` 周均值进行聚合，没有使用 `D4_total` 或 `usable_for_D4` mask；因此 D2 veto / not-evaluable 窗口只要仍能计算 raw risk，就可能进入周均值。','weekly mean 会显著稀释短时但严重的 D4<3 episodes；均值也不直接对应 classification threshold。','连续 viridis 1–5 没有突出 D4=3 的科学决策边界，也没有明确 not-evaluable 空白。','`W-MON` 被称为 ISO week ending，术语不够严谨。'],
['主文建议保留“热图思想”但彻底重构：使用 `D4_total` 或 `D4_raw.where(usable_for_D4)`。','主面板改为每周 `P(D4_raw<3 | usable)` 或 q05/q25 burden，而不是 mean；阈值 3 直接进入图形语义。','not-evaluable 用灰色/NA 显示，并在同一行附 coverage。','与 FigD4_5 合并为“asymmetry burden + evaluability + calibration quality”三联图。','周定义统一为 W-SUN 或去掉 ISO 字样。']),
('7 FigD4_5 — Status/evaluability：科学角色重要，但 weekly mode 不够顶刊','FigD4_5_status_evaluability.png',
['把 status 与 evaluability 分开是非常正确的 DQR 构念设计，尤其能解释 DO14“高可评价但低一致性”和 ORP13“低一致性且证据覆盖不足”。','Panel (b) evaluable fraction 是主文应保留的重要证据。'],
['Panel (a) 对每周 hourly status 取 mode：如果 49% 为 asymmetry、51% 为 borderline，整周只显示 borderline，会压缩重要低分负担。','tie 时 `mode().iloc[0]` 的选择没有科学优先级。','green–amber–red 主要依赖色彩类别，色盲可访问性不足。','周标签存在 W-MON / ISO starting 的语义问题。'],
['用 weekly asymmetry fraction 替代 mode；必要时同时显示 borderline fraction。','与 FigD4_4 合并，Panel a=low-score burden，Panel b=evaluable fraction，Panel c=variable×regime calibration-quality/fallback share。','使用色盲安全 palette，并用灰色明确 NA。','该合并图应成为主文最核心的现场长期结果图之一。']),
('8 FigD4_6 — Context independence：作为审计图有价值，但应退出主文结果核心','FigD4_6_context_independence.png',
['能够证明 D4_raw 与 D1 解释层分离，体现“观测事实不被上游先验重写”的科学合同。'],
['Panel (a) 使用历史 `D4_forDQR_provisional`，而最终科学合同已经明确 D4_final = D4_raw；主文继续突出 provisional 容易让读者误以为 D1 fuse 仍是正式评分环节。','Panel (b) log count 的四类 D1 state 主要属于软件审计/接口治理证据，不是 D4 科学机制结果。','图占用一整幅主文 figure，但提供的新增水科学信息有限。'],
['下沉到 Supplementary / Methods audit。','若保留，改成“final integrity check”：展示 Δ=D4_final−D4_raw 的分布/最大绝对差=0，并用简洁 contract schematic 说明 D1/D2/D5 角色。','主文可用一张方法框架图替代当前 FigD4_6，并将 numeric independence 写入 Methods/QA table。']),
('9 FigD4_7 — ROC/PR：现有图中最接近顶刊主文水平，但验证证据仍需升级','FigD4_7_validation_roc_pr.png',
['drift/step AUROC≈0.93–0.94、freeze≈0.80、spike≈0.56，机制边界清晰且没有掩饰 spike 弱项。','ROC 与 PR 同时给出，比只报告 ROC 更适合 fault detection 评价。'],
['ROC legend 给 AUC，但 PR legend 未给 AUPRC；图上无 95% CI ribbon/point interval。','当前验证属于内部 injection/chronological stress test，不是 independently adjudicated field accuracy。','未在该图中直接表现 equal common-process FAR=0.041 这一最能证明“pair asymmetry而非all-change detector”的关键 specificity 结果。','尚未展示 target/peer direction、per-pair heterogeneity、terminal/future holdout。','PR chance line固定为0.5必须与实际正例 prevalence 严格一致并在图注说明。'],
['保留为主文核心图并升级为 3–4 panels：ROC、PR、common-mode FAR/positive-control response、per-pair/target-peer forest。','曲线附 block-bootstrap 95% CI；legend 同时报告 AUROC/AUPRC。','将 spike 明确标为“secondary/non-target mechanism”，避免读者把 0.56 视为整体算法失败。','开发期 mapping 完全冻结后，再在 terminal/future block 重画最终性能图；Water Research 投稿还需要现场事件 truth。']),
('10 FigD4_8 — Ablation：具有方法学价值，但 pooled metric 需要机制化','FigD4_8_validation_ablation.png',
['同时考察 unilateral discrimination 与 synchronous FAR，是比单一 AUC 更成熟的 multi-objective ablation。','当前已经给 CI error bars，是现有 8 图中统计表达较规范的一幅。'],
['pooled unilateral AUC 会掩盖 drift/step/freeze/spike 对各组件的不同依赖。','no_var/no_cp 可能出现 pooled AUC 略高于 full，若无机制分层很容易被审稿人解释为组件冗余。','deadband 的真正作用是 low-dynamic process-floor 保护，而现有通用 injection mixture 并不能充分检验它。','v2.0 addendum 已显示当前 hourly change-point 对 0/10/30/60/180 min lag 的 severity monotonicity 为 0，现图未把这一关键分辨率边界显式呈现。'],
['保留主文但改为“机制×组件”热图/forest：每个 fault mechanism 分别显示 ΔAUROC/ΔAUPRC 相对 full；common-mode 单独显示 ΔFAR。','增加 low-dynamic DO/ORP negative-control panel 来证明 deadband 的必要性。','增加 change-point lag-response panel；若未解决 10–30 min 分辨率，则在图中明确标注 resolution floor，或将 Q_cp 解释降级为 coarse structural-change evidence。'])
]

for title,fn,strengths,issues,recs in reviews:
    doc.add_heading(title,1)
    add_original_figure(doc,fn,title.split('—')[0].strip())
    doc.add_heading('现有优势',2); add_bullets(doc,strengths)
    doc.add_heading('关键科学/绘图问题',2); add_bullets(doc,issues)
    doc.add_heading('建议修改',2); add_num(doc,recs)

# New figures

doc.add_heading('11 建议新增的“真正有论文价值”的图',1)
doc.add_paragraph('下列新增图不是为了增加图数，而是填补现有 8 图无法回答的顶刊审稿问题。优先级按“没有这张图，核心 claim 是否容易被质疑”排序。')
add_table(doc,['优先级','建议新增图','要回答的科学问题','核心面板设计','主文价值'],[
('A1','D4 方法与证据链框架图','D4究竟测什么？四类风险、benchmark、fallback、D1/D2/D5如何解耦？','de-periodised residual → 24 h homologous pair window → W1/KS、Theil–Sen、IQR、CP → quantile mapping → D4_raw；旁路标注 D1 admission、D2 observability、D5 attribution。','极高：目前缺少能让审稿人一页理解 D4 构念的主图。'),
('A2','Calibration evidence / provenance 图','ORP 分数的 regime-specific calibration 到底有多强？哪些依赖 fallback？','variable×regime heatmap：exact_stratum_size；第二面板 mapping_scope/quality；第三面板 bootstrap q90/q97.5 uncertainty 或 score-migration。','极高：直接回应 ORP sparse support 和跨 regime 可比性问题。'),
('A3','Field case-study 图','低 D4 是真实过程差异还是 sensor-specific anomaly？','DO14、ORP13 代表性 episode：raw/residual、四 Q、D4_raw、D1/D2、process log、maintenance/reference truth；标 onset/recovery。','极高：Water Research 尤其需要，当前内部注入不能替代现场真实性。'),
('A4','Asymmetry burden–evaluability–calibration 三联图','长期低分是否持续？证据覆盖是否足够？阈值是否来自 adequate calibration？','合并并替换当前 FigD4_4/5：weekly low-score fraction、evaluable fraction、fallback/quality fraction。','极高：成为现场结果的主图。'),
('A5','Mechanism specificity 图','D4能否区分 unilateral 与 common-mode？方向是否可定位？','target/peer drift/step/freeze forest + equal common FAR + unequal/opposite positive controls。','高：比单纯 ROC 曲线更贴合 D4 的“asymmetry detector”构念。'),
('B1','Event persistence / recurrence 图','ORP13/DO14 的低分到底是多次短事件还是少数长事件？','episode duration ECDF/raincloud + event count + total exposure hours；可按 pair/variable 分层。','高：把周热图视觉印象转成正式 event-level evidence。'),
('B2','Temporal holdout / external validation 图','性能是否仅对开发期有效？','development calibration frozen → terminal test → future period / second plant；AUROC/AUPRC/FAR forest。','顶刊决定性：具备新数据后应进入主文。'),
('B3','Change-point lag resolution 图','Q_cp 能定量到什么时间尺度？','true lag vs detected timing difference / anomaly severity；0、≥1 h 等有效分辨率区间。','JPC高价值；若当前 rho=0，则先作为 limitation/sensitivity。'),
('C1','Downstream utility 图','D4 作为 DQR 子维度是否真正改善后续建模？','按 D4 strata 选择/加权数据，比较 SUMO/软测量/预测模型参数稳定性与 test error。','全DQR论文价值极高；若D4独立方法短文则可放补充或后续研究。')
], [1.4,3.5,4.3,6.4,2.3])

# Recommended manuscript architecture

doc.add_heading('12 推荐的最终主文图组架构',1)
add_table(doc,['建议主文图号','来源','建议内容','当前动作'],[
('Fig.1','新增','D4 scientific construct + calibration/observability/attribution contract schematic','必须新增'),
('Fig.2','重构 FigD4_2','pair-level D4 profile + mechanism-resolved subscore with block-bootstrap CI','保留但重绘'),
('Fig.3','合并 FigD4_4 + FigD4_5 + 新增 calibration strip','weekly asymmetry burden + evaluability + calibration quality','核心重构'),
('Fig.4','新增','DO14/ORP13 field case studies + process/maintenance truth','强烈建议新增'),
('Fig.5','升级 FigD4_7','ROC/PR + common-mode FAR + target/peer direction/per-pair uncertainty','保留升级'),
('Fig.6','升级 FigD4_8','mechanism-specific ablation + low-dynamic negative control + lag-resolution','保留升级'),
('Fig.7','新增（有数据后）','terminal/future/second-plant validation 或 downstream utility','顶刊增强')
], [2.2,4.1,7.4,3.8])

doc.add_heading('13 建议下沉至 Supplementary Information 的图',1)
add_table(doc,['补充图','来源','理由'],[
('Fig. S1','当前 FigD4_1 全 7 pair 长期轨迹','信息完整但太密；主文应只展示代表 case。'),
('Fig. S2','当前 FigD4_3 slope hexbin','作为原始 trend distribution 支撑有价值，但主文需用 effect size/CI。'),
('Fig. S3','当前 FigD4_6 independence diagnostic','属于接口/审计证明，不宜占主文核心结果位。'),
('Fig. S4','完整 pair×subscore distributions','支撑 Fig.2 的汇总指标与审稿复核。'),
('Fig. S5','window length / threshold / event-min-duration sensitivity','回答参数依赖。'),
('Fig. S6','calibration fallback / hierarchical shrinkage sensitivity','ORP support 不足时的稳健性。')
], [2.4,4.0,9.0])

# Journal-specific

doc.add_heading('14 不同期刊的图件策略',1)
add_table(doc,['期刊','图件必须强化的证据','当前 D4 的匹配度','投稿前最低补充'],[
('Water Research','现场水处理意义、真实工艺事件、传感器/过程因果边界、对运行/数据可信度的实际影响','中等','Field case-study + process context；主文弱化纯软件审计图。'),
('Environmental Modelling & Software','构念清晰、校准/不确定性、严格测试、比较、可重复性和数据/代码 lineage','较高','Calibration provenance + temporal holdout + denominator contract + sensitivity。'),
('Journal of Process Control','动态监测、检测延迟、common-mode rejection、方向识别、change-point resolution','中高','Lag-resolution、detection delay、common-process controls；Q_cp 当前能力需先收口。')
], [3.2,7.2,3.4,5.2])

# Priority work packages

doc.add_heading('15 建议的绘图整改优先级',1)
add_table(doc,['等级','整改事项','原因'],[
('P0','统一 denominator：FigD4_2(a)、FigD4_3、FigD4_4 明确/改为 usable_for_D4 正式口径','这是科学语义问题，不是美观问题；应在任何投稿前优先修复。'),
('P0','修正 FigD4_1 event shading 为正式 event windows；修正 W-MON/ISO 周标签','避免图形与事件/时间定义合同不一致。'),
('P0','新增 calibration-quality/provenance 图，并把 ORP fallback 进入图注/结果解释','直接关系 ORP 跨 regime 评分可比性。'),
('P1','合并 FigD4_4/5 为 burden–coverage–calibration 主图','减少冗余并提高信息密度。'),
('P1','升级 FigD4_7/8：CI、AUPRC、common-mode、mechanism-specific、per-pair','形成顶刊级验证证据链。'),
('P1','新增 DO14/ORP13 field case-study','Water Research 类应用顶刊不可替代。'),
('P2','将 FigD4_3/6 下沉 supplement；当前 FigD4_1 全轨迹下沉 supplement','让主文围绕核心 claim，而非完整项目输出。'),
('P2','统一色盲安全、7 pt 文字、0.8–1.0 pt 主线；投稿用 PDF/EPS 矢量','出版规范收尾。')
], [1.5,8.8,9.2])

# Detailed code findings

doc.add_heading('16 绘图代码层面的具体修改建议',1)
code_rows=[
('figure_m1','事件阴影','不要 `events["timestamp"].iloc[::6] + 6h`；读取 D4_event_windows 的 start_ts/end_ts。'),
('figure_m2','分母','`means` 应基于 `main[main.usable_for_D4]`；panel a/b 使用相同正式 denominator。'),
('figure_m2','权重','不要硬编码 weights；从 cfg.weights 注入绘图函数。'),
('figure_m3','样本口径','raw detector output 与 main_scores 通过 timestamp+pair_id join，按 usable_for_D4 过滤或标 raw descriptive。'),
('figure_d1','热图分母','将 D4_raw 替换为 D4_total 或 `D4_raw.where(usable_for_D4)`；NA 显式灰色。'),
('figure_d1','聚合','优先 weekly low-score fraction/q05/q25，不用 mean 作为唯一长期指标。'),
('figure_d1/d2','周定义','若写 ISO week，使用 Monday–Sunday（如 W-SUN）并核对标签；否则直接写 week ending Monday。'),
('figure_d2','状态','用 weekly asymmetry fraction 代替 mode；coverage 保留。'),
('figure_d3','最终合同','主文改成 final D4_raw protection/integrity，历史 provisional 移 supplement。'),
('figure_v1','验证统计','AUPRC、prevalence、bootstrap CI、common-mode controls、target/peer/per-pair。'),
('figure_v2','消融','从 raw metric bar 改为 Δmetric relative-to-full + CI；按 mechanism 分层。'),
('figure_style','可访问性','采用色盲安全 palette；重要类别同时使用线型/纹理/符号。'),
('figure_style','印刷','legend/tick 尽量 ≥7 pt，关键曲线 0.8–1.0 pt；保留 PDF vector 作为投稿源文件。'),
]
add_table(doc,['函数','问题域','修改'],code_rows,[3.0,3.1,12.2])

# Final decision

doc.add_heading('17 最终专家结论',1)
doc.add_paragraph('D4 当前绘图组件已经具备较好的 publication engineering：统一尺寸、字体、panel 标识、矢量/高分辨率双轨输出以及 SHA lineage QA 都是明显优势。真正限制顶刊发表的不是“图不漂亮”，而是部分图形仍没有严格锁定正式可评价分母和最终 D4 科学合同，且缺少 calibration evidence strength、field truth 与 external validity 三类决定性证据。')
doc.add_paragraph('因此，不建议继续在现有 8 图基础上简单“增加字体、换配色、加注释”后直接投稿。应进行一次以 scientific claims 为中心的主文重构：将现有 8 图压缩/重组为 4 个有效主图来源（FigD4_2、FigD4_4/5、FigD4_7、FigD4_8），新增方法框架、calibration provenance 和 field case-study；将完整 residual trajectories、slope density 和 D1-context independence 下沉补充材料。完成 P0/P1 后，D4 图组才真正接近 Environmental Modelling & Software 的方法学严谨性要求；若再补足现场 case truth，则可进一步满足 Water Research 的应用水科学证据要求；若希望投 Journal of Process Control，则 change-point lag/detection-delay 的时间分辨率必须先形成可量化的有效区间。')

# References

doc.add_heading('参考依据与建议文献',1)
refs=[
'GitHub D4 formal figure bundle, `outputs/figures`, main branch, accessed 2026-08-12.',
'GitHub D4 `src/d4/figures.py`, `figure_style.py`, `outputs/qa/figure_bundle_audit.json`, and `docs/D4_EXPERT_AUDIT_2026-07.md`, main branch.',
'Elsevier. Artwork sizing. Current author policy. Recommended finished lettering ~7 pt, subscripts/superscripts no smaller than ~6 pt.',
'Elsevier. Artwork types / Artwork FAQs. Vector PDF/EPS preferred for graphs; recommended prominent line weights approximately 1 pt; 300/500/1000 dpi guidance for halftone/combination/line art.',
'Elsevier. Artwork and media instructions. Guidance on color-vision accessibility and redundant visual encoding.',
'Water Research journal scope (Elsevier): emphasizes applied relevance to the anthropogenic water cycle and clear water-research connection.',
'Environmental Modelling & Software journal scope (Elsevier): emphasizes rigorous model/software development, substantial testing/evaluation, comparison and credibility.',
'Spindler A. Structural redundancy of data from wastewater treatment systems. Water Research. 2014;57:193–201. doi:10.1016/j.watres.2014.03.042.',
'Villez K, Vanrolleghem PA, Corominas L. Optimal flow sensor placement on wastewater treatment plants. Water Research. 2016;101:75–83. doi:10.1016/j.watres.2016.05.068.',
'Schneider MY et al. Beyond signal quality: The value of unmaintained pH, dissolved oxygen, and oxidation–reduction potential sensors for remote performance monitoring. Water Research. 2019;161:639–651. doi:10.1016/j.watres.2019.06.007.',
'Ba-Alawi AH et al. Intelligent sensor validation for sustainable influent quality monitoring in wastewater treatment plants using stacked denoising autoencoders. Journal of Water Process Engineering. 2021;43:102206. doi:10.1016/j.jwpe.2021.102206.',
'Künsch HR. The Jackknife and the Bootstrap for General Stationary Observations. Annals of Statistics. 1989;17(3):1217–1241. doi:10.1214/aos/1176347265.',
'Killick R, Fearnhead P, Eckley IA. Optimal Detection of Changepoints With a Linear Computational Cost. JASA. 2012;107(500):1590–1598. doi:10.1080/01621459.2012.737745.'
]
for i,ref in enumerate(refs,1):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.3); p.paragraph_format.first_line_indent=Cm(-0.3)
    p.add_run(f'{i}. {ref}')

# Header/footer
for section in doc.sections:
    hp=section.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=hp.add_run('D4 Parallel-redundancy Temporal Consistency｜Figure Publication Audit'); rr.font.size=Pt(8); rr.font.color.rgb=RGBColor(100,100,100)
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

# Prevent figure rows/images from being split weirdly is handled by image paragraphs; save.
doc.save(DOCX)
print(DOCX)
